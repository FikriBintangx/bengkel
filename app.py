import os
import re
import time
from docx import Document
from flask import Flask, request, render_template_string, send_file, redirect, url_for, jsonify
from google import genai
from google.genai import types as genai_types

from pypdf import PdfReader, PdfWriter

import random as _random
import builtins

def print(*args, **kwargs):
    try:
        builtins.print(*args, **kwargs)
    except OSError:
        pass


# Setiap kata punya DAFTAR sinonim → dipilih acak per kemunculan
# None dalam list = pertahankan kata asli (agar tidak 100% selalu diganti → lebih manusiawi)
INDONESIAN_SYNONYMS_MULTI = {

    # Verba
    "menggunakan":     ["memakai", "memanfaatkan", "menerapkan"],
    "memakai":         ["menggunakan", "memanfaatkan"],
    "menunjukkan":     ["memperlihatkan", "mengindikasikan", "menandakan"],
    "memperlihatkan":  ["menunjukkan", "mengindikasikan"],
    "membantu":        [None, "mempermudah", "mendukung", "menunjang"],
    "mempermudah":     [None, "membantu", "mendukung"],
    "melakukan":       ["menjalankan", "melaksanakan", "mengerjakan"],
    "menjalankan":     ["melakukan", "melaksanakan", "mengerjakan"],
    "melaksanakan":    ["melakukan", "menjalankan", "mengerjakan"],
    "menyatakan":      ["mengungkapkan", "mengemukakan", "mengutarakan"],
    "mengungkapkan":   ["menyatakan", "mengemukakan", "mengutarakan"],
    "mengemukakan":    ["menyatakan", "mengungkapkan", "mengutarakan"],
    "menjelaskan":     ["menerangkan", "menguraikan", "mendeskripsikan", "memaparkan"],
    "menerangkan":     ["menjelaskan", "menguraikan", "memaparkan"],
    "menguraikan":     ["menjelaskan", "menerangkan", "memaparkan"],
    "diperlukan":      ["dibutuhkan", "diharuskan", "dipersyaratkan"],
    "dibutuhkan":      ["diperlukan", "diharuskan", "dipersyaratkan"],
    "meningkatkan":    ["memaksimalkan", "mengoptimalkan", "memperbaiki", "mendongkrak"],
    "memaksimalkan":   ["meningkatkan", "mengoptimalkan"],
    "merancang":       ["mendesain", "membangun", "membuat"],
    "mendesain":       ["merancang", "membangun", "membuat"],
    "membuat":         ["menyusun", "merancang", "membangun"],
    "menyusun":        ["membuat", "merancang", "membangun"],
    "mengatasi":       ["menyelesaikan", "memecahkan", "mengurai"],
    "menyelesaikan":   ["mengatasi", "memecahkan", "menuntaskan"],
    "memecahkan":      ["mengatasi", "menyelesaikan", "menuntaskan"],
    "ditemukan":       ["didapatkan", "diperoleh", "dihasilkan"],
    "didapatkan":      ["ditemukan", "diperoleh", "dihasilkan"],
    "diperoleh":       ["didapatkan", "ditemukan", "dihasilkan"],
    "dikembangkan":    ["dibangun", "dirancang", "dibuat"],
    "dibangun":        ["dikembangkan", "dirancang", "dibuat"],
    "digunakan":       ["dimanfaatkan", "diterapkan", "dipakai"],
    "dimanfaatkan":    ["digunakan", "diterapkan", "dipakai"],
    "diterapkan":      ["digunakan", "dimanfaatkan", "dipakai"],
    "mengidentifikasi":["mengenali", "mendeteksi", "menetapkan"],
    "mengenali":       ["mengidentifikasi", "mendeteksi"],
    "menganalisis":    ["mengkaji", "menelaah", "mengevaluasi"],
    "mengkaji":        ["menganalisis", "menelaah", "mengevaluasi"],
    "menelaah":        ["menganalisis", "mengkaji", "mengevaluasi"],
    "membahas":        ["mengulas", "mendiskusikan", "mengkaji"],
    "mengulas":        ["membahas", "mendiskusikan", "mengkaji"],
    "menerapkan":      ["mengimplementasikan", "menggunakan", "memanfaatkan"],
    "mengimplementasikan": ["menerapkan", "menggunakan", "memanfaatkan"],
    "mengolah":        ["memproses", "mengelola", "mengerjakan"],
    "memproses":       ["mengolah", "mengelola", "mengerjakan"],
    "memperoleh":      ["mendapatkan", "mendapat", "menerima"],
    "mendapatkan":     ["memperoleh", "mendapat", "menerima"],
    "menghasilkan":    ["memproduksi", "menghasilkan", "menciptakan"],
    "memproduksi":     ["menghasilkan", "menciptakan"],
    "menentukan":      ["menetapkan", "memutuskan", "mengidentifikasi"],
    "menetapkan":      ["menentukan", "memutuskan"],
    "mengintegrasikan":["menggabungkan", "menyatukan", "memadukan"],
    "menggabungkan":   ["mengintegrasikan", "menyatukan", "memadukan"],
    "memvalidasi":     ["memverifikasi", "menguji", "mengecek"],
    "memverifikasi":   ["memvalidasi", "menguji", "mengecek"],
    "menguji":         ["memvalidasi", "memverifikasi", "mengecek"],
    "mengecek":        ["menguji", "memvalidasi", "memverifikasi"],
    "mendukung":       ["menunjang", "membantu"],
    "menunjang":       ["mendukung", "membantu"],
    # Kata hubung & keterangan waktu (sering luput, padahal banyak di skripsi)
    "karena":          ["sebab", "dikarenakan", "lantaran"],
    "sebab":           ["karena", "dikarenakan", "lantaran"],
    "sehingga":        ["sehingga", "akibatnya", "maka"],
    "maka":            ["oleh karena itu", "sehingga", "dengan demikian"],
    "ketika":          ["saat", "tatkala", "sewaktu", "pada saat"],
    "saat":            ["ketika", "tatkala", "sewaktu"],
    "setelah":         ["sesudah", "usai", "pasca"],
    "sesudah":         ["setelah", "usai", "selesai"],
    "sebelum":         ["mendahului", "prior to", "sebelumnya"],
    "selama":          ["sepanjang", "dalam kurun", "semasa"],
    "hingga":          ["sampai", "sampai dengan", "sampai pada"],
    "sampai":          ["hingga", "sampai dengan"],
    "sejak":           ["semenjak", "mulai dari"],
    "kemudian":        ["selanjutnya", "lalu", "berikutnya"],
    "selanjutnya":     ["kemudian", "lalu", "berikutnya"],
    "lalu":            ["kemudian", "selanjutnya", "berikutnya"],
    "juga":            ["pula", "turut", "ikut"],
    "pula":            ["juga", "turut"],
    "bahkan":          ["malahan", "justru", "lebih jauh lagi"],
    "malahan":         ["bahkan", "justru"],
    "justru":          ["bahkan", "malahan", "sebaliknya"],
    "walaupun":        ["meskipun", "biarpun", "kendati"],
    "meskipun":        ["walaupun", "biarpun", "kendati"],
    "biarpun":         ["meskipun", "walaupun", "kendati"],
    "agar":            ["supaya", "guna", "demi"],
    "supaya":          ["agar", "guna", "demi"],
    "apabila":         ["jika", "bila", "kalau", "andaikata"],
    "jika":            ["apabila", "bila", "kalau"],
    "bila":            ["apabila", "jika", "kalau"],
    "kalau":           [None, "jika", "apabila"],
    "namun":           ["akan tetapi", "meskipun demikian", "kendati demikian", "hanya saja"],
    "tetapi":          ["namun", "akan tetapi", "meskipun demikian"],
    "bahwa":           [None, None, None],  # lindungi kata gramatikal penting
    # Kata keterangan tempat & derajat
    "sangat":          ["amat", "begitu", "cukup", "sungguh"],
    "amat":            ["sangat", "begitu", "sungguh"],
    "cukup":           [None, "lumayan", "relatif"],
    "relatif":         [None, "cukup", "agak"],
    "agak":            [None, "relatif", "sedikit"],
    "lebih":           [None, "lebih"],
    "paling":          [None, "ter-", "paling"],
    "banyak":          ["sejumlah", "berbagai", "beragam"],
    "berbagai":        ["beragam", "bermacam-macam", "sejumlah"],
    "beberapa":        ["sejumlah", "sebagian", "berbagai"],
    "sejumlah":        ["beberapa", "berbagai", "beragam"],
    "semua":           ["seluruh", "keseluruhan", "segenap"],
    "seluruh":         ["semua", "keseluruhan", "segenap"],
    "setiap":          ["masing-masing", "tiap", "tiap-tiap"],
    "masing-masing":   ["setiap", "tiap", "tiap-tiap"],
    "tiap":            ["setiap", "masing-masing"],
    "hanya":           [None, "saja", "semata"],
    "saja":            [None, "hanya", "semata"],
    "sudah":           ["telah", "sudah"],
    "telah":           ["sudah", "telah"],
    "sedang":          [None, "tengah", "dalam proses"],
    "tengah":          [None, "sedang"],
    "akan":            [None, "hendak", "bakal"],
    "hendak":          [None, "akan", "bakal"],
    "perlu":           [None, "butuh", "memerlukan"],
    "butuh":           [None, "perlu", "memerlukan"],
    "baru":            [None, "anyar", "mutakhir"],
    "lama":            [None, "lama", "lawas"],
    "besar":           ["besar", "luas", "signifikan"],
    "kecil":           [None, "minor", "terbatas"],
    "tinggi":          [None, "tinggi", "tingkat tinggi"],
    "rendah":          [None, "rendah", "minim"],
    "baik":            [None, "bagus", "unggul"],
    "buruk":           [None, "kurang baik", "tidak optimal"],
    # Nomina
    "metode":          ["pendekatan", "cara", "teknik"],
    "pendekatan":      ["metode", "cara", "teknik"],
    "cara":            ["metode", "pendekatan", "teknik"],
    "teknik":          ["metode", "pendekatan", "cara"],
    "adalah":          ["merupakan", "ialah", "yakni"],
    "merupakan":       ["adalah", "ialah", "yakni"],
    "ialah":           ["adalah", "merupakan", "yakni"],
    "yakni":           ["adalah", "merupakan", "yaitu"],
    "yaitu":           ["yakni", "adalah", "merupakan"],
    "berbasis":        ["berlandaskan", "berbasiskan", "berdasarkan"],
    "berlandaskan":    ["berbasis", "berdasarkan", "bertumpu pada"],
    "berdasarkan":     ["berlandaskan", "mengacu pada", "berpedoman pada"],
    "sangat":          ["amat", "begitu", "sungguh"],
    "amat":            ["sangat", "begitu", "sungguh"],
    # -- KATA DENGAN PILIHAN NATURAL (None = pertahankan kata asli) --
    "dapat":           [None, None, "bisa", "mampu"],   # 50% tetap "dapat"
    "bisa":            [None, "dapat", "mampu"],
    "mampu":           [None, "dapat", "bisa"],
    "analisis":        [None, "kajian", "evaluasi", "penilaian"],
    "kajian":          [None, "analisis", "evaluasi"],
    "telaah":          [None, "analisis", "kajian", "evaluasi"],
    "evaluasi":        [None, "analisis", "kajian", "penilaian"],
    "implementasi":    ["penerapan", "realisasi", "eksekusi"],
    "penerapan":       ["implementasi", "realisasi", "eksekusi"],
    "realisasi":       ["implementasi", "penerapan", "eksekusi"],
    "proses":          ["tahapan", "langkah", "prosedur"],
    "tahapan":         ["proses", "langkah", "prosedur"],
    "langkah":         ["tahapan", "proses", "prosedur"],
    "prosedur":        ["proses", "tahapan", "langkah"],
    "tujuan":          ["sasaran", "target", "maksud"],
    "sasaran":         ["tujuan", "target", "maksud"],
    "target":          ["tujuan", "sasaran", "maksud"],
    "namun":           ["akan tetapi", "meskipun demikian", "kendati demikian"],
    "tetapi":          ["namun", "akan tetapi", "meskipun demikian"],
    "tentang":         [None, "mengenai", "terkait"],  # hapus "perihal" tidak natural
    "mengenai":        [None, "tentang", "terkait"],
    "terkait":         [None, "tentang", "mengenai"],
    "perihal":         ["tentang", "mengenai", "terkait"],  # reverse: normalkan kembali
    "cepat":           ["pesat", "cepat", "sigap"],
    "pesat":           ["cepat", "sigap"],
    "mudah":           ["efisien", "praktis", "sederhana"],
    "pengembangan":    ["perancangan", "pembangunan", "pembuatan"],
    "perancangan":     ["pengembangan", "pembangunan", "pembuatan"],
    "pembangunan":     ["pengembangan", "perancangan", "pembuatan"],
    "penting":         ["krusial", "vital", "signifikan"],
    "krusial":         ["penting", "vital", "signifikan"],
    "vital":           ["penting", "krusial", "signifikan"],
    "signifikan":      ["penting", "krusial", "vital"],
    "perkembangan":    ["kemajuan", "peningkatan", "progres"],
    "kemajuan":        ["perkembangan", "peningkatan", "progres"],
    "peningkatan":     ["kemajuan", "perkembangan", "progres"],
    "efektif":         ["optimal", "tepat guna", "manjur"],
    "optimal":         ["efektif", "tepat guna"],
    "efisien":         ["hemat", "tepat guna", "praktis"],
    "akurat":          ["presisi", "tepat", "cermat"],
    "presisi":         ["akurat", "tepat", "cermat"],
    "tepat":           ["akurat", "presisi", "cermat"],
    "cermat":          ["akurat", "presisi", "tepat"],
    "otomatis":        ["secara otomatis", "mandiri", "otonom"],
    "mengurangi":      ["meminimalisir", "menekan", "memangkas"],
    "meminimalisir":   ["mengurangi", "menekan", "memangkas"],
    "kesalahan":       ["kekeliruan", "error"],           # 'galat' dihapus - terlalu teknis/kaku
    "kekeliruan":      ["kesalahan", "error"],
    "data":            [None, None, None, None, "informasi"],   # 'catatan' dihapus sebagai sinonim data
    "informasi":       [None, None, None, "data"],              # 'keterangan' dihapus sebagai sinonim informasi
    "pengguna":        [None, "pemakai", "user"],
    "pemakai":         [None, "pengguna", "user"],
    "kinerja":         [None, "performa", "prestasi"],
    "performa":        [None, "kinerja", "prestasi"],
    "laporan":         [None, None, "dokumen", "catatan", "rekap"],
    "rekap":           [None, "laporan", "dokumen", "catatan"],
    # 'hambatan' ≠ 'masalah' (beda makna: obstacle vs problem) → pisahkan
    "masalah":         [None, None, "permasalahan", "kendala"],
    "permasalahan":    [None, "masalah", "kendala"],
    "kendala":         [None, "masalah", "permasalahan"],
    "hambatan":        [None, None],   # pertahankan 'hambatan' as-is, jangan ganti ke masalah
    "hasil":           [None, None, "temuan", "output"],  # 50% tetap "hasil"
    "temuan":          [None, "hasil"],
    "luaran":          ["hasil", "output"],    # reverse: normalkan kembali
    "keluaran":        [None, "hasil", "output"],
    "pengujian":       ["uji coba", "tes", "verifikasi"],
    "penelitian":      ["studi", "riset", "kajian"],
    "studi":           ["penelitian", "riset", "kajian"],
    "riset":           ["penelitian", "studi", "kajian"],
    "variabel":        ["parameter", "faktor", "indikator"],
    "parameter":       ["variabel", "faktor", "indikator"],
    "faktor":          ["variabel", "parameter", "aspek"],
    "aspek":           ["faktor", "variabel", "dimensi"],
    "fungsi":          ["peran", "tugas", "kegunaan"],
    "peran":           ["fungsi", "tugas", "kegunaan"],
    "kegunaan":        ["fungsi", "peran", "manfaat"],
    "manfaat":         ["kegunaan", "faedah", "nilai guna"],
    "faedah":          ["manfaat", "kegunaan", "nilai guna"],
    "kualitas":        ["mutu", "standar", "kelayakan"],
    "mutu":            ["kualitas", "standar", "kelayakan"],
    "keamanan":        ["sekuriti", "perlindungan", "proteksi"],
    "perlindungan":    ["keamanan", "sekuriti", "proteksi"],
    "pengelolaan":     [None, None, "manajemen", "pengurusan"],
    "manajemen":       [None, None, "pengelolaan", "pengurusan"],
    "administrasi":    [None, "pengelolaan", "manajemen"],
    # 'efisien' jangan diganti jadi 'hemat' (beda konteks akademik)
    "efisien":         [None, "tepat guna", "praktis"],
    "kebutuhan":       [None, None, "keperluan"],              # 'persyaratan' ≠ 'kebutuhan'
    "keperluan":       ["kebutuhan", "persyaratan"],
    "kapasitas":       ["kemampuan", "kapabilitas", "daya"],
    "kemampuan":       ["kapasitas", "kapabilitas", "kompetensi"],
    "komponen":        ["bagian", "elemen", "unsur"],
    "elemen":          ["komponen", "bagian", "unsur"],
    "unsur":           ["komponen", "elemen", "bagian"],
    "fitur":           ["fungsi", "kemampuan", "fasilitas"],
    "fasilitas":       ["fitur", "sarana", "kemampuan"],
    "arsitektur":      ["rancangan", "desain", "struktur"],
    "desain":          ["rancangan", "arsitektur", "model"],
    "rancangan":       ["desain", "arsitektur", "model"],
    "model":           ["rancangan", "desain", "pola"],
    "pola":            ["model", "rancangan", "bentuk"],
    "struktur":        ["susunan", "arsitektur", "kerangka"],
    "susunan":         ["struktur", "kerangka", "tatanan"],
    "alur":            ["aliran", "prosedur", "tahapan"],
    "aliran":          ["alur", "prosedur", "urutan"],

    # --- NATURALISASI: kata terlalu formal → lebih sederhana ---
    # Kata bombastis yang bikin "fingerprint" seragam
    "akselerasi":          ["percepatan", "perkembangan", "kemajuan"],
    "mengakselerasi":      ["mempercepat", "mendorong", "meningkatkan"],
    "mitigasi":            ["pencegahan", "penanggulangan", "penanganan"],
    "kapabilitas":         ["kemampuan", "kapasitas", "kompetensi"],
    "instrumen":           ["alat", "perangkat", "sarana"],
    "memformulasi":        ["merumuskan", "menyusun", "membuat"],
    "memformulasikan":     ["merumuskan", "menyusun", "membuat"],
    "trajektori":          ["arah", "jalur", "perkembangan"],
    "intervensi":          ["tindakan", "penanganan", "langkah"],
    "masif":               ["besar-besaran", "luas", "menyeluruh"],
    "masifnya":            ["luasnya", "banyaknya", "pesatnya"],
    "komprehensif":        ["menyeluruh", "lengkap", "mendalam"],
    "terpadu":             ["terintegrasi", "gabungan", "menyatu"],
    "terintegrasi":        ["terpadu", "gabungan", "terhubung"],
    "mengintegrasikan":    ["memadukan", "menggabungkan", "menyatukan"],
    "integrasi":           ["keterpaduan", "penggabungan", "penyatuan"],
    "real-time":           ["secara langsung", "seketika", "waktu nyata"],
    "monitoring":          ["pemantauan", "pengawasan", "pengecekan"],
    "validasi":            ["pengesahan", "verifikasi", "pembuktian"],
    "verifikasi":          ["validasi", "pengecekan", "konfirmasi"],
    "komputasi":           ["pengolahan data", "pemrosesan", "perhitungan"],
    "akurasi":             ["keakuratan", "ketepatan", "presisi"],
    "klasifikasi":         ["pengelompokan", "pengkategorian", "pembagian"],
    "identifikasi":        ["pengenalan", "penentuan", "penemuan"],
    "optimasi":            ["pengoptimalan", "peningkatan", "penyempurnaan"],
    "optimalisasi":        ["pengoptimalan", "peningkatan", "penyempurnaan"],
    "mengoptimalkan":      ["meningkatkan", "memperbaiki", "memaksimalkan"],
    "signifikan":          ["bermakna", "penting", "berarti"],
    "signifikansi":        ["makna", "arti penting", "kepentingan"],
    "kontribusi":          ["sumbangan", "peran", "andil"],
    "berkontribusi":       ["berperan", "menyumbang", "memberikan andil"],
    "implementasi":        ["penerapan", "pelaksanaan", "penggunaan"],
    "mengimplementasikan": ["menerapkan", "melaksanakan", "menggunakan"],
    "formulasi":           ["rumusan", "penyusunan", "pembentukan"],
    "konsistensi":         ["kestabilan", "keajegan", "kesamaan"],
    "konsisten":           ["stabil", "ajeg", "tetap"],
    "relevan":             ["sesuai", "terkait", "berkaitan"],
    "relevansi":           ["kesesuaian", "keterkaitan", "hubungan"],
    "inovasi":             ["pembaruan", "terobosan", "pengembangan baru"],
    "berinovasi":          ["memperbarui", "menciptakan terobosan"],
    "kolaborasi":          ["kerja sama", "kerjasama", "koordinasi"],
    "berkolaborasi":       ["bekerja sama", "berkoordinasi"],
    "koordinasi":          ["kerja sama", "kolaborasi", "keterpaduan"],
    "konvensional":        ["tradisional", "manual", "lama"],
    "digitalisasi":        ["pendigitalan", "transformasi digital"],
    "transformasi":        ["perubahan", "pembaruan", "transisi"],
    "otomatisasi":         ["pengotomatisan", "pemrosesan otomatis"],
    "notifikasi":          ["pemberitahuan", "peringatan", "pengumuman"],
    "dashboard":           ["dasbor", "panel utama", "halaman utama"],
    "inputan":             ["masukan", "data masukan", "input"],
    "output":              ["keluaran", "hasil", "luaran"],
    "update":              ["pembaruan", "perbaharuan", "perubahan"],
    "realtime":            ["secara langsung", "waktu nyata", "seketika"],
    "tata kelola":         ["pengelolaan", "manajemen", "pengaturan"],
    "pemangku":            ["pemegang", "pihak", "pelaku"],
    "krusial":             ["penting", "mendasar", "vital"],

    # --- Extra naturalizer: kata yg masih muncul dari Gemini engine ---
    "diarsitekturi":       ["dirancang", "dikembangkan", "dibangun"],
    "mengarsitekturi":     ["merancang", "membangun", "mengembangkan"],
    "konstelasi":          ["kondisi", "situasi", "keadaan"],
    "mengondisikan":       ["menyesuaikan", "mengatur", "mengarahkan"],
    "mengonseptualisasi":  ["merumuskan", "menyusun konsep", "menggambarkan"],
    "konseptualisasi":     ["perumusan", "penyusunan konsep", "gambaran"],
    "mentranskripsi":      ["mencatat", "merekam", "menyalin"],
    "mentransformasi":     ["mengubah", "mengonversi", "mengalihkan"],
    "mengkalkulasi":       ["menghitung", "menentukan nilai", "mengukur"],
    "kalkulasi":           ["perhitungan", "pengukuran", "penentuan nilai"],
    "rekognisi":           ["pengenalan", "identifikasi", "deteksi"],
    "menganalisis":        ["mengkaji", "menelaah", "membahas"],
    "menganalisa":         ["mengkaji", "menelaah", "membahas"],
    "urgensi":             ["kepentingan", "kebutuhan mendesak", "prioritas"],
    "urgen":               ["mendesak", "penting", "perlu segera"],
    "esensial":            ["penting", "mendasar", "pokok"],
    "fundamental":         ["mendasar", "pokok", "inti"],
    "substansial":         ["berarti", "signifikan", "penting"],
    "substansi":           ["isi", "inti", "pokok"],
    "konkret":             ["nyata", "jelas", "spesifik"],
    "kompetitif":          ["bersaing", "unggul", "andal"],
    "progresif":           ["berkembang", "maju", "meningkat"],
    "prospektif":          ["menjanjikan", "berpotensi", "potensial"],
    "holistik":            ["menyeluruh", "lengkap", "terpadu"],
    "pragmatis":           ["praktis", "realistis", "fungsional"],

    # ── REVERSE-NATURALIZER: kata tidak alami → kembalikan ke kata umum ─────────
    # Kata-kata ini MUNGKIN muncul dari engine parafrase → langsung kembalikan ke natural
    "sanggup":             ["dapat", "bisa", "mampu"],
    "rekaman":             ["data", "informasi", "catatan"],
    "luaran":              ["hasil", "output"],
    "keluaran":            [None, "hasil", "output"],
    "memfasilitasi":       ["membantu", "mempermudah", "mendukung"],
    "dipersyaratkan":      ["diperlukan", "dibutuhkan"],
    "mendongkrak":         ["meningkatkan", "mengoptimalkan", None],
    "ajeg":                ["konsisten", "stabil", "tetap"],
    "memproduksi":         [None, "menghasilkan", "membuat", "menciptakan"],
    "memutuskan":          [None, "menentukan", "menetapkan"],
    "dimensi":             [None, "aspek", "faktor", "sisi"],
    "penemuan":            [None, "temuan", "hasil", "identifikasi"],
    "diarsitekturi":       ["dirancang", "dikembangkan", "dibangun"],
    "mengarsitekturi":     ["merancang", "membangun", "mengembangkan"],
    "tepat guna":          [None, "efisien", "efektif", "praktis"],
    "unjuk kerja":         [None, "kinerja", "performa"],
    "nilai guna":          [None, "manfaat", "kegunaan"],
    "faedah":              ["manfaat", "kegunaan", None],
    "sekuriti":            ["keamanan", "perlindungan", None],
    "kapabilitas":         ["kemampuan", "kapasitas", None],
    "masifnya":            ["luasnya", "banyaknya", "pesatnya"],
}


# ── KATA AKADEMIK BEROPERASI TINGGI (PROTEKSI 90% TETAP AS-IS) ─────────────────
# Kata-kata ini jika diparafrase malah merusak konteks skripsi / jurnal akademik.
# Kita berikan peluang 90% untuk tetap dipertahankan (No Rewrite Decision).
ACADEMIC_HIGH_PROTECT = {
    # Hanya proteksi kata teknis yang BENAR-BENAR tidak boleh diganti
    # (istilah baku, nama teknologi, singkatan)
    "transaksi", "efektivitas",
}

def _get_synonym(word_lower: str) -> str:
    """Randomly pick one synonym from INDONESIAN_SYNONYMS_MULTI.
    If None is chosen, keep the original word (humanization: not every word gets replaced).
    Also respects ACADEMIC_HIGH_PROTECT.
    """
    if word_lower in ACADEMIC_HIGH_PROTECT:
        # 65% chance to keep truly-protected words as-is (was 90% — too conservative)
        if _random.random() < 0.65:
            return word_lower

    choices = INDONESIAN_SYNONYMS_MULTI.get(word_lower)
    if not choices:
        return word_lower
    picked = _random.choice(choices)
    if picked is None:
        return word_lower   # pertahankan kata asli -> lebih natural seperti manusia
    return picked



# --- HUMANIZATION LAYER: variasikan pembuka kalimat agar tidak ada "AI fingerprint" ---
# Pola pembuka yang terlalu sering dipakai AI → diganti dengan variasi acak
SENTENCE_OPENER_MULTI = {
    # ── "Sebagai..." patterns ─────────────────────────────────────────────────
    r'^(Sebagai salah satu)':       ["Salah satu", "Merupakan salah satu", "Termasuk salah satu", "Menjadi salah satu"],
    r'^(Sebagai sebuah)':           ["Merupakan sebuah", "Berupa sebuah", "Menjadi sebuah"],
    r'^(Sebagai bagian dari)':      ["Bagian dari", "Dalam rangka", "Termasuk dalam"],
    r'^(Sebagai upaya)':            ["Dalam upaya", "Guna", "Demi"],
    r'^(Sebagai bentuk)':           ["Dalam bentuk", "Merupakan bentuk", "Wujud dari"],

    # ── "Melalui..." patterns ─────────────────────────────────────────────────
    r'^(Melalui penerapan)':        ["Dengan menerapkan", "Lewat penerapan", "Penerapan"],
    r'^(Melalui penggunaan)':       ["Dengan menggunakan", "Lewat penggunaan", "Penggunaan"],
    r'^(Melalui proses)':           ["Dalam proses", "Lewat proses", "Dengan melalui proses"],
    r'^(Melalui pendekatan)':       ["Dengan pendekatan", "Lewat pendekatan", "Pendekatan"],
    r'^(Melalui analisis)':         ["Dengan menganalisis", "Dari hasil analisis", "Berdasarkan analisis"],
    r'^(Melalui penelitian)':       ["Dari hasil penelitian", "Berdasarkan studi", "Lewat riset"],

    # ── "Dalam..." patterns ───────────────────────────────────────────────────
    r'^(Dalam konteks ini)':        ["Terkait hal ini", "Berkaitan dengan ini", "Dalam hal ini", "Sehubungan dengan ini"],
    r'^(Dalam konteks)':            ["Terkait", "Berkaitan dengan", "Dalam hal", "Sehubungan dengan"],
    r'^(Dalam rangka)':             ["Guna", "Demi", "Untuk", "Dengan tujuan"],
    r'^(Dalam penelitian ini)':     ["Pada studi ini", "Di dalam riset ini", "Studi ini"],
    r'^(Dalam hal ini)':            ["Terkait hal ini", "Berkenaan dengan ini", "Sehubungan dengan itu"],
    r'^(Dalam upaya)':              ["Guna", "Demi", "Sebagai upaya", "Untuk"],
    r'^(Dalam proses)':             ["Pada tahapan", "Di dalam proses", "Saat proses"],

    # ── "Secara..." patterns ──────────────────────────────────────────────────
    r'^(Secara konseptual)':        ["Secara teoritis", "Dari sisi konsep", "Pada tataran konsep", "Dari perspektif teori"],
    r'^(Secara keseluruhan)':       ["Secara umum", "Pada dasarnya", "Bila dilihat secara utuh", "Jika dilihat menyeluruh"],
    r'^(Secara khusus)':            ["Lebih spesifik", "Terkhusus", "Khususnya", "Secara spesifik"],
    r'^(Secara umum)':              ["Pada umumnya", "Umumnya", "Secara garis besar", "Pada dasarnya"],
    r'^(Secara signifikan)':        ["Secara nyata", "Cukup signifikan", "Dengan jelas", "Secara bermakna"],
    r'^(Secara teoritis)':          ["Dari sudut pandang teori", "Secara konseptual", "Menurut teori"],
    r'^(Secara empiris)':           ["Berdasarkan data", "Dari temuan lapangan", "Secara faktual"],
    r'^(Secara praktis)':           ["Dalam praktiknya", "Secara nyata", "Pada kenyataannya"],

    # ── "Hal ini / Hal tersebut" patterns ────────────────────────────────────
    r'^(Hal ini menunjukkan bahwa)':  ["Ini mengindikasikan bahwa", "Fakta ini memperlihatkan bahwa", "Kondisi ini menandakan bahwa"],
    r'^(Hal ini menunjukkan)':        ["Ini menunjukkan", "Kondisi ini memperlihatkan", "Fakta ini mengindikasikan", "Ini mengindikasikan"],
    r'^(Hal ini disebabkan)':         ["Kondisi ini disebabkan", "Ini terjadi karena", "Penyebabnya adalah", "Hal ini diakibatkan oleh"],
    r'^(Hal ini dikarenakan)':        ["Ini disebabkan oleh", "Kondisi ini terjadi karena", "Penyebabnya ialah"],
    r'^(Hal tersebut menunjukkan)':   ["Kondisi tersebut memperlihatkan", "Ini mengindikasikan", "Temuan itu menandakan"],
    r'^(Hal tersebut)':               ["Kondisi tersebut", "Situasi ini", "Keadaan ini", "Hal itu"],
    r'^(Hal ini)':                    ["Kondisi ini", "Keadaan ini", "Ini", "Situasi ini"],

    # ── "Dengan demikian / Dengan..." patterns ───────────────────────────────
    r'^(Dengan demikian)':          ["Oleh karena itu", "Maka dari itu", "Karenanya", "Dengan begitu"],
    r'^(Dengan kata lain)':         ["Artinya", "Maksudnya", "Dengan demikian", "Singkatnya"],
    r'^(Dengan adanya)':            ["Hadirnya", "Keberadaan", "Berkat adanya"],
    r'^(Dengan menggunakan)':       ["Melalui penggunaan", "Memanfaatkan", "Dengan memakai"],
    r'^(Dengan memanfaatkan)':      ["Melalui pemanfaatan", "Menggunakan", "Dengan memakai"],
    r'^(Dengan menerapkan)':        ["Melalui penerapan", "Lewat implementasi", "Dengan mengimplementasikan"],
    r'^(Dengan demikian, dapat)':   ["Sehingga dapat", "Maka dapat", "Oleh karena itu dapat"],

    # ── "Berdasarkan..." patterns ─────────────────────────────────────────────
    r'^(Berdasarkan hasil penelitian)': ["Berlandaskan temuan studi", "Dari hasil riset", "Mengacu pada temuan penelitian"],
    r'^(Berdasarkan hasil)':        ["Dari temuan", "Berlandaskan hasil", "Mengacu pada hasil"],
    r'^(Berdasarkan analisis)':     ["Dari hasil analisis", "Berlandaskan kajian", "Mengacu pada analisis"],
    r'^(Berdasarkan uraian)':       ["Berdasarkan pemaparan", "Berlandaskan uraian", "Dari penjelasan"],
    r'^(Berdasarkan)':              ["Berlandaskan", "Mengacu pada", "Merujuk pada", "Sesuai dengan"],

    # ── "Penelitian ini / Studi ini" patterns ─────────────────────────────────
    r'^(Penelitian ini bertujuan untuk)': ["Studi ini dirancang guna", "Riset ini dimaksudkan untuk", "Tujuan studi ini adalah"],
    r'^(Penelitian ini)':           ["Studi ini", "Riset ini", "Kajian ini"],
    r'^(Studi ini)':                ["Penelitian ini", "Riset ini", "Kajian ini"],

    # ── "Oleh karena itu / Oleh sebab itu" patterns ──────────────────────────
    r'^(Oleh karena itu)':          ["Dengan demikian", "Maka dari itu", "Karenanya", "Sehubungan dengan itu"],
    r'^(Oleh sebab itu)':           ["Oleh karena itu", "Maka dari itu", "Dengan demikian"],

    # ── "Selain itu / Di samping itu" patterns ───────────────────────────────
    r'^(Selain itu)':               ["Di samping itu", "Lebih lanjut", "Tak hanya itu", "Di luar itu"],
    r'^(Di samping itu)':           ["Selain itu", "Lebih lanjut", "Tak hanya itu"],
    r'^(Lebih lanjut)':             ["Selain itu", "Lebih jauh", "Lebih dari itu"],
    r'^(Lebih jauh)':               ["Lebih lanjut", "Selain itu", "Lebih dari itu"],

    # ── "Adapun / Sehubungan / Berkenaan" patterns ───────────────────────────
    r'^(Adapun)':                   ["Sehubungan dengan itu", "Terkait hal tersebut", "Mengenai"],
    r'^(Sehubungan dengan)':        ["Berkaitan dengan", "Terkait dengan", "Mengenai"],
    r'^(Berkenaan dengan)':         ["Terkait dengan", "Sehubungan dengan", "Mengenai"],

    # ── "Perlu diketahui / Perlu dicatat" patterns ───────────────────────────
    r'^(Perlu diketahui)':          ["Patut diperhatikan", "Perlu dicatat", "Penting untuk diketahui"],
    r'^(Perlu dicatat)':            ["Patut dicatat", "Perlu diperhatikan", "Penting untuk dicatat"],
    r'^(Perlu diperhatikan)':       ["Patut diperhatikan", "Perlu dicatat", "Penting untuk diperhatikan"],

    # ── "Dapat disimpulkan / Dapat dikatakan" patterns ───────────────────────
    r'^(Dapat disimpulkan bahwa)':  ["Dapat diambil kesimpulan bahwa", "Simpulan yang diperoleh adalah", "Dari uraian tersebut disimpulkan"],
    r'^(Dapat disimpulkan)':        ["Dapat diambil kesimpulan", "Disimpulkan", "Kesimpulannya"],
    r'^(Dapat dikatakan bahwa)':    ["Dapat dinyatakan bahwa", "Artinya", "Hal ini berarti"],
    r'^(Dapat dikatakan)':          ["Dapat dinyatakan", "Bisa dikatakan", "Dapat disebutkan"],

    # ── "Menurut..." patterns ─────────────────────────────────────────────────
    r'^(Menurut para ahli)':        ["Menurut kalangan pakar", "Para pakar menyatakan", "Menurut para peneliti"],
    r'^(Menurut)':                  ["Sesuai pendapat", "Berdasarkan pandangan", "Sebagaimana diungkapkan"],

    # ── "Tidak hanya / Tidak saja" patterns ──────────────────────────────────
    r'^(Tidak hanya)':              ["Bukan sekadar", "Tak sebatas", "Lebih dari sekadar"],
    r'^(Tidak hanya itu)':          ["Bukan hanya itu", "Tak hanya itu saja", "Lebih dari itu"],
}

import re as _re_opener

def _humanize_opener(text: str) -> str:
    """Replace repetitive AI sentence openers with a random alternative."""
    for pattern, alternatives in SENTENCE_OPENER_MULTI.items():
        m = _re_opener.match(pattern, text)
        if m:
            chosen = _random.choice(alternatives)
            rest = text[m.end():]
            # Capitalize first letter of rest if needed
            if rest and rest[0].islower():
                rest = rest[0].upper() + rest[1:]
            return chosen + " " + rest
    return text


# --- Frasa yang SAMA SEKALI tidak boleh disentuh ---
EXCLUDE_PHRASES = [
    # Metadata akademik - WAJIB DILINDUNGI
    "program studi", "jenjang studi", "fakultas", "universitas", "jurusan",
    "nim", "nip", "dosen pembimbing", "dosen penguji", "ketua sidang",
    # ── ISTILAH BAKU SISTEM INFORMASI / TEKNOLOGI INFORMASI ──────────────────
    "teknologi informasi", "sistem informasi akuntansi", "sistem informasi manajemen",
    "sistem informasi", "sistem penggajian", "sistem pembayaran",
    "sistem pendukung keputusan", "decision support system", "dss",
    "tarif efektif rata-rata", "tarif efektif",
    "daftar pustaka", "daftarpustaka",
    # Teknologi & Tools
    "use case diagram", "sequence diagram", "activity diagram", "class diagram",
    "entity relationship diagram", "erd", "use case",
    "waterfall", "agile", "scrum", "kanban",
    "database", "data base", "data warehouse",
    "black box testing", "white box testing", "unit testing",
    "php", "mysql", "postgresql", "mongodb", "sqlite",
    "react", "vue", "angular", "laravel", "codeigniter", "django", "flask",
    "inertia.js", "node.js", "javascript", "python", "java", "kotlin",
    "haar cascade", "lbph", "uml", "api", "rest api", "xml", "json", "html", "css",
    "cloud computing", "artificial intelligence", "machine learning", "deep learning",
    "neural network", "natural language processing", "computer vision",
    "internet of things", "iot",
    # Istilah domain IT yang sering salah diganti
    "data transaksi", "data keuangan", "data pengguna", "data pelanggan",
    "input data", "output data",
    "basis data", "basis pengetahuan",
    "pengelolaan keuangan", "administrasi keuangan",
    "kas masuk", "kas keluar",
    "laporan keuangan",
    "pengambilan keputusan",
    "rekonsiliasi bank",
    # Istilah statistik/akademik
    "pph 21", "pajak penghasilan", "pajak pertambahan nilai", "ppn",
    "skala likert", "uji validitas", "uji reliabilitas", "uji normalitas",
    "regresi linier", "regresi berganda", "korelasi pearson", "anova",
    "mean absolute error", "root mean square error", "mae", "rmse", "mse",
    "f-measure", "precision", "recall", "accuracy",
    # Istilah akademik umum yang tidak boleh diubah
    "literatur review", "tinjauan pustaka",
    "latar belakang", "rumusan masalah", "batasan masalah",
    "tujuan penelitian", "manfaat penelitian",
    # Tools Microsoft & Office
    "microsoft excel", "microsoft word", "microsoft access",
    "google sheets", "google docs",
    # Istilah UI/UX & sistem yang tidak boleh diubah
    "dashboard", "admin", "administrator",
    "input", "output",
]

# ── COLLOCATION BLACKLIST: frasa tidak natural → auto-koreksi ke frasa baku ──
# Format: (pola_regex, pengganti_natural)
# Dijalankan sebagai POST-PROCESSING setelah synonym replacement
FORBIDDEN_COLLOCATIONS = [
    # Teknologi + kata salah
    (r'\bteknologi\s+(keterangan|catatan|rekaman)\b', 'teknologi informasi'),
    (r'\bteknologi\s+data\b', 'teknologi informasi'),
    # Sistem + kata salah  
    (r'\bsistem\s+keterangan\b', 'sistem informasi'),
    (r'\bsistem\s+catatan\b', 'sistem informasi'),
    # Data/informasi compound yang sering salah
    (r'\bcatatan\s+transaksi\b', 'data transaksi'),
    (r'\bketerangan\s+transaksi\b', 'data transaksi'),
    (r'\bcatatan\s+keuangan\b', 'data keuangan'),
    (r'\bketerangan\s+keuangan\b', 'informasi keuangan'),
    # Pengambilan keputusan (bukan "membangun" atau "membuat" keputusan)
    (r'\bmembangun\s+keputusan\b', 'mengambil keputusan'),
    (r'\bmembuat\s+keputusan\b', 'mengambil keputusan'),
    (r'\bmenyusun\s+keputusan\b', 'mengambil keputusan'),
    # Pengelolaan (bukan pengaturan) keuangan
    (r'\bpengaturan\s+keuangan\b', 'pengelolaan keuangan'),
    (r'\bpengaturan\s+kas\b', 'pengelolaan kas'),
    # Laporan/catatan kas
    (r'\bcatatan\s+kas\s+masuk\b', 'laporan kas masuk'),
    (r'\bcatatan\s+kas\s+keluar\b', 'laporan kas keluar'),
    # User / pengguna akhir
    (r'\bpengguna\s+akhir\b', 'pengguna'),
    (r'\buser\s+akhir\b', 'pengguna'),
    # Istilah output yang tidak natural
    (r'\bluaran\s+sistem\b', 'output sistem'),
    (r'\bkeluaran\s+sistem\b', 'output sistem'),
    # Kata tidak natural untuk konteks akademik
    (r'\bsungguh\s+berisiko\b', 'sangat berisiko'),
    (r'\bsungguh\s+penting\b', 'sangat penting'),
    (r'\bsigap\b', 'cepat'),
    (r'\bajeg\b', 'konsisten'),
    (r'\brekaman\b', 'data'),
    (r'\bluaran\b', 'hasil'),
    (r'\btemuan\s+riset\b', 'hasil penelitian'),
    # Tambahan Review V5
    (r'\bprogres\s+teknologi\b', 'perkembangan teknologi'),
    (r'\btahapan\s+pelaporan\b', 'proses pelaporan'),
    (r'\bcatatan\s+perusahaan\b', 'data perusahaan'),
    (r'\bcatatan\s+transaksi\b', 'data transaksi'),
    # memfasilitasi → tidak natural, kembalikan ke membantu/mendukung
    (r'\bmemfasilitasi\b', 'membantu'),
    (r'\bmemproduksi\s+informasi\b', 'menghasilkan informasi'),
    (r'\bbasis\s+catatan\b', 'basis data'),
    # ── TAMBAHAN V9: Perbaikan berdasarkan review dokumen nyata ──────────────
    # galat → kesalahan (galat terlalu teknis/kaku)
    (r'\bgalat\b', 'kesalahan'),
    (r'\bmemangkas\s+galat\b', 'mengurangi kesalahan'),
    # cermat guna → efisien (tidak natural dalam bahasa Indonesia)
    (r'\bcermat\s+guna\b', 'efisien'),
    (r'\btepat\s+guna\b', 'efisien'),
    # manjur → efektif
    (r'\bmanjur\b', 'efektif'),
    # hemat dalam konteks sistem/teknologi → efisien
    (r'\blebih\s+hemat\s+(dalam|untuk)\b', 'lebih efisien dalam'),
    (r'\bsecara\s+hemat\b', 'secara efisien'),
    # sekuriti → keamanan (bahasa Indonesia baku)
    (r'\bsekuriti\b', 'keamanan'),
    # catatan (jika bukan laporan) → data
    (r'\bcatatan\s+(keuangan|kas|bank|transaksi|perusahaan|sistem)\b', r'data \1'),
    # keterangan → informasi (dalam konteks sistem informasi)
    (r'\bketerangan\s+terstruktur\b', 'informasi terstruktur'),
    (r'\bketerangan\s+akurat\b', 'informasi akurat'),
    # otonom (monitoring context) → otomatis
    (r'\bsecara\s+otonom\b', 'secara otomatis'),
    (r'\bmemantau\s+(.{1,30})\s+otonom\b', r'memantau \1 secara otomatis'),
    # tuntutan (kebutuhan context) → kebutuhan (gunakan lambda karena ada grup)
    (r'\btuntutan\s+(pengguna|bisnis|perusahaan|user)\b', lambda m: 'kebutuhan ' + m.group(1)),

    # sanggup → dapat/bisa (sanggup terlalu informal)
    (r'\bsanggup\b', 'dapat'),
]


def _apply_collocation_fix(text: str) -> str:
    """Apply FORBIDDEN_COLLOCATIONS blacklist to auto-correct unnatural phrases.
    Supports both string replacements and callable (lambda) replacements.
    """
    result = text
    for pattern, replacement in FORBIDDEN_COLLOCATIONS:
        if callable(replacement):
            # Lambda/function replacement (for rules with regex groups)
            def _call_replacer(m, fn=replacement):
                rep = fn(m)
                matched = m.group(0)
                if matched[0].isupper() and rep:
                    return rep[0].upper() + rep[1:]
                return rep
            result = re.sub(pattern, _call_replacer, result, flags=re.IGNORECASE)
        else:
            def _str_replacer(m, rep=replacement):
                matched = m.group(0)
                if matched[0].isupper():
                    return rep[0].upper() + rep[1:]
                return rep
            result = re.sub(pattern, _str_replacer, result, flags=re.IGNORECASE)
    return result


# Kata sambung jamak yang sering rusak saat replace
CONJUNCTIONS_PROTECT = [
    "oleh karena itu", "akan tetapi", "meskipun demikian", "dengan demikian",
    "sehubungan dengan", "berkenaan dengan", "berkaitan dengan",
    "selain itu", "di samping itu", "lebih lanjut", "lebih jauh",
]

# --- Parafrase level frasa (lebih kuat dari sekadar sinonim kata) ---
PHRASE_PARAPHRASE = {
    # ── Pembuka kalimat penelitian ──────────────────────────────────────────
    "penelitian ini bertujuan untuk": "studi ini dirancang guna",
    "penelitian ini bertujuan": "studi ini memiliki sasaran",
    "tujuan dari penelitian ini": "sasaran dari studi ini",
    "tujuan penelitian ini": "sasaran studi ini",
    "dalam penelitian ini": "dalam studi ini",
    "hasil penelitian ini": "temuan studi ini",
    "berdasarkan hasil penelitian": "berlandaskan temuan studi",
    "berdasarkan penelitian": "berlandaskan studi",
    "penelitian ini dilakukan": "studi ini dijalankan",
    "penelitian yang dilakukan": "studi yang dijalankan",
    "penelitian ini menggunakan": "studi ini memanfaatkan",
    "penelitian ini membahas": "studi ini mengulas",
    "tujuan dari studi ini": "sasaran dari penelitian ini",
    "studi ini dilakukan": "penelitian ini dijalankan",

    # ── Frasa umum akademik ──────────────────────────────────────────────────
    "dapat disimpulkan bahwa": "dapat ditarik kesimpulan bahwa",
    "dapat disimpulkan": "dapat diambil kesimpulan",
    "dengan demikian dapat": "maka dari itu dapat",
    "sehingga dapat disimpulkan": "sehingga dapat diambil kesimpulan",
    "berdasarkan uraian di atas": "berdasarkan pemaparan tersebut",
    "berdasarkan uraian tersebut": "berlandaskan pemaparan di atas",
    "sebagaimana yang telah dijelaskan": "sebagaimana yang telah dipaparkan",
    "sebagaimana telah disebutkan": "seperti yang telah diuraikan",
    "seperti yang telah dijelaskan": "sebagaimana telah dipaparkan",
    "pada penelitian sebelumnya": "pada studi terdahulu",
    "pada studi sebelumnya": "pada penelitian terdahulu",
    "yang telah dilakukan sebelumnya": "yang pernah dilaksanakan sebelumnya",
    "yang dilakukan sebelumnya": "yang pernah dilaksanakan",
    "menurut para ahli": "menurut kalangan pakar",
    "menurut beberapa ahli": "menurut sejumlah pakar",
    "para ahli menyatakan": "kalangan pakar mengungkapkan",
    "sebagaimana yang dimaksud": "sesuai dengan yang dinyatakan",
    "sebagaimana dimaksud": "sesuai dengan yang dinyatakan",
    "perlu diperhatikan bahwa": "patut dicermati bahwa",
    "perlu diketahui bahwa": "perlu dipahami bahwa",
    "yang perlu diperhatikan": "yang patut dicermati",
    "dengan memperhatikan": "dengan mempertimbangkan",
    "mengacu pada": "berpedoman pada",
    "mengacu kepada": "berpedoman kepada",
    "sesuai dengan": "sejalan dengan",
    "berkaitan dengan": "berhubungan dengan",
    "sehubungan dengan hal": "terkait dengan hal",
    "dalam hal ini": "terkait hal ini",
    "yang pada dasarnya": "yang pada intinya",
    "pada dasarnya": "pada intinya",
    "pada hakikatnya": "pada intinya",
    "yang bersangkutan": "yang terkait",

    # ── Penggunaan sistem ────────────────────────────────────────────────────
    "sistem ini dapat": "perangkat ini mampu",
    "sistem ini digunakan": "perangkat ini dimanfaatkan",
    "sistem yang digunakan": "perangkat yang dimanfaatkan",
    "sistem yang dikembangkan": "perangkat lunak yang dibangun",
    "sistem yang dibangun": "perangkat lunak yang dikembangkan",
    "sistem yang dirancang": "perangkat lunak yang dikembangkan",
    "digunakan untuk membantu": "dimanfaatkan guna mempermudah",
    "digunakan untuk mempermudah": "dimanfaatkan guna membantu",
    "digunakan dalam": "dimanfaatkan dalam",
    "digunakan pada": "dimanfaatkan pada",
    "dapat digunakan": "dapat dimanfaatkan",
    "yang dapat membantu": "yang dapat mempermudah",
    "untuk membantu": "guna mempermudah",
    "dalam membantu": "dalam mempermudah",
    "membantu dalam": "mempermudah dalam",
    "membantu proses": "mempermudah tahapan",
    "perangkat lunak yang": "aplikasi yang",
    "mengembangkan sistem": "membangun perangkat",
    "pembangunan sistem": "pengembangan perangkat",
    "perancangan sistem": "desain perangkat",
    "pengujian sistem": "evaluasi perangkat",

    # ── Masalah dan solusi ───────────────────────────────────────────────────
    "permasalahan yang ada": "kendala yang dihadapi",
    "masalah yang dihadapi": "kendala yang ada",
    "masalah yang ada": "permasalahan yang dihadapi",
    "untuk mengatasi masalah": "guna menyelesaikan kendala",
    "untuk mengatasi permasalahan": "guna mengatasi kendala",
    "dalam mengatasi masalah": "dalam menyelesaikan permasalahan",
    "solusi yang ditawarkan": "solusi yang diajukan",
    "solusi dari permasalahan": "jawaban atas kendala",
    "untuk memecahkan masalah": "guna menyelesaikan persoalan",
    "upaya mengatasi": "langkah menyelesaikan",
    "dalam rangka mengatasi": "guna menyelesaikan",
    "yang menjadi masalah": "yang menjadi kendala",

    # ── Metodologi ───────────────────────────────────────────────────────────
    "metode yang digunakan": "pendekatan yang dimanfaatkan",
    "metode yang digunakan dalam": "pendekatan yang digunakan dalam",
    "metode pengembangan": "pendekatan pengembangan",
    "dengan menggunakan metode": "dengan memanfaatkan pendekatan",
    "menggunakan metode": "memanfaatkan pendekatan",
    "tahapan pengembangan": "langkah-langkah pengembangan",
    "proses pengembangan": "tahapan perancangan",
    "proses perancangan": "langkah-langkah perancangan",
    "proses pembuatan": "tahapan pembuatan",
    "dalam proses": "dalam tahapan",
    "melalui proses": "melalui tahapan",
    "metode penelitian yang": "pendekatan riset yang",
    "pengumpulan data": "penghimpunan data",
    "pengolahan data": "pemrosesan data",
    "analisis data": "kajian data",
    "teknik pengumpulan": "cara penghimpunan",
    "teknik analisis": "pendekatan analisis",
    "objek penelitian": "subjek studi",
    "sampel penelitian": "sampel studi",
    "populasi penelitian": "populasi studi",

    # ── Tinjauan pustaka / Landasan teori ────────────────────────────────────
    "landasan teori": "kajian teori",
    "tinjauan pustaka": "kajian literatur",
    "kajian teori": "tinjauan teoritis",
    "menurut pendapat": "berdasarkan pandangan",
    "dikemukakan oleh": "disampaikan oleh",
    "disampaikan oleh": "dikemukakan oleh",
    "berpendapat bahwa": "mengungkapkan bahwa",
    "mendefinisikan": "mengartikan",
    "didefinisikan sebagai": "diartikan sebagai",
    "pengertian dari": "definisi dari",
    "dapat didefinisikan": "dapat diartikan",
    "teori yang digunakan": "teori yang dimanfaatkan",
    "konsep dasar": "konsep inti",
    "dasar teori": "landasan teoritis",

    # ── Kesimpulan dan saran ─────────────────────────────────────────────────
    "berdasarkan pengujian yang dilakukan": "berlandaskan pengujian yang dijalankan",
    "pengujian yang dilakukan": "pengujian yang dijalankan",
    "hasil pengujian menunjukkan": "hasil pengujian memperlihatkan",
    "hasil pengujian menunjukkan bahwa": "temuan pengujian memperlihatkan bahwa",
    "menunjukkan bahwa sistem": "memperlihatkan bahwa perangkat",
    "telah berhasil": "telah sukses",
    "berhasil diimplementasikan": "sukses diterapkan",
    "berhasil dikembangkan": "sukses dibangun",
    "berhasil dirancang": "sukses didesain",
    "dapat disarankan": "direkomendasikan",
    "saran yang diberikan": "rekomendasi yang diajukan",
    "untuk penelitian selanjutnya": "bagi studi berikutnya",
    "penelitian lebih lanjut": "studi lebih mendalam",
    "hasil penelitian menunjukkan": "temuan studi memperlihatkan",
    "hasil penelitian membuktikan": "temuan riset membuktikan",
    "hasil penelitian menunjukkan bahwa": "temuan studi memperlihatkan bahwa",

    # ── Kata penghubung dan penyambung ──────────────────────────────────────
    "hal ini menunjukkan": "hal ini memperlihatkan",
    "hal ini menunjukkan bahwa": "hal ini memperlihatkan bahwa",
    "hal tersebut menunjukkan": "hal tersebut memperlihatkan",
    "hal ini disebabkan": "hal ini diakibatkan",
    "hal tersebut disebabkan": "hal tersebut diakibatkan",
    "hal ini dikarenakan": "hal ini disebabkan oleh fakta",
    "hal ini karena": "hal ini disebabkan",
    "karena hal tersebut": "disebabkan hal tersebut",
    "tidak hanya itu": "lebih dari itu",
    "selain daripada itu": "di samping itu",
    "dalam kaitannya": "dalam hubungannya",
    "terkait dengan hal": "berhubungan dengan hal",
    "yang dimaksud dengan": "yang diartikan sebagai",

    # ── Manfaat dan nilai ────────────────────────────────────────────────────
    "manfaat yang diperoleh": "kegunaan yang didapat",
    "manfaat dari": "kegunaan dari",
    "memberikan manfaat": "memberikan kegunaan",
    "sangat bermanfaat": "amat berguna",
    "bermanfaat bagi": "berguna bagi",
    "memberikan kemudahan": "memberikan kelancaran",
    "mempermudah pekerjaan": "membantu pekerjaan",
    "mempermudah proses": "membantu tahapan",
    "meningkatkan efisiensi": "memaksimalkan efisiensi",
    "meningkatkan efektivitas": "memaksimalkan efektivitas",
    "meningkatkan kualitas": "memaksimalkan kualitas",
    "meningkatkan kinerja": "memaksimalkan performa",
    "meningkatkan produktivitas": "mendongkrak produktivitas",
    "memberikan dampak": "menimbulkan dampak",
    "dampak yang ditimbulkan": "pengaruh yang dihasilkan",
    "dampak positif": "pengaruh yang baik",
    "dampak negatif": "pengaruh yang kurang baik",

    # ── Kata kerja umum akademik ─────────────────────────────────────────────
    "dapat dilihat bahwa": "terlihat bahwa",
    "dapat dilihat dari": "terlihat dari",
    "terlihat bahwa": "tampak bahwa",
    "nampak bahwa": "tampak bahwa",
    "dapat diketahui bahwa": "diketahui bahwa",
    "diketahui bahwa": "terungkap bahwa",
    "perlu dipahami bahwa": "penting untuk dipahami bahwa",
    "harus dipahami bahwa": "perlu dipahami bahwa",
    "tidak dapat dipungkiri": "tidak bisa disangkal",
    "tidak dapat dihindari": "tidak bisa dielakkan",
    "yang telah disebutkan": "yang telah diuraikan",
    "yang sudah dijelaskan": "yang telah dipaparkan",
    "yang dijelaskan di atas": "yang dipaparkan sebelumnya",
}

def _is_direct_quote_or_citation(sentence: str) -> bool:
    """Returns True if sentence is a direct quote or inline citation — must not be paraphrased."""
    s = sentence.strip()
    # Kalimat yang dimulai dan/atau diakhiri tanda kutip
    if (s.startswith('"') and s.endswith('"')) or (s.startswith("'") and s.endswith("'")):
        return True
    # Kutipan dengan tanda petik khusus Indonesia
    if s.startswith('\u201c') or s.startswith('\u2018'):
        return True
    # Sitasi inline: (Nama, 2020) atau (Nama 2020) atau [1] atau [12]
    if re.match(r'^\([A-Z][^)]+,?\s*\d{4}[a-z]?\)', s):
        return True
    if re.match(r'^\[\d+\]', s):
        return True
    # Baris yang hampir seluruhnya sitasi: banyak (Nama, Tahun) di dalamnya
    citations = re.findall(r'\([A-Z][a-zA-Z ]+,?\s*\d{4}\)', s)
    if len(citations) >= 2 and len(citations) * 30 > len(s):
        return True
    return False


# ── NUMBER → WORDS CONVERTER (Indonesian) ─────────────────────────────────────
_SATUAN_ID = [
    '', 'satu', 'dua', 'tiga', 'empat', 'lima',
    'enam', 'tujuh', 'delapan', 'sembilan', 'sepuluh',
    'sebelas', 'dua belas', 'tiga belas', 'empat belas', 'lima belas',
    'enam belas', 'tujuh belas', 'delapan belas', 'sembilan belas'
]
_PULUHAN_ID = [
    '', '', 'dua puluh', 'tiga puluh', 'empat puluh', 'lima puluh',
    'enam puluh', 'tujuh puluh', 'delapan puluh', 'sembilan puluh'
]

def _int_to_id_words(n: int) -> str:
    """Convert integer 0-999 to Indonesian words."""
    if n < 0:
        return 'minus ' + _int_to_id_words(-n)
    if n == 0:
        return 'nol'
    if n < 20:
        return _SATUAN_ID[n]
    if n < 100:
        tens = _PULUHAN_ID[n // 10]
        ones = _SATUAN_ID[n % 10]
        return (tens + ' ' + ones).strip()
    if n < 1000:
        hundreds = n // 100
        prefix = 'seratus' if hundreds == 1 else _SATUAN_ID[hundreds] + ' ratus'
        rest = _int_to_id_words(n % 100) if n % 100 else ''
        return (prefix + ' ' + rest).strip()
    return str(n)  # fallback untuk angka besar

def _number_to_words_in_sentence(text: str) -> str:
    """Replace standalone small numbers (1-99) with Indonesian words probabilistically.
    Angka > 99, tahun (1900-2099), dan nomor versi dibiarkan as-is.
    Hanya ~60% angka yang dikonversi agar tetap natural.
    """
    def replacer(m):
        raw = m.group(0)
        # Jangan ubah jika diawali/diakhiri huruf (singkatan, kode)
        before = text[max(0, m.start()-1):m.start()]
        after = text[m.end():m.end()+1]
        if before.isalpha() or after.isalpha():
            return raw
        n = int(raw)
        # Lindungi tahun, versi, dan angka besar
        if 1900 <= n <= 2099:
            return raw
        if n > 99:
            return raw
        # 60% kemungkinan dikonversi
        if _random.random() < 0.6:
            return _int_to_id_words(n)
        return raw
    return re.sub(r'\b(\d+)\b', replacer, text)


# ── SENTENCE RESTRUCTURER ──────────────────────────────────────────────────────
# Mengubah STRUKTUR kalimat, bukan hanya sinonimnya.
# Teknik: (1) passive↔active, (2) object-fronting, (3) temporal inversion

# Pola kalimat pasif yang umum di skripsi Indonesia
_PASSIVE_PATTERNS = [
    # "X digunakan untuk Y" → "Y menggunakan X" / "Untuk Y, X dimanfaatkan"
    (re.compile(
        r'^(.+?)\s+(digunakan|dimanfaatkan|diterapkan|dipakai)\s+untuk\s+(.+?)$',
        re.IGNORECASE
    ), lambda m: f"Untuk {m.group(3)}, {m.group(1)} {m.group(2).replace('di','me',1)}"),

    # "X dilakukan untuk Y" → "Untuk Y, dilakukan X"
    (re.compile(
        r'^(.+?)\s+(dilakukan|dijalankan|dilaksanakan)\s+untuk\s+(.+?)$',
        re.IGNORECASE
    ), lambda m: f"Untuk {m.group(3)}, {m.group(2)} {m.group(1)}"),

    # "Penelitian ini menggunakan X untuk Y" → "X digunakan dalam studi ini guna Y"
    (re.compile(
        r'^(Penelitian ini|Studi ini|Riset ini)\s+(menggunakan|memanfaatkan|menerapkan)\s+(.+?)\s+untuk\s+(.+?)$',
        re.IGNORECASE
    ), lambda m: f"{m.group(3).capitalize()} {m.group(2).replace('meng','di').replace('meman','diman').replace('mener','diter')} dalam {m.group(1).lower()} guna {m.group(4)}"),

    # "X dapat digunakan" → "Penggunaan X dapat dilakukan"
    (re.compile(
        r'^(.{10,60})\s+(dapat\s+digunakan|dapat\s+dimanfaatkan|dapat\s+diterapkan)$',
        re.IGNORECASE
    ), lambda m: f"Pemanfaatan {m.group(1).lower()} {m.group(2).replace('dapat di','bisa di')}"),
]

# Pola temporal/conditional inversion
_TEMPORAL_PATTERNS = [
    # "Setelah X, Y" → "Y setelah X"
    (re.compile(r'^(Setelah|Sesudah|Ketika|Saat|Apabila|Jika|Bila)\s+(.{10,60}?),\s+(.{10,})$', re.IGNORECASE),
     lambda m: f"{m.group(3).capitalize()}, {m.group(1).lower()} {m.group(2)}"),

    # "Dengan menggunakan X, Y" → "Y dengan menggunakan X"
    (re.compile(r'^(Dengan\s+(?:menggunakan|memanfaatkan|menerapkan)\s+.{5,50}?),\s+(.{10,})$', re.IGNORECASE),
     lambda m: f"{m.group(2).capitalize()} {m.group(1).lower()}"),

    # "Dalam rangka X, Y" → "Y dalam rangka X"
    (re.compile(r'^(Dalam rangka\s+.{5,50}?),\s+(.{10,})$', re.IGNORECASE),
     lambda m: f"{m.group(2).capitalize()} {m.group(1).lower()}"),

    # "Berdasarkan X, Y" → "Y berdasarkan X"
    (re.compile(r'^(Berdasarkan|Berlandaskan|Mengacu pada)\s+(.{5,60}?),\s+(.{10,})$', re.IGNORECASE),
     lambda m: f"{m.group(3).capitalize()} {m.group(1).lower()} {m.group(2)}"),
]

def _restructure_sentence(sentence: str, aggression: float = 0.6) -> str:
    """Probabilistically restructure a sentence for maximum Turnitin evasion.
    Higher aggression → more likely to restructure.
    Only fires on sufficiently long sentences (>= 40 chars).
    """
    if len(sentence) < 40:
        return sentence
    # Selalu jalankan jika pola cocok untuk meminimalkan Turnitin similarity secara maksimal

    # Coba passive patterns dulu (lebih kuat)
    for pattern, transform in _PASSIVE_PATTERNS:
        m = pattern.match(sentence)
        if m:
            try:
                result = transform(m)
                if result and len(result) > 20:
                    return result
            except Exception:
                pass

    # Coba temporal inversion
    for pattern, transform in _TEMPORAL_PATTERNS:
        m = pattern.match(sentence)
        if m:
            try:
                result = transform(m)
                if result and len(result) > 20:
                    return result
            except Exception:
                pass

    return sentence


def _shuffle_paragraph_sentences(text: str) -> str:
    """Rotates supporting sentences inside a paragraph to break Turnitin's n-gram matching.
    Locks the 1st sentence (topic) and last sentence (conclusion) in place.
    Shuffles the middle sentences if there are 3 or more sentences.
    """
    if not text or len(text.strip()) < 50:
        return text
    
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    if len(sentences) < 3:
        return text
        
    first = sentences[0]
    last = sentences[-1]
    middle = sentences[1:-1]
    
    # Acak urutan kalimat penjelas di tengah
    _random.shuffle(middle)
    
    return first + " " + " ".join(middle) + " " + last


# ── SIMILARITY SCORE ESTIMATOR ─────────────────────────────────────────────────
def _estimate_change_score(original: str, paraphrased: str) -> float:
    """Estimate % of words changed between original and paraphrased text.
    Returns value 0.0-100.0. Used as a rough similarity-reduction indicator.
    """
    if not original or not paraphrased:
        return 0.0
    orig_words = re.findall(r'\b\w+\b', original.lower())
    para_words = re.findall(r'\b\w+\b', paraphrased.lower())
    if not orig_words:
        return 0.0
    # Count words that appear in paraphrased but NOT in original (new words added)
    orig_set = set(orig_words)
    changed = sum(1 for w in para_words if w not in orig_set)
    # Score = % of paraphrased words that are "new"
    score = (changed / max(len(para_words), 1)) * 100
    return round(min(score, 100.0), 1)


    # Score = % of paraphrased words that are "new"
    score = (changed / max(len(para_words), 1)) * 100
    return round(min(score, 100.0), 1)


# ── TEMPLATE & DECLARATION CLEANER ───────────────────────────────────────────
_TEMPLATE_JUNK_PATTERNS = [
    r'diajukan\s+untuk\s+memenuhi\s+salah\s+satu\s+syarat',
    r'memperoleh\s+gelar\s+sarjana',
    r'program\s+studi\s+sistem\s+informasi',
    r'menyatakan\s+bahwa\s+skripsi\s+ini\s+adalah\s+karya\s+saya',
    r'pernyataan\s+keaslian\s+skripsi',
    r'tugas\s+akhir\s+ini\s+diajukan\s+sebagai',
    r'persetujuan\s+dosen\s+pembimbing',
    r'persyaratan\s+akademis\s+pada\s+universitas',
    r'dosen\s+pembimbing\s+skripsi',
    r'seluruh\s+isi\s+skripsi\s+ini\s+sepenuhnya\s+menjadi',
    r'hak\s+cipta\s+milik\s+universitas',
]

def _is_template_junk_line(text: str) -> bool:
    """Returns True if the line is administrative or boilerplate template junk that can be skipped."""
    t = text.lower().strip()
    if len(t) < 15:
        return False
    return any(re.search(pat, t) for pat in _TEMPLATE_JUNK_PATTERNS)


# ── HUMAN ALTERNATIVE SPELLING (TYPO EVASION) ───────────────────────────────
# Turnitin is highly rigid on spelling. Using common alternative spellings
# breaks matches entirely while humans easily read them.
_ALTERNATIVE_SPELLINGS = {
    "sistem":       "sistim",
    "analisis":     "analisa",
    "efektif":      "efektip",
    "efisien":      "episien",
    "kualitas":     "kwalitas",
    "metode":       "metoda",
    "objek":        "obyek",
    "aktivitas":    "aktifitas",
    "aktif":        "aktip",
    "provinsi":     "propinsi",
    "standardisasi":"standarisasi",
    "teoretis":     "teoritis",
    "praktik":      "praktek",
    "jadwal":       "jadual",
    "subjek":       "subyek",
    "frekuensi":    "frekwensi",
}

def _apply_alternative_spellings(text: str) -> str:
    """Randomly applies alternative spellings to ~4% of occurrences of matching words."""
    def replacer(m):
        word = m.group(0)
        lower = word.lower()
        if lower in _ALTERNATIVE_SPELLINGS:
            # 4% probability to apply alternative spelling
            if _random.random() < 0.04:
                alt = _ALTERNATIVE_SPELLINGS[lower]
                if word[0].isupper():
                    return alt[0].upper() + alt[1:]
                return alt
        return word
    return re.sub(r'\b\w+\b', replacer, text)


def _is_protected_line(text: str) -> bool:
    """Returns True if this line should NOT be paraphrased at all."""
    t = text.strip().lower()
    # Gambar, Tabel, Bagan, Lampiran, Figure, Chart captions
    if re.match(r'^(gambar|tabel|bagan|lampiran|grafik|diagram|figure|table|chart|grafik)\s+[\dIVXivx]+', t, re.IGNORECASE):
        return True
    # Semua huruf kapital = judul bab / cover
    if text.strip().isupper() and len(text.strip()) > 3:
        return True
    # Judul dokumen (baris dimulai dengan kata kunci metadata kampus)
    meta_kw = ["program studi", "jurusan", "fakultas", "universitas", "nim", "nip",
               "dosen pembimbing", "dosen penguji", "disusun oleh", "disusun:",
               "diajukan oleh", "oleh :", "oleh:", "tahun akademik"]
    if any(t.startswith(kw) for kw in meta_kw):
        return True
    # Baris yang hanya berisi angka / kode
    if re.match(r'^[\d\s\-\./,]+$', text.strip()):
        return True
    # Sitasi inline (Author, YYYY) atau [n]
    if re.match(r'^\([A-Z][^)]+,\s*\d{4}\)$', text.strip()):
        return True
    # Nomor halaman atau header/footer (angka pendek saja)
    if re.match(r'^\d{1,3}$', text.strip()):
        return True
    # Daftar isi entry (titik-titik + halaman): "Kata Pengantar .... v"
    if re.search(r'\.{3,}\s*[ivxIVX\d]+\s*$', text.strip()):
        return True
    # Entry daftar pustaka (dimulai tahun atau nama belakang + koma + tahun)
    if re.match(r'^[A-Z][^,]+,\s*(\d{4}|[A-Z])', text.strip()):
        if len(text.strip()) < 150:  # daftar pustaka biasanya pendek
            return True
    # Footnote / endnote marker
    if re.match(r'^\d+\s+[A-Z]', text.strip()) and len(text.strip()) < 60:
        return True
    # Baris heading BAB
    if re.match(r'^(BAB|CHAPTER)\s+[IVXivx\d]+', text.strip(), re.IGNORECASE):
        return True
    return False


# ── DYNAMIC REWRITE LEVEL: tingkat agresivitas per bagian dokumen ────────────
# 0.0 = jangan ubah apapun | 1.0 = seagresif mungkin
SECTION_AGGRESSION = {
    # Bagian cover & administratif
    "cover":              0.0,
    "halaman judul":      0.0,
    "lembar pengesahan":  0.0,
    "lembar persetujuan": 0.0,
    "pernyataan":         0.0,
    "abstrak":            0.55,
    "abstract":           0.55,
    # Kata pengantar
    "kata pengantar":     0.15,
    "kata sambutan":      0.15,
    "prakata":            0.15,
    # Daftar
    "daftar isi":         0.0,
    "daftar gambar":      0.0,
    "daftar tabel":       0.0,
    "daftar lampiran":    0.0,
    "daftar singkatan":   0.0,
    # BAB I - Pendahuluan  (dinaikkan: latar belakang sering copy-heavy)
    "bab i":              0.72,
    "bab 1":              0.72,
    "pendahuluan":        0.72,
    # BAB II - Tinjauan Pustaka / Landasan Teori (PENYUMBANG SIMILARITY TERBESAR → max agresif)
    "bab ii":             0.97,
    "bab 2":              0.97,
    "tinjauan pustaka":   0.97,
    "landasan teori":     0.97,
    "kajian pustaka":     0.97,
    "kajian literatur":   0.97,
    "tinjauan teoritis":  0.97,
    # BAB III - Metodologi  (dinaikkan: banyak kalimat baku yang perlu diubah)
    "bab iii":            0.78,
    "bab 3":              0.78,
    "metodologi":         0.78,
    "metode penelitian":  0.78,
    "perancangan sistem": 0.78,
    "desain penelitian":  0.78,
    # BAB IV - Hasil & Pembahasan  (dinaikkan: banyak kalimat template)
    "bab iv":             0.70,
    "bab 4":              0.70,
    "hasil dan pembahasan": 0.70,
    "hasil penelitian":   0.70,
    "pembahasan":         0.70,
    # BAB V - Kesimpulan & Saran  (dinaikkan: kalimat penutup sering mirip)
    "bab v":              0.65,
    "bab 5":              0.65,
    "kesimpulan":         0.65,
    "saran":              0.65,
    "penutup":            0.65,
    # Akhir dokumen
    "daftar pustaka":     0.0,
    "daftar referensi":   0.0,
    "lampiran":           0.0,
}



# ── PARAPHRASE SESSION: Consistency Checker + Frequency Penalty + Dynamic Rewrite Level ───
class ParaphraseSession:
    """Document-level session that enforces:
    - synonym_memory  : consistent synonym choices across paragraphs
    - synonym_freq    : frequency penalty to avoid overused synonyms
    - current_aggression: dynamic rewrite level per BAB section (0.0-1.0)
    """
    MAX_FREQ = 4   # max times a single synonym may appear per document (turunkan dari 8 → rotasi lebih cepat)

    def __init__(self):
        self.synonym_memory: dict = {}   # original_word -> chosen_synonym
        self.synonym_freq: dict  = {}    # synonym -> count
        self.current_aggression: float = 0.6   # default moderate
        self.current_section: str = ""          # nama section saat ini (for logging)

    def set_section(self, heading_text: str):
        """Update aggression level based on detected heading/section name.
        Resets synonym_memory when entering a high-aggression BAB (>= 0.7)
        so each BAB gets fresh synonym choices, while frequency counter persists
        to prevent global repetition.
        """
        h = heading_text.lower().strip()
        old_aggression = self.current_aggression
        best_match = None
        for key, level in SECTION_AGGRESSION.items():
            if key in h:
                if best_match is None or len(key) > len(best_match):
                    best_match = key
                    self.current_aggression = level
                    self.current_section = key
        # Jika masuk ke BAB dengan aggression tinggi (BAB II = tinjauan pustaka),
        # reset memory agar bisa menggunakan variasi sinonim yang berbeda
        if best_match and self.current_aggression >= 0.7 and old_aggression < 0.7:
            self.synonym_memory = {}  # fresh choices untuk BAB agresif
        # Jika heading tidak dikenali, pertahankan aggression sebelumnya


    def get_synonym(self, word_lower: str) -> str:
        """Session-aware synonym picker with memory + frequency penalty + dynamic aggression."""
        # Dynamic aggression: sesuaikan threshold HIGH_PROTECT per section
        # aggression=0.0 -> protect_threshold=0.99 (hampir tidak ada perubahan)
        # aggression=0.92 -> protect_threshold=0.54 (agresif, sering berubah)
        if word_lower in ACADEMIC_HIGH_PROTECT:
            protect_threshold = 1.0 - (self.current_aggression * 0.5)
            if _random.random() < protect_threshold:
                return word_lower

        # Jika aggression nol (cover, daftar pustaka), skip synonym replacement
        if self.current_aggression == 0.0:
            return word_lower

        # 1. If we've already chosen a synonym for this word, reuse it
        #    UNLESS that synonym is now overused.
        if word_lower in self.synonym_memory:
            previous = self.synonym_memory[word_lower]
            freq = self.synonym_freq.get(previous, 0)
            if freq < self.MAX_FREQ:
                # still within budget — keep the previous choice
                self.synonym_freq[previous] = freq + 1
                return previous
            # overused — fall through to pick a new one

        choices = INDONESIAN_SYNONYMS_MULTI.get(word_lower)
        if not choices:
            return word_lower

        # 2. Build a list of candidates that are NOT overused
        candidates = []
        for c in choices:
            if c is None:
                candidates.append(word_lower)   # "keep original" option
            elif self.synonym_freq.get(c, 0) < self.MAX_FREQ:
                candidates.append(c)

        # If all synonyms are overused, fall back to keeping original
        if not candidates:
            return word_lower

        chosen = _random.choice(candidates)

        # 3. Store the decision for future paragraphs
        self.synonym_memory[word_lower] = chosen
        self.synonym_freq[chosen] = self.synonym_freq.get(chosen, 0) + 1
        return chosen



def offline_paraphrase(text, session: 'ParaphraseSession | None' = None):
    if not text:
        return text

    # Jika baris ini dilindungi, kembalikan as-is
    if _is_protected_line(text):
        return text

    # Jika baris ini adalah kalimat administrasi / templat skripsi, skip/kembalikan kosong
    # (agar tidak masuk ke output dokumen parafrase jika sangat mirip template kampus)
    if _is_template_junk_line(text):
        return ""

    sentences = re.split(r'(?<=[.!?])\s+', text)
    processed_sentences = []

    # Ambil aggression dari session (untuk restructurer)
    aggression = session.current_aggression if session else 0.6

    for sentence in sentences:
        if not sentence.strip():
            continue

        # === PROTEKSI: Kutipan langsung & sitasi inline — skip sepenuhnya ===
        if _is_direct_quote_or_citation(sentence):
            processed_sentences.append(sentence)
            continue

        # STEP 0: Parafrase level frasa dulu (paling kuat)
        temp2 = sentence
        for phrase_orig, phrase_para in sorted(PHRASE_PARAPHRASE.items(), key=lambda x: len(x[0]), reverse=True):
            pattern = re.compile(re.escape(phrase_orig), re.IGNORECASE)
            def phrase_replacer(m, orig=phrase_orig, para=phrase_para):
                matched = m.group(0)
                # Preserve capitalization of first letter
                if matched[0].isupper():
                    return para[0].upper() + para[1:]
                return para
            temp2 = pattern.sub(phrase_replacer, temp2)

        # Re-do proteksi konjungsi dengan cara yang lebih bersih
        conj_holders2 = {}
        for ci, conj in enumerate(sorted(CONJUNCTIONS_PROTECT, key=len, reverse=True)):
            key = f"__CONJ_{ci}__"
            new_t = re.sub(re.escape(conj), key, temp2, flags=re.IGNORECASE)
            if new_t != temp2:
                conj_holders2[key] = conj
                temp2 = new_t

        # Lindungi exclude phrases
        placeholders = {}
        sorted_phrases = sorted(EXCLUDE_PHRASES, key=len, reverse=True)
        for idx, phrase in enumerate(sorted_phrases):
            pattern = re.compile(r'\b' + re.escape(phrase) + r'\b', re.IGNORECASE)
            def replacer(m, k=f"__EX_{idx}_{id(phrase)}__"):
                placeholders[k] = m.group(0)
                return k
            temp2 = pattern.sub(replacer, temp2)

        # Ganti sinonim kata per kata (HANYA kata lengkap, case-aware, session-aware)
        def replace_word(m):
            word = m.group(0)
            lower = word.lower()
            # Use session-aware picker if session is provided, else plain random
            syn = session.get_synonym(lower) if session else _get_synonym(lower)
            if syn == lower:
                return word  # pertahankan kata asli
            # Preserve capitalization
            if word[0].isupper():
                return syn[0].upper() + syn[1:]
            elif word.isupper() and len(word) > 1:
                return syn.upper()
            return syn

        result = re.sub(r'\b\w+(?:-\w+)*\b', replace_word, temp2)

        # Kembalikan konjungsi
        for key, val in conj_holders2.items():
            result = result.replace(key, val)

        # Kembalikan exclude phrases
        for key, val in placeholders.items():
            result = result.replace(key, val)

        # --- Post-processing perbaikan spasi & duplikasi ---
        # Hapus duplikasi kata sambung, cth: "oleh oleh karena itu"
        result = re.sub(r'\b(\w+)\s+\1\b', r'\1', result)
        # Perbaiki spasi di sekitar tanda baca
        result = re.sub(r'\s+([,.:;!?])', r'\1', result)
        result = re.sub(r'([,.:;])\s*([^\s])', r'\1 \2', result)
        # Perbaiki spasi antara kata dan tanda kurung
        result = re.sub(r'\(\s+', '(', result)
        result = re.sub(r'\s+\)', ')', result)
        # Hapus spasi ganda
        result = re.sub(r'\s{2,}', ' ', result).strip()

        # --- COLLOCATION VALIDATOR: auto-koreksi frasa tidak natural ---
        result = _apply_collocation_fix(result)

        # --- HUMANIZATION LAYER: Acak pola pembuka kalimat ---
        result = _humanize_opener(result)

        # --- NUMBER → WORDS: Ubah angka kecil menjadi kata (probabilistik) ---
        result = _number_to_words_in_sentence(result)

        # --- SENTENCE RESTRUCTURER: Ubah struktur kalimat (passive↔active, inversion) ---
        result = _restructure_sentence(result, aggression=aggression)

        # --- ALTERNATIVE SPELLINGS: Gunakan variasi ejaan tidak baku (4% rate) ---
        result = _apply_alternative_spellings(result)

        processed_sentences.append(result)

    paragraph_text = " ".join(processed_sentences)
    
    # --- SENTENCE SHUFFLER: Acak posisi kalimat pendukung di dalam paragraf ---
    if aggression >= 0.7:  # Hanya untuk Bab dengan tingkat agresivitas tinggi
        paragraph_text = _shuffle_paragraph_sentences(paragraph_text)

    return paragraph_text


app = Flask(__name__)
if os.environ.get('VERCEL'):
    UPLOAD_FOLDER = '/tmp'
else:
    UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

PARAPHRASE_SYSTEM_PROMPT = """Anda adalah asisten akademik profesional yang berspesialisasi dalam menulis ulang karya ilmiah (skripsi/tesis) untuk menurunkan persentase kemiripan (Turnitin) di tingkat universitas.

MASTER RULES - PARAFRASE BERKUALITAS (ANTI-HALUSINASI)

1. Selalu pertahankan makna asli 100%.
2. Jangan pernah menambahkan informasi baru.
3. Jangan pernah mengurangi informasi penting.
4. Jangan membuat asumsi.
5. Jangan berhalusinasi.
6. Jangan mengarang fakta.
7. Jangan mengubah kesimpulan penulis.
8. Jangan mengubah tujuan kalimat.
9. Jangan mengubah konteks pembahasan.
10. Jangan mengubah hubungan sebab akibat.

11. Parafrase seluruh kalimat, bukan hanya mengganti sinonim.
12. Ubah struktur kalimat secara menyeluruh.
13. Ubah pola penyampaian ide.
14. Ubah urutan frasa bila tetap alami.
15. Gunakan variasi struktur aktif dan pasif bila sesuai.
16. Pecah kalimat panjang bila diperlukan.
17. Gabungkan kalimat pendek bila lebih alami.
18. Hindari pola kalimat yang identik.
19. Hindari mempertahankan susunan frasa asli.
20. Gunakan variasi konjungsi.

21. Hindari sinonim yang terdengar dipaksakan.
22. Gunakan bahasa Indonesia yang alami.
23. Gunakan diksi akademik yang wajar.
24. Hindari kata-kata berlebihan.
25. Hindari repetisi.
26. Hindari kalimat kaku.
27. Hindari gaya bahasa AI.
28. Hindari frasa klise.
29. Hindari pengulangan pola.
30. Hindari kalimat yang terlalu panjang.

31. Pertahankan semua angka.
32. Pertahankan seluruh persentase.
33. Pertahankan seluruh tanggal.
34. Pertahankan seluruh tahun.
35. Pertahankan seluruh satuan.
36. Pertahankan seluruh simbol matematika.
37. Pertahankan seluruh rumus.
38. Pertahankan seluruh variabel.
39. Pertahankan seluruh kode.
40. Pertahankan seluruh identifier.

41. Jangan mengubah nama orang.
42. Jangan mengubah nama penulis.
43. Jangan mengubah nama institusi.
44. Jangan mengubah nama perusahaan.
45. Jangan mengubah nama organisasi.
46. Jangan mengubah nama aplikasi.
47. Jangan mengubah nama framework.
48. Jangan mengubah nama library.
49. Jangan mengubah nama database.
50. Jangan mengubah nama algoritma.

51. Jangan mengubah istilah teknis yang tidak memiliki padanan resmi.
52. Jika memiliki padanan resmi, gunakan secara konsisten.
53. Jangan menerjemahkan istilah teknis sembarangan.
54. Pertahankan singkatan resmi.
55. Pertahankan akronim resmi.
56. Pertahankan nama metode.
57. Pertahankan nama model.
58. Pertahankan nama standar.
59. Pertahankan nama protokol.
60. Pertahankan nomenklatur ilmiah.

61. Jangan mengubah nomor referensi.
62. Jangan mengubah format sitasi.
63. Jangan menghapus sitasi.
64. Jangan menambah sitasi baru.
65. Jangan memindahkan posisi sitasi bila mengubah makna.
66. Jangan mengubah urutan daftar referensi.
67. Jangan memalsukan referensi.
68. Jangan membuat DOI palsu.
69. Jangan membuat jurnal palsu.
70. Jangan membuat kutipan palsu.

71. Jika teks ambigu, pertahankan ambiguitasnya.
72. Jangan memperjelas informasi yang tidak ada.
73. Jangan menyederhanakan fakta penting.
74. Jangan memperluas pembahasan.
75. Jangan memberikan opini.
76. Jangan memberikan interpretasi.
77. Jangan menyisipkan saran.
78. Jangan menambahkan contoh baru.
79. Jangan menambahkan ilustrasi.
80. Jangan menambahkan analogi.

81. Pertahankan hubungan logis antar kalimat.
82. Pertahankan urutan argumen.
83. Pertahankan alur pembahasan.
84. Pertahankan fokus paragraf.
85. Pertahankan topik utama.
86. Pertahankan detail penting.
87. Pertahankan ruang lingkup pembahasan.
88. Pertahankan definisi.
89. Pertahankan klasifikasi.
90. Pertahankan terminologi.

91. Gunakan ejaan sesuai PUEBI / EYD V.
92. Perbaiki typo bila ada.
93. Perbaiki tanda baca bila diperlukan.
94. Hindari kapitalisasi yang salah.
95. Hindari spasi ganda.
96. Hindari karakter asing yang tidak perlu.
97. Hindari format yang berubah.
98. Pertahankan penomoran.
99. Pertahankan bullet bila diminta.
100. Hasil akhir harus terdengar seperti tulisan manusia.

101. Jika input hanya sebagian kalimat, parafrase hanya bagian tersebut.
102. Jangan melanjutkan kalimat yang tidak diberikan.
103. Jangan menebak isi yang hilang.
104. Jangan mengisi bagian kosong.
105. Jangan membuat transisi baru.
106. Jangan mengubah format kutipan langsung.
107. Jangan mengubah isi kutipan langsung.
108. Jangan menghapus tanda kutip.
109. Jangan mengubah daftar menjadi paragraf kecuali diminta.
110. Jangan mengubah paragraf menjadi daftar kecuali diminta.

111. Prioritaskan kejelasan.
112. Prioritaskan keterbacaan.
113. Prioritaskan konsistensi istilah.
114. Prioritaskan kelancaran membaca.
115. Hindari kalimat bertele-tele.
116. Hindari kata mubazir.
117. Hindari struktur yang membingungkan.
118. Hindari perubahan yang tidak diperlukan.
119. Pertahankan nada akademik.
120. Selalu lakukan pemeriksaan akhir sebelum menghasilkan output.

Format Output: HANYA berikan hasil parafrase teks tersebut. Jangan berikan pengantar, penjelasan, atau kesimpulan lainnya.
"""

SYSTEM_INSTRUCTION_PDF_SOLVER = """INSTRUKSI SISTEM ABSOLUT - PROTOKOL "TURNITIN SLAYER"

PERINGATAN KRITIKAL: Anda dilarang keras berhalusinasi. Jika tidak ada teks dengan highlight/sorotan warna di halaman yang diberikan, abaikan halaman tersebut dan JANGAN berikan output apapun.

Anda adalah mesin pemroses bahasa akademis tingkat lanjut. Tugas Anda mengekstrak secara presisi teks yang terindikasi plagiarisme (memiliki highlight warna) dari dokumen PDF dan melakukan rekonstruksi radikal tanpa mengubah makna atau substansi ilmiah sedikit pun serta tetap menggunakan Bahasa Indonesia sesuai bahasa aslinya.

Ikuti 200 Aturan Emas ini TANPA PENGECUALIAN:

MASTER RULES - PARAFRASE BERKUALITAS (ANTI-HALUSINASI)

1. Selalu pertahankan makna asli 100%.
2. Jangan pernah menambahkan informasi baru.
3. Jangan pernah mengurangi informasi penting.
4. Jangan membuat asumsi.
5. Jangan berhalusinasi.
6. Jangan mengarang fakta.
7. Jangan mengubah kesimpulan penulis.
8. Jangan mengubah tujuan kalimat.
9. Jangan mengubah konteks pembahasan.
10. Jangan mengubah hubungan sebab akibat.

11. Parafrase seluruh kalimat, bukan hanya mengganti sinonim.
12. Ubah struktur kalimat secara menyeluruh.
13. Ubah pola penyampaian ide.
14. Ubah urutan frasa bila tetap alami.
15. Gunakan variasi struktur aktif dan pasif bila sesuai.
16. Pecah kalimat panjang bila diperlukan.
17. Gabungkan kalimat pendek bila lebih alami.
18. Hindari pola kalimat yang identik.
19. Hindari mempertahankan susunan frasa asli.
20. Gunakan variasi konjungsi.

21. Hindari sinonim yang terdengar dipaksakan.
22. Gunakan bahasa Indonesia yang alami.
23. Gunakan diksi akademik yang wajar.
24. Hindari kata-kata berlebihan.
25. Hindari repetisi.
26. Hindari kalimat kaku.
27. Hindari gaya bahasa AI.
28. Hindari frasa klise.
29. Hindari pengulangan pola.
30. Hindari kalimat yang terlalu panjang.

31. Pertahankan semua angka.
32. Pertahankan seluruh persentase.
33. Pertahankan seluruh tanggal.
34. Pertahankan seluruh tahun.
35. Pertahankan seluruh satuan.
36. Pertahankan seluruh simbol matematika.
37. Pertahankan seluruh rumus.
38. Pertahankan seluruh variabel.
39. Pertahankan seluruh kode.
40. Pertahankan seluruh identifier.

41. Jangan mengubah nama orang.
42. Jangan mengubah nama penulis.
43. Jangan mengubah nama institusi.
44. Jangan mengubah nama perusahaan.
45. Jangan mengubah nama organisasi.
46. Jangan mengubah nama aplikasi.
47. Jangan mengubah nama framework.
48. Jangan mengubah nama library.
49. Jangan mengubah nama database.
50. Jangan mengubah nama algoritma.

51. Jangan mengubah istilah teknis yang tidak memiliki padanan resmi.
52. Jika memiliki padanan resmi, gunakan secara konsisten.
53. Jangan menerjemahkan istilah teknis sembarangan.
54. Pertahankan singkatan resmi.
55. Pertahankan akronim resmi.
56. Pertahankan nama metode.
57. Pertahankan nama model.
58. Pertahankan nama standar.
59. Pertahankan nama protokol.
60. Pertahankan nomenklatur ilmiah.

61. Jangan mengubah nomor referensi.
62. Jangan mengubah format sitasi.
63. Jangan menghapus sitasi.
64. Jangan menambah sitasi baru.
65. Jangan memindahkan posisi sitasi bila mengubah makna.
66. Jangan mengubah urutan daftar referensi.
67. Jangan memalsukan referensi.
68. Jangan membuat DOI palsu.
69. Jangan membuat jurnal palsu.
70. Jangan membuat kutipan palsu.

71. Jika teks ambigu, pertahankan ambiguitasnya.
72. Jangan memperjelas informasi yang tidak ada.
73. Jangan menyederhanakan fakta penting.
74. Jangan memperluas pembahasan.
75. Jangan memberikan opini.
76. Jangan memberikan interpretasi.
77. Jangan menyisipkan saran.
78. Jangan menambahkan contoh baru.
79. Jangan menambahkan ilustrasi.
80. Jangan menambahkan analogi.

81. Pertahankan hubungan logis antar kalimat.
82. Pertahankan urutan argumen.
83. Pertahankan alur pembahasan.
84. Pertahankan fokus paragraf.
85. Pertahankan topik utama.
86. Pertahankan detail penting.
87. Pertahankan ruang lingkup pembahasan.
88. Pertahankan definisi.
89. Pertahankan klasifikasi.
90. Pertahankan terminologi.

91. Gunakan ejaan sesuai PUEBI / EYD V.
92. Perbaiki typo bila ada.
93. Perbaiki tanda baca bila diperlukan.
94. Hindari kapitalisasi yang salah.
95. Hindari spasi ganda.
96. Hindari karakter asing yang tidak perlu.
97. Hindari format yang berubah.
98. Pertahankan penomoran.
99. Pertahankan bullet bila diminta.
100. Hasil akhir harus terdengar seperti tulisan manusia.

101. Jika input hanya sebagian kalimat, parafrase hanya bagian tersebut.
102. Jangan melanjutkan kalimat yang tidak diberikan.
103. Jangan menebak isi yang hilang.
104. Jangan mengisi bagian kosong.
105. Jangan membuat transisi baru.
106. Jangan mengubah format kutipan langsung.
107. Jangan mengubah isi kutipan langsung.
108. Jangan menghapus tanda kutip.
109. Jangan mengubah daftar menjadi paragraf kecuali diminta.
110. Jangan mengubah paragraf menjadi daftar kecuali diminta.

111. Prioritaskan kejelasan.
112. Prioritaskan keterbacaan.
113. Prioritaskan konsistensi istilah.
114. Prioritaskan kelancaran membaca.
115. Hindari kalimat bertele-tele.
116. Hindari kata mubazir.
117. Hindari struktur yang membingungkan.
118. Hindari perubahan yang tidak diperlukan.
119. Pertahankan nada akademik.
120. Selalu lakukan pemeriksaan akhir sebelum menghasilkan output.

=========================
DOCUMENT PRESERVATION RULES
=========================

121. Pertahankan seluruh struktur dokumen.
122. Jangan mengubah layout halaman.
123. Jangan mengubah ukuran kertas.
124. Jangan mengubah margin.
125. Jangan mengubah orientasi halaman.
126. Jangan mengubah section break.
127. Jangan mengubah page break.
128. Jangan mengubah header.
129. Jangan mengubah footer.
130. Jangan mengubah nomor halaman.

131. Pertahankan seluruh format font.
132. Pertahankan jenis font.
133. Pertahankan ukuran font.
134. Pertahankan warna font.
135. Pertahankan Bold.
136. Pertahankan Italic.
137. Pertahankan Underline.
138. Pertahankan Highlight.
139. Pertahankan Strikethrough.
140. Pertahankan seluruh efek teks.

141. Pertahankan seluruh alignment.
142. Pertahankan indentasi.
143. Pertahankan line spacing.
144. Pertahankan paragraph spacing.
145. Pertahankan tab stop.
146. Pertahankan numbering.
147. Pertahankan bullet.
148. Pertahankan multilevel list.
149. Pertahankan style heading.
150. Pertahankan style bawaan dokumen.

151. Jangan menghapus gambar.
152. Jangan mengganti gambar.
153. Jangan memindahkan posisi gambar.
154. Jangan mengubah ukuran gambar.
155. Jangan mengubah resolusi gambar.
156. Jangan mengubah caption gambar.
157. Jangan mengubah SmartArt.
158. Jangan mengubah Shape.
159. Jangan mengubah WordArt.
160. Jangan mengubah watermark.

161. Pertahankan seluruh tabel.
162. Jangan mengubah jumlah kolom.
163. Jangan mengubah jumlah baris.
164. Jangan merge cell.
165. Jangan split cell.
166. Jangan mengubah border tabel.
167. Jangan mengubah shading tabel.
168. Jangan mengubah ukuran tabel.
169. Jangan mengubah posisi tabel.
170. Jangan mengubah isi tabel kecuali teks target.

171. Pertahankan seluruh hyperlink.
172. Pertahankan seluruh bookmark.
173. Pertahankan seluruh cross-reference.
174. Pertahankan seluruh footnote.
175. Pertahankan seluruh endnote.
176. Pertahankan seluruh comment.
177. Pertahankan seluruh field Word.
178. Pertahankan seluruh equation.
179. Pertahankan seluruh object.
180. Pertahankan seluruh embedded file.

181. Jangan mengubah TOC (Table of Contents).
182. Jangan mengubah daftar gambar.
183. Jangan mengubah daftar tabel.
184. Jangan mengubah index.
185. Jangan mengubah citation manager.
186. Jangan mengubah bibliography.
187. Jangan mengubah style referensi.
188. Jangan mengubah format daftar pustaka.
189. Jangan mengubah nomor referensi.
190. Jangan mengubah urutan referensi.

191. Hanya ubah teks yang memang ditargetkan.
192. Seluruh bagian lain WAJIB identik.
193. Jangan menghapus whitespace yang penting.
194. Jangan menambah halaman baru.
195. Jangan menghapus halaman.
196. Jangan mengubah urutan halaman.
197. Jangan mengubah posisi objek.
198. Jangan mengubah posisi caption.
199. Jangan mengubah format dokumen.
200. Output akhir harus semirip mungkin dengan dokumen asli selain isi teks yang diparafrase.

Format Output:
Wajib memberikan output dengan format Markdown berikut untuk setiap temuan plagiasi (tanpa pengantar, penutup, atau komentar apa pun):

### Halaman [Nomor Halaman]
**Teks Asli:** [Teks asli yang tersorot plagiasi di PDF]
**Hasil Parafrase:** [Teks baru hasil rekonstruksi/parafrase]
"""

HTML_TEMPLATE = """
<html lang="id">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=1.0, user-scalable=no">
    <title>Ruang Kreasi - Academic Writing & Paraphrasing Suite</title>
    <link href="https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700;800&family=Plus+Jakarta+Sans:wght@300;400;500;600;700&display=swap" rel="stylesheet">
    <style>
        :root {
            --bg-color: #fafafa;
            --container-bg: #ffffff;
            --card-bg: #ffffff;
            --text-color: #111827;
            --secondary-text: #6b7280;
            --border-color: #e5e7eb;
            --border-hover: #9ca3af;
            --hover-bg: #f3f4f6;
            --accent-color: #111827;
            --accent-gradient: #111827;
            --button-gradient: #111827;
            --button-hover: #374151;
            --card-shadow: 0 1px 3px rgba(0, 0, 0, 0.05), 0 1px 2px rgba(0, 0, 0, 0.02);
            --sidebar-width: 260px;
            --storage-width: 340px;
        }

        /* Menyembunyikan scrollbar secara global */
        * {
            scrollbar-width: none !important;
            -ms-overflow-style: none !important;
        }
        *::-webkit-scrollbar {
            display: none !important;
            width: 0 !important;
            height: 0 !important;
        }

        body.dark-mode {
            --bg-color: #0b0c0e;
            --container-bg: #121316;
            --card-bg: #121316;
            --text-color: #f3f4f6;
            --secondary-text: #9ca3af;
            --border-color: #1f2937;
            --border-hover: #374151;
            --hover-bg: #1e2025;
            --accent-color: #ffffff;
            --accent-gradient: #ffffff;
            --button-gradient: #ffffff;
            --button-hover: #e5e7eb;
            --card-shadow: 0 4px 20px rgba(0, 0, 0, 0.4);
        }

        html, body {
            margin: 0;
            padding: 0;
            width: 100%;
            height: 100%;
            overflow: hidden;
            background-color: var(--bg-color);
            color: var(--text-color);
            font-family: 'Plus Jakarta Sans', -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
            transition: background-color 0.3s, color 0.3s;
        }

        .main-layout {
            display: grid;
            grid-template-columns: var(--sidebar-width) 1fr var(--storage-width);
            width: 100%;
            height: 100%;
            overflow: hidden;
        }

        /* Mobile top navbar - hidden on desktop */
        .mobile-topbar {
            display: none;
        }

        /* Left Sidebar Styling */
        .left-sidebar {
            background-color: var(--container-bg);
            border-right: 1px solid var(--border-color);
            display: flex;
            flex-direction: column;
            justify-content: space-between;
            height: 100%;
            padding: 24px;
            box-sizing: border-box;
            overflow-y: auto;
        }

        .left-sidebar.collapsed {
            width: 0px !important;
            min-width: 0px !important;
            max-width: 0px !important;
            padding: 0px !important;
            margin: 0px !important;
            border: none !important;
            overflow: hidden !important;
        }

        .right-sidebar.collapsed {
            width: 0px !important;
            min-width: 0px !important;
            max-width: 0px !important;
            padding: 0px !important;
            margin: 0px !important;
            border: none !important;
            overflow: hidden !important;
        }

        /* Loading Overlay Styling */
        .loading-overlay {
            display: none;
            position: fixed;
            top: 0;
            left: 0;
            width: 100vw;
            height: 100vh;
            background: rgba(255, 255, 255, 0.7);
            backdrop-filter: blur(8px);
            -webkit-backdrop-filter: blur(8px);
            z-index: 9999;
            align-items: center;
            justify-content: center;
            opacity: 0;
            transition: opacity 0.3s ease;
        }

        body.dark-mode .loading-overlay {
            background: rgba(11, 12, 14, 0.85);
        }

        .loading-overlay.active {
            display: flex;
            opacity: 1;
        }

        .loading-container {
            text-align: center;
            display: flex;
            flex-direction: column;
            align-items: center;
            gap: 20px;
        }

        /* Premium pulsing glow spinner */
        .glow-spinner {
            width: 80px;
            height: 80px;
            border: 4px solid transparent;
            border-top: 4px solid var(--text-color);
            border-radius: 50%;
            animation: spin 1s cubic-bezier(0.5, 0.1, 0.1, 0.9) infinite;
            position: relative;
            box-shadow: 0 0 15px rgba(17, 24, 39, 0.1);
        }

        body.dark-mode .glow-spinner {
            border-top-color: #f3f4f6;
            box-shadow: 0 0 25px rgba(243, 244, 246, 0.2);
        }

        .glow-spinner::before {
            content: '';
            position: absolute;
            top: 6px;
            left: 6px;
            right: 6px;
            bottom: 6px;
            border: 4px solid transparent;
            border-bottom: 4px solid var(--secondary-text);
            border-radius: 50%;
            animation: spin-reverse 1.5s linear infinite;
        }

        .loading-title {
            font-family: 'Outfit', sans-serif;
            font-weight: 700;
            font-size: 1.5rem;
            color: var(--text-color);
            margin: 0;
            letter-spacing: -0.02em;
        }

        .loading-subtitle {
            font-family: 'Plus Jakarta Sans', sans-serif;
            font-weight: 500;
            color: var(--secondary-text);
            font-size: 1rem;
            animation: pulse-opacity 2s infinite ease-in-out;
            min-height: 24px;
        }

        @keyframes spin {
            0% { transform: rotate(0deg); }
            100% { transform: rotate(360deg); }
        }

        @keyframes spin-reverse {
            0% { transform: rotate(360deg); }
            100% { transform: rotate(0deg); }
        }

        @keyframes pulse-opacity {
            0%, 100% { opacity: 0.6; }
            50% { opacity: 1; }
        }

        .logo-area {
            display: flex;
            align-items: center;
            font-family: 'Outfit', sans-serif;
            font-weight: 800;
            font-size: 1.25rem;
            color: var(--text-color);
            margin-bottom: 35px;
            gap: 10px;
        }

        .logo-area svg {
            color: var(--text-color);
        }

        .nav-menu {
            display: flex;
            flex-direction: column;
            gap: 6px;
        }

        .nav-link {
            display: flex;
            align-items: center;
            gap: 12px;
            padding: 12px 16px;
            color: var(--secondary-text);
            text-decoration: none;
            font-size: 0.9rem;
            font-weight: 600;
            border-radius: 10px;
            transition: all 0.2s ease;
            cursor: pointer;
        }

        .nav-link:hover {
            color: var(--text-color);
            background-color: var(--hover-bg);
        }

        .nav-link.active {
            color: var(--text-color);
            background-color: var(--hover-bg);
        }

        .sidebar-bottom {
            margin-top: auto;
            padding-top: 24px;
            border-top: 1px solid var(--border-color);
        }

        .user-profile {
            display: flex;
            align-items: center;
            gap: 12px;
            margin-bottom: 20px;
        }

        .avatar {
            width: 40px;
            height: 40px;
            border-radius: 50%;
            background-color: var(--border-color);
            display: flex;
            align-items: center;
            justify-content: center;
            font-weight: 700;
            color: var(--text-color);
        }

        .user-info {
            flex: 1;
        }

        .user-name {
            font-weight: 700;
            font-size: 0.9rem;
        }

        .user-plan {
            font-size: 0.75rem;
            color: var(--secondary-text);
        }

        .storage-usage {
            display: flex;
            flex-direction: column;
            gap: 8px;
        }

        .storage-text {
            display: flex;
            justify-content: space-between;
            font-size: 0.75rem;
            color: var(--secondary-text);
            font-weight: 600;
        }

        .progress-bar {
            width: 100%;
            height: 6px;
            background-color: var(--border-color);
            border-radius: 3px;
            overflow: hidden;
        }

        .progress {
            height: 100%;
            background-color: var(--text-color);
            border-radius: 3px;
        }

        /* Center Content Styling */
        .main-content {
            height: 100%;
            overflow-y: auto;
            padding: 40px;
            box-sizing: border-box;
        }

        .header {
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 35px;
        }

        h1 {
            font-family: 'Outfit', sans-serif;
            font-weight: 800;
            font-size: 2.2rem;
            margin: 0;
            letter-spacing: -0.03em;
            color: var(--text-color);
        }

        p.subtitle {
            color: var(--secondary-text);
            margin: 5px 0 0 0;
            font-size: 1.05rem;
            font-weight: 500;
        }

        .theme-toggle-btn {
            background: none;
            border: 1px solid var(--border-color);
            width: 42px;
            height: 42px;
            border-radius: 50%;
            display: flex;
            align-items: center;
            justify-content: center;
            color: var(--text-color);
            cursor: pointer;
            transition: all 0.2s ease;
        }

        .theme-toggle-btn:hover {
            background-color: var(--hover-bg);
            border-color: var(--border-hover);
        }

        .tabs {
            display: flex;
            border-bottom: 1px solid var(--border-color);
            margin-bottom: 35px;
            gap: 6px;
            overflow-x: auto;
            padding-bottom: 4px;
        }

        .tab-btn {
            background: none;
            border: 1px solid transparent;
            padding: 10px 18px;
            font-size: 0.85rem;
            font-weight: 600;
            color: var(--secondary-text);
            cursor: pointer;
            transition: all 0.2s ease;
            white-space: nowrap;
            border-radius: 10px;
            display: flex;
            align-items: center;
            gap: 8px;
        }

        .tab-btn:hover {
            color: var(--text-color);
            background: var(--hover-bg);
        }

        .tab-btn.active {
            color: var(--container-bg);
            background: var(--text-color);
            border: 1px solid var(--text-color);
        }

        .tab-content {
            display: none;
        }

        .tab-content.active {
            display: block;
        }

        .card {
            background: var(--container-bg);
            border: 1px solid var(--border-color);
            border-radius: 16px;
            padding: 40px;
            box-shadow: var(--card-shadow);
            margin-bottom: 30px;
        }

        .form-group {
            margin-bottom: 28px;
        }

        label {
            display: block;
            font-weight: 700;
            margin-bottom: 8px;
            font-size: 0.75rem;
            color: var(--text-color);
            text-transform: uppercase;
            letter-spacing: 0.08em;
        }

        .api-input-container {
            display: flex;
            gap: 12px;
            align-items: center;
        }

        input[type="text"], input[type="number"], select {
            flex-grow: 1;
            padding: 12px 16px;
            background: var(--container-bg);
            border: 1px solid var(--border-color);
            border-radius: 10px;
            color: var(--text-color);
            font-family: inherit;
            box-sizing: border-box;
            font-size: 0.95rem;
            transition: all 0.2s ease;
        }

        input[type="text"]:focus, input[type="number"]:focus, select:focus {
            outline: none;
            border-color: var(--text-color);
        }

        .btn-check-api {
            padding: 12px 20px;
            border: 1px solid var(--border-color);
            background: var(--container-bg);
            color: var(--text-color);
            font-weight: 700;
            font-size: 0.85rem;
            border-radius: 10px;
            cursor: pointer;
            transition: all 0.2s ease;
            white-space: nowrap;
        }

        .btn-check-api:hover {
            background: var(--hover-bg);
            border-color: var(--border-hover);
        }

        .api-status-badge {
            margin-top: 10px;
            font-size: 0.85rem;
            font-weight: 600;
            display: none;
            padding: 10px 14px;
            border-radius: 10px;
            border: 1px solid transparent;
        }

        .file-upload-wrapper {
            position: relative;
            border: 2px dashed var(--border-color);
            border-radius: 12px;
            padding: 35px 20px;
            text-align: center;
            background: var(--bg-color);
            transition: all 0.2s ease;
            cursor: pointer;
        }

        .file-upload-wrapper:hover {
            background: var(--hover-bg);
            border-color: var(--border-hover);
        }

        .file-upload-wrapper input[type="file"] {
            position: absolute;
            left: 0;
            top: 0;
            opacity: 0;
            width: 100%;
            height: 100%;
            cursor: pointer;
        }

        .upload-icon {
            font-size: 2rem;
            margin-bottom: 10px;
            display: block;
            transition: transform 0.2s ease;
        }

        .file-name-label {
            color: var(--secondary-text);
            font-size: 0.85rem;
            font-weight: 600;
            word-break: break-all;
        }

        .btn-submit {
            display: flex;
            align-items: center;
            justify-content: center;
            gap: 10px;
            width: 100%;
            padding: 14px;
            border: 1px solid var(--text-color);
            border-radius: 10px;
            background: var(--text-color);
            color: var(--container-bg);
            font-weight: 700;
            font-size: 0.95rem;
            cursor: pointer;
            transition: all 0.2s ease;
            margin-top: 15px;
        }

        .btn-submit:hover {
            background: var(--button-hover);
            border-color: var(--button-hover);
        }

        .btn-submit.loading {
            background: var(--border-color);
            border-color: var(--border-color);
            color: var(--secondary-text);
            cursor: not-allowed;
        }

        .spinner {
            width: 16px;
            height: 16px;
            border: 2px solid rgba(0, 0, 0, 0.2);
            border-top-color: var(--text-color);
            border-radius: 50%;
            animation: spin 0.8s linear infinite;
            display: inline-block;
        }

        @keyframes spin {
            to { transform: rotate(360deg); }
        }

        .btn-download {
            display: inline-flex;
            align-items: center;
            gap: 8px;
            padding: 12px 24px;
            border: 1px solid var(--text-color);
            border-radius: 10px;
            background: var(--container-bg);
            color: var(--text-color);
            font-weight: 700;
            font-size: 0.9rem;
            text-decoration: none;
            transition: all 0.2s ease;
        }

        .btn-download:hover {
            background: var(--hover-bg);
        }

        .alert-box {
            background: #eafaf1;
            border: 1px solid #d4ebdf;
            color: #0f5132;
            border-radius: 10px;
            padding: 14px;
            margin-bottom: 25px;
            font-size: 0.85rem;
            font-weight: 600;
            display: flex;
            align-items: center;
            gap: 10px;
        }

        .rules-title {
            font-size: 0.85rem;
            font-weight: 800;
            letter-spacing: 0.05em;
            color: var(--text-color);
            text-transform: uppercase;
            margin-bottom: 12px;
            margin-top: 15px;
        }

        .rules-container {
            border: 1px solid var(--border-color);
            background: var(--container-bg);
            padding: 20px;
            border-radius: 12px;
            font-size: 0.85rem;
            color: var(--secondary-text);
        }
        
        .rules-container ul {
            margin: 0;
            padding-left: 20px;
        }

        .rules-container li {
            margin-bottom: 12px;
            line-height: 1.5;
        }

        .features-grid {
            display: grid;
            grid-template-columns: repeat(4, 1fr);
            gap: 16px;
            margin-top: 30px;
        }

        .feature-card {
            border: 1px solid var(--border-color);
            background: var(--container-bg);
            padding: 20px;
            border-radius: 12px;
            text-align: left;
            transition: all 0.2s ease;
        }

        .feature-card:hover {
            border-color: var(--border-hover);
        }

        .feature-icon {
            font-size: 1.5rem;
            margin-bottom: 10px;
            display: block;
        }

        .feature-title {
            font-weight: 700;
            font-size: 0.85rem;
            margin-bottom: 4px;
        }

        .feature-desc {
            font-size: 0.75rem;
            color: var(--secondary-text);
            line-height: 1.4;
        }

        /* Right Sidebar Styling */
        .right-sidebar {
            background-color: var(--container-bg);
            border-left: 1px solid var(--border-color);
            height: 100%;
            padding: 24px;
            box-sizing: border-box;
            display: flex;
            flex-direction: column;
            overflow-y: auto;
        }

        .sidebar-header h2 {
            font-family: 'Outfit', sans-serif;
            font-size: 1.15rem;
            margin: 0 0 8px 0;
            font-weight: 700;
            color: var(--text-color);
        }

        .sidebar-header p {
            margin: 0 0 25px 0;
            font-size: 0.8rem;
            color: var(--secondary-text);
            line-height: 1.5;
        }

        .sidebar-section {
            margin-bottom: 25px;
        }

        .sidebar-section h3 {
            font-size: 0.75rem;
            text-transform: uppercase;
            letter-spacing: 0.08em;
            color: var(--text-color);
            margin: 0 0 12px 0;
            font-weight: 800;
            border-left: 3px solid var(--text-color);
            padding-left: 8px;
        }

        .file-list {
            display: flex;
            flex-direction: column;
            gap: 10px;
        }

        .file-item {
            display: flex;
            justify-content: space-between;
            align-items: center;
            background: var(--bg-color);
            border: 1px solid var(--border-color);
            border-radius: 10px;
            padding: 8px 12px;
            font-size: 0.85rem;
            transition: all 0.2s ease;
        }

        .file-item:hover {
            border-color: var(--border-hover);
            background: var(--hover-bg);
        }

        .file-name {
            font-weight: 500;
            color: var(--text-color);
            white-space: nowrap;
            overflow: hidden;
            text-overflow: ellipsis;
            max-width: 170px;
            cursor: pointer;
        }

        .file-actions {
            display: flex;
            align-items: center;
            gap: 12px;
        }

        .file-actions a, .file-actions button {
            background: none;
            border: none;
            color: var(--secondary-text);
            cursor: pointer;
            display: flex;
            align-items: center;
            justify-content: center;
            padding: 4px;
            border-radius: 6px;
            transition: all 0.15s ease;
        }

        .file-actions a:hover, .file-actions button:hover {
            color: var(--text-color);
            background-color: var(--hover-bg);
        }

        .empty-files {
            font-size: 0.8rem;
            color: var(--secondary-text);
            font-style: italic;
            text-align: center;
            padding: 12px 0;
            background: var(--bg-color);
            border-radius: 10px;
            border: 1px dashed var(--border-color);
        }

        .view-all-btn {
            display: flex;
            align-items: center;
            justify-content: center;
            gap: 8px;
            width: 100%;
            padding: 12px;
            border: 1px solid var(--border-color);
            background: var(--container-bg);
            color: var(--text-color);
            border-radius: 10px;
            font-size: 0.85rem;
            font-weight: 700;
            cursor: pointer;
            transition: all 0.2s ease;
            margin-top: auto;
        }

        .view-all-btn:hover {
            background-color: var(--hover-bg);
        }

        /* Pop up menu styling */
        .file-action-menu {
            border: 1px solid var(--border-color) !important;
            box-shadow: var(--card-shadow) !important;
            background: var(--container-bg) !important;
            border-radius: 10px !important;
            overflow: hidden;
        }

        .file-action-menu button {
            color: var(--text-color) !important;
            transition: all 0.15s ease !important;
        }

        .file-action-menu button:hover {
            background: var(--hover-bg) !important;
        }

        footer {
            margin-top: 40px;
            font-size: 0.8rem;
            color: var(--secondary-text);
            text-align: center;
        }

        /* ========================================
           PORTRAIT & MOBILE RESPONSIVE LAYOUT
           ======================================== */
        @media (max-width: 1024px) {

            /* Allow body to scroll on mobile */
            html, body {
                overflow: auto !important;
                height: auto !important;
                min-height: 100% !important;
            }

            /* Show mobile top navigation bar */
            .mobile-topbar {
                display: flex !important;
                align-items: center;
                justify-content: space-between;
                padding: 14px 18px;
                background: var(--container-bg);
                border-bottom: 1px solid var(--border-color);
                position: sticky;
                top: 0;
                z-index: 100;
                box-sizing: border-box;
                width: 100%;
            }

            .mobile-topbar-logo {
                font-family: 'Outfit', sans-serif;
                font-weight: 800;
                font-size: 1.1rem;
                color: var(--text-color);
                display: flex;
                align-items: center;
                gap: 8px;
            }

            .mobile-topbar-actions {
                display: flex;
                gap: 8px;
                align-items: center;
            }

            /* Collapse the desktop 3-column grid to single column */
            .main-layout {
                display: block !important;
                width: 100% !important;
                height: auto !important;
                overflow: visible !important;
            }

            /* Hide desktop sidebars - replaced by mobile drawer */
            .left-sidebar {
                position: fixed !important;
                left: 0 !important;
                top: 0 !important;
                width: min(80vw, 300px) !important;
                height: 100dvh !important;
                z-index: 1000 !important;
                transform: translateX(-110%) !important;
                transition: transform 0.3s cubic-bezier(0.4,0,0.2,1) !important;
                box-shadow: 5px 0 20px rgba(0,0,0,0.2) !important;
                overflow-y: auto !important;
            }

            .left-sidebar.active {
                transform: translateX(0) !important;
            }

            .right-sidebar {
                position: fixed !important;
                right: 0 !important;
                top: 0 !important;
                width: min(85vw, 340px) !important;
                height: 100dvh !important;
                z-index: 1000 !important;
                transform: translateX(110%) !important;
                transition: transform 0.3s cubic-bezier(0.4,0,0.2,1) !important;
                box-shadow: -5px 0 20px rgba(0,0,0,0.2) !important;
                overflow-y: auto !important;
            }

            .right-sidebar.active {
                transform: translateX(0) !important;
            }

            /* Overlay backdrop behind drawers */
            .sidebar-overlay {
                display: none;
                position: fixed;
                inset: 0;
                background: rgba(0,0,0,0.45);
                z-index: 999;
            }
            .sidebar-overlay.active {
                display: block;
            }

            /* Full-width main content, no fixed height */
            .main-content {
                width: 100% !important;
                height: auto !important;
                overflow: visible !important;
                padding: 18px !important;
                box-sizing: border-box !important;
            }

            /* Hide desktop header buttons (moved to mobile topbar) */
            .header {
                display: none !important;
            }

            /* Tabs scrollable horizontally */
            .tabs {
                flex-wrap: nowrap !important;
                overflow-x: auto !important;
                -webkit-overflow-scrolling: touch;
                padding-bottom: 8px !important;
                margin-bottom: 20px !important;
            }

            .tab-btn {
                font-size: 0.78rem !important;
                padding: 8px 14px !important;
                flex-shrink: 0;
            }

            /* Smaller titles */
            h1 {
                font-size: 1.6rem !important;
            }

            .card {
                padding: 18px !important;
                border-radius: 14px !important;
            }

            .features-grid {
                grid-template-columns: repeat(2, 1fr) !important;
            }

            /* Stack API input row vertically */
            .api-input-container {
                flex-direction: column !important;
                align-items: stretch !important;
            }
        }

        @media (max-width: 480px) {
            .features-grid {
                grid-template-columns: 1fr !important;
            }
            .mobile-topbar {
                padding: 12px 14px !important;
            }
            .main-content {
                padding: 14px !important;
            }
            h1 {
                font-size: 1.4rem !important;
            }
        }
    </style>
</head>
<body>
    <!-- MOBILE TOP BAR (hidden on desktop) -->
    <div class="mobile-topbar">
        <div class="mobile-topbar-logo">
            <svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><path d="M13 2L3 14h9l-1 8 10-12h-9l1-8z"/></svg>
            Ruang Kreasi
        </div>
        <div class="mobile-topbar-actions">
            <button type="button" class="theme-toggle-btn" onclick="toggleLeftSidebar()" title="Menu Navigasi">
                <svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><line x1="3" y1="6" x2="21" y2="6"/><line x1="3" y1="12" x2="21" y2="12"/><line x1="3" y1="18" x2="21" y2="18"/></svg>
            </button>
            <button type="button" class="theme-toggle-btn" onclick="toggleRightSidebar()" title="Penyimpanan Server">
                <svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M22 19a2 2 0 0 1-2 2H4a2 2 0 0 1-2-2V5a2 2 0 0 1 2-2h5l2 3h9a2 2 0 0 1 2 2z"/></svg>
            </button>
            <button type="button" class="theme-toggle-btn" onclick="toggleDarkMode()" title="Ganti Tema">
                <svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><circle cx="12" cy="12" r="5"/><line x1="12" y1="1" x2="12" y2="3"/><line x1="12" y1="21" x2="12" y2="23"/><line x1="4.22" y1="4.22" x2="5.64" y2="5.64"/><line x1="18.36" y1="18.36" x2="19.78" y2="19.78"/><line x1="1" y1="12" x2="3" y2="12"/><line x1="21" y1="12" x2="23" y2="12"/><line x1="4.22" y1="19.78" x2="5.64" y2="18.36"/><line x1="18.36" y1="5.64" x2="19.78" y2="4.22"/></svg>
            </button>
        </div>
    </div>

    <!-- SIDEBAR OVERLAY (tap to close on mobile) -->
    <div class="sidebar-overlay" id="sidebar-overlay" onclick="closeMobileDrawers()"></div>

    <div class="main-layout">
        <!-- LEFT SIDEBAR -->
        <div class="left-sidebar">
            <div>
                <div class="logo-area">
                    <svg width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5">
                        <path d="M13 2L3 14h9l-1 8 10-12h-9l1-8z"/>
                    </svg>
                    <span>Ruang Kreasi</span>
                </div>
                
                <nav class="nav-menu">
                    <a class="nav-link active" onclick="switchTab('manual')">
                        <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><rect x="3" y="3" width="7" height="7"/><rect x="14" y="3" width="7" height="7"/><rect x="14" y="14" width="7" height="7"/><rect x="3" y="14" width="7" height="7"/></svg>
                        Dashboard
                    </a>
                    <a class="nav-link" onclick="switchTab('compress')">
                        <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/><polyline points="14 2 14 8 20 8"/></svg>
                        PDF Tools
                    </a>
                    <a class="nav-link" onclick="switchTab('ai')">
                        <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><polygon points="12 2 2 22 22 22"/></svg>
                        AI Tools
                    </a>
                    <a class="nav-link" onclick="alert('Fitur riwayat disinkronkan di server.')">
                        <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><circle cx="12" cy="12" r="10"/><polyline points="12 6 12 12 16 14"/></svg>
                        History
                    </a>
                    <a class="nav-link" onclick="switchTab('compare_docs')">
                        <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M16 4h2a2 2 0 0 1 2 2v14a2 2 0 0 1-2 2H6a2 2 0 0 1-2-2V6a2 2 0 0 1 2-2h2"/><rect x="8" y="2" width="8" height="4" rx="1" ry="1"/></svg>
                        Compare Docs
                    </a>
                    <a class="nav-link" onclick="alert('Anda berada di dasbor penyimpanan server.')">
                        <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M22 19a2 2 0 0 1-2 2H4a2 2 0 0 1-2-2V5a2 2 0 0 1 2-2h5l2 3h9a2 2 0 0 1 2 2z"/></svg>
                        Server Storage
                    </a>
                    <a class="nav-link" onclick="alert('Pengaturan akun disinkronkan otomatis.')">
                        <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><circle cx="12" cy="12" r="3"/><path d="M19.4 15a1.65 1.65 0 0 0 .33 1.82l.06.06a2 2 0 1 1-2.83 2.83l-.06-.06a1.65 1.65 0 0 0-1.82-.33 1.65 1.65 0 0 0-1 1.51V21a2 2 0 0 1-4 0v-.09A1.65 1.65 0 0 0 9 19.4a1.65 1.65 0 0 0-1.82.33l-.06.06a2 2 0 1 1-2.83-2.83l.06-.06a1.65 1.65 0 0 0 .33-1.82 1.65 1.65 0 0 0-1.51-1H3a2 2 0 0 1 0-4h.09A1.65 1.65 0 0 0 4.6 9a1.65 1.65 0 0 0-.33-1.82l-.06-.06a2 2 0 1 1 2.83-2.83l.06.06a1.65 1.65 0 0 0 1.82.33H9a1.65 1.65 0 0 0 1-1.51V3a2 2 0 0 1 4 0v.09a1.65 1.65 0 0 0 1 1.51 1.65 1.65 0 0 0 1.82-.33l.06-.06a2 2 0 1 1 2.83 2.83l-.06.06a1.65 1.65 0 0 0-.33 1.82V9a1.65 1.65 0 0 0 1.51 1H21a2 2 0 0 1 0 4h-.09a1.65 1.65 0 0 0-1.51 1z"/></svg>
                        Settings
                    </a>
                </nav>
            </div>
            
            <div class="sidebar-bottom">
                <div class="user-profile">
                    <div class="avatar">S</div>
                    <div class="user-info">
                        <div class="user-name">Star</div>
                        <div class="user-plan">Premium User</div>
                    </div>
                </div>
                <div class="storage-usage">
                    <div class="storage-text">
                        <span>Penyimpanan</span>
                        <span>32.4 GB / 100 GB</span>
                    </div>
                    <div class="progress-bar">
                        <div class="progress" style="width: 32%;"></div>
                    </div>
                </div>
            </div>
        </div>

        <!-- MAIN CONTENT -->
        <div class="main-content">
            <div class="header">
                <div>
                    <h1>Ruang Kreasi</h1>
                    <p class="subtitle">Alat Produktivitas Dokumen & Bypass Turnitin Premium</p>
                </div>
                <div style="display: flex; gap: 8px; align-items: center; z-index: 100; position: relative;">
                    <!-- Sembunyikan/Tampilkan Panel Kiri -->
                    <button type="button" class="theme-toggle-btn" onclick="toggleLeftSidebar()" title="Tampilkan/Sembunyikan Panel Kiri">
                        <svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><rect x="3" y="3" width="18" height="18" rx="2" ry="2"/><line x1="9" y1="3" x2="9" y2="21"/></svg>
                    </button>
                    <!-- Sembunyikan/Tampilkan Panel Kanan -->
                    <button type="button" class="theme-toggle-btn" onclick="toggleRightSidebar()" title="Tampilkan/Sembunyikan Panel Kanan">
                        <svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><rect x="3" y="3" width="18" height="18" rx="2" ry="2"/><line x1="15" y1="3" x2="15" y2="21"/></svg>
                    </button>
                    <!-- Toggle Tema Dark/Light -->
                    <button type="button" class="theme-toggle-btn" onclick="toggleDarkMode()" title="Ganti Tema">
                        <svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><circle cx="12" cy="12" r="5"/><line x1="12" y1="1" x2="12" y2="3"/><line x1="12" y1="21" x2="12" y2="23"/><line x1="4.22" y1="4.22" x2="5.64" y2="5.64"/><line x1="18.36" y1="18.36" x2="19.78" y2="19.78"/><line x1="1" y1="12" x2="3" y2="12"/><line x1="21" y1="12" x2="23" y2="12"/><line x1="4.22" y1="19.78" x2="5.64" y2="18.36"/><line x1="18.36" y1="5.64" x2="19.78" y2="4.22"/></svg>
                    </button>
                </div>
            </div>

            <!-- Horizontal Tabs -->
            <div class="tabs">
                <button class="tab-btn active" onclick="switchTab('manual')">
                    <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><path d="M12 20h9"/><path d="M16.5 3.5a2.121 2.121 0 0 1 3 3L7 19l-4 1 1-4L16.5 3.5z"/></svg>
                    Turnitin Manual
                </button>
                <button class="tab-btn" onclick="switchTab('ai')">
                    <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><polygon points="12 2 2 22 22 22"/><path d="M12 2v20"/></svg>
                    Turnitin AI Pro
                </button>
                <button class="tab-btn" onclick="switchTab('compress')">
                    <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><rect x="3" y="3" width="18" height="18" rx="2" ry="2"/><line x1="9" y1="9" x2="15" y2="15"/><line x1="15" y1="9" x2="9" y2="15"/></svg>
                    Kompresor PDF
                </button>
                <button class="tab-btn" onclick="switchTab('pdf2word')">
                    <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/><polyline points="14 2 14 8 20 8"/></svg>
                    PDF ke Word
                </button>
                <button class="tab-btn" onclick="switchTab('word2pdf')">
                    <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/><polyline points="14 2 14 8 20 8"/></svg>
                    Word ke PDF
                </button>
                <button class="tab-btn" onclick="switchTab('pdf_solver')">
                    <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><circle cx="12" cy="12" r="10"/><path d="M9.09 9a3 3 0 0 1 5.83 1c0 2-3 3-3 3"/><line x1="12" y1="17" x2="12.01" y2="17"/></svg>
                    PDF Solver (AI)
                </button>
                <button class="tab-btn" onclick="switchTab('compare_docs')">
                    <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><path d="M16 4h2a2 2 0 0 1 2 2v14a2 2 0 0 1-2 2H6a2 2 0 0 1-2-2V6a2 2 0 0 1 2-2h2"/><rect x="8" y="2" width="8" height="4" rx="1" ry="1"/></svg>
                    Bandingkan Dokumen
                </button>
            </div>

            <!-- SUCCESS & ERROR MESSAGES -->
            {% if success_msg %}
                <div class="alert-box">
                    <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><polyline points="20 6 9 17 4 12"/></svg>
                    <span>{{ success_msg }}</span>
                </div>
            {% endif %}

            {% if error_msg %}
                <div class="error-box" style="background: #fdf2f2; border: 1px solid #fde8e8; color: #9b1c1c; border-radius: 10px; padding: 14px; margin-bottom: 25px; font-size: 0.85rem; font-weight: 600; display: flex; align-items: center; gap: 10px;">
                    <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><circle cx="12" cy="12" r="10"/><line x1="12" y1="8" x2="12" y2="12"/><line x1="12" y1="16" x2="12.01" y2="16"/></svg>
                    <span>{{ error_msg }}</span>
                </div>
            {% endif %}

            {% if result_file %}
                <div style="margin-bottom: 35px; text-align: center; background: #eafaf1; padding: 25px; border: 1px solid #d4ebdf; border-radius: 16px;">
                    <p style="margin: 0 0 15px 0; font-size: 0.95rem; font-weight: 600; color: #0f5132;">Proses Selesai! Hasil Anda siap diunduh.</p>
                    <div style="display: flex; gap: 12px; justify-content: center; align-items: center; flex-wrap: wrap;">
                        <a href="/download/{{ result_file }}" class="btn-download">
                            <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"/><polyline points="7 10 12 15 17 10"/><line x1="12" y1="15" x2="12" y2="3"/></svg>
                            Download Hasil
                        </a>
                        {% if orig_file %}
                        <form action="/process_compare_docs" method="POST" style="margin: 0; display: inline;">
                            <input type="hidden" name="server_original_doc" value="{{ orig_file }}">
                            <input type="hidden" name="server_paraphrased_doc" value="{{ result_file }}">
                            <button type="submit" class="btn-download" style="background: #ffffff; color: #171717; border: 1px solid #d4d4d4; cursor: pointer;">
                                <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><path d="M1 12s4-8 11-8 11 8 11 8-4 8-11 8-11-8-11-8z"/><circle cx="12" cy="12" r="3"/></svg>
                                Lihat Preview Perbandingan
                            </button>
                        </form>
                        {% endif %}
                    </div>
                </div>
            {% endif %}

            <!-- TAB CONTAINER CARD -->
            <div class="card">
                <!-- TAB 1: TURNITIN MANUAL -->
                <div id="manual" class="tab-content active">
                    <form action="/process_manual" method="POST" enctype="multipart/form-data">
                        {% if has_active_replacements %}
                        <div style="background: rgba(139, 92, 246, 0.08); border: 1px solid rgba(139, 92, 246, 0.2); padding: 18px; border-radius: 14px; margin-bottom: 25px; display: flex; align-items: center; justify-content: space-between;">
                            <div>
                                <strong style="font-size: 0.95rem; color: #8b5cf6; display: flex; align-items: center; gap: 8px;">
                                    <svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><path d="M13 2L3 14h9l-1 8 10-12h-9l1-8z"/></svg>
                                    Sesi Revisi Aktif Ditemukan
                                </strong>
                                <p style="margin: 6px 0 0 0; font-size: 0.8rem; color: var(--secondary-text);">Ada {{ active_replacements_count }} kata/kalimat hasil parafrase PDF Solver siap diaplikasikan.</p>
                            </div>
                            <label style="display: flex; align-items: center; gap: 10px; cursor: pointer; font-size: 0.85rem; font-weight: 700; color: var(--text-color); padding: 8px 14px; background: var(--hover-bg); border-radius: 10px; border: 1px solid var(--border-color); transition: all 0.2s;">
                                <input type="checkbox" name="use_active_session" id="use_active_session" onchange="toggleManualRefUpload()" style="margin: 0; transform: scale(1.1); accent-color: var(--text-color);"> Gunakan Sesi
                            </label>
                        </div>
                        {% endif %}

                        <div class="form-group">
                            <label>DOKUMEN SKRIPSI ASLI (.DOCX)</label>
                            <div class="file-upload-wrapper">
                                <span class="upload-icon">📄</span>
                                <div class="file-name-label" id="manual-orig-label">Pilih file atau seret ke sini</div>
                                <input type="file" name="original_doc" accept=".docx" required onchange="updateLabel(this, 'manual-orig-label')">
                            </div>
                        </div>

                        <div class="form-group" id="manual_ref_group">
                            <label>DOKUMEN ACUAN PARAFRASE (.DOCX)</label>
                            <div class="file-upload-wrapper">
                                <span class="upload-icon">✍️</span>
                                <div class="file-name-label" id="manual-ref-label">Pilih file atau seret ke sini</div>
                                <input type="file" name="reference_doc" id="manual_ref_input" accept=".docx" required onchange="updateLabel(this, 'manual-ref-label')">
                            </div>
                        </div>

                        <div class="form-group">
                            <label>NAMA FILE OUTPUT <span style="font-weight: normal; color: var(--secondary-text); font-size: 0.8rem; text-transform: none; letter-spacing: 0;">(Opsional - kosongkan untuk nama otomatis)</span></label>
                            <div style="display: flex; align-items: center; gap: 8px;">
                                <input type="text" name="custom_filename" id="manual_custom_filename" placeholder="Contoh: Skripsi_Final_Revisi" style="width: 100%; font-family: inherit;" oninput="sanitizeFilename(this)">
                                <span style="color: var(--secondary-text); font-size: 0.85rem; white-space: nowrap;">.docx</span>
                            </div>
                        </div>

                        <button type="submit" class="btn-submit">Proses & Ganti</button>
                    </form>
                </div>

                <!-- TAB 2: TURNITIN AI PRO -->
                <div id="ai" class="tab-content">
                    <form action="/process_ai" method="POST" enctype="multipart/form-data">
                        <div class="form-group">
                            <label>MESIN PARAFRASE</label>
                            <select name="engine" id="engine_select" onchange="toggleAPIKeyField(); toggleLocalOptions()">
                                <option value="gemini">Gemini AI (Membutuhkan Google Key)</option>
                                <option value="openrouter">OpenRouter AI (Kompatibel Free &amp; Credits Key)</option>
                                <option value="local">Parafrase Lokal (100% Gratis &amp; Offline - Tanpa Key)</option>
                            </select>
                        </div>

                        <div class="form-group" id="local_options_group" style="display: none; background: rgba(16,185,129,0.07); border: 1px solid rgba(16,185,129,0.2); border-radius: 12px; padding: 15px;">
                            <label style="color: #10b981;">⚙️ OPSI LOKAL OFFLINE</label>
                            <div style="display: flex; align-items: center; gap: 12px; margin-top: 8px;">
                                <label style="font-size: 0.85rem; font-weight: 600; color: var(--text-color); white-space: nowrap;">Jumlah Pass:</label>
                                <select name="local_passes" id="local_passes" style="flex: 1; font-family: inherit;">
                                    <option value="1">1x Pass (Cepat)</option>
                                    <option value="2" selected>2x Pass (Direkomendasikan)</option>
                                    <option value="3">3x Pass (Agresif - Lebih Lama)</option>
                                </select>
                            </div>
                            <p style="font-size: 0.78rem; color: var(--secondary-text); margin: 8px 0 0 0;">💡 2-3x pass memproses ulang hasil sehingga similarity bisa turun lebih dalam (target &lt;19%)</p>
                        </div>

                        <div class="form-group" id="gemini_options_group" style="background: rgba(59,130,246,0.07); border: 1px solid rgba(59,130,246,0.2); border-radius: 12px; padding: 15px;">
                            <label style="color: #3b82f6;">🤖 OPSI GEMINI AI</label>
                            <div style="display: flex; align-items: center; gap: 12px; margin-top: 8px;">
                                <label style="font-size: 0.85rem; font-weight: 600; color: var(--text-color); white-space: nowrap;">Model:</label>
                                <select name="gemini_model" id="gemini_model_select" style="flex: 1; font-family: inherit;">
                                    <option value="gemini-2.5-flash-lite" selected>gemini-2.5-flash-lite (Cepat &amp; Hemat)</option>
                                    <option value="gemini-2.5-flash">gemini-2.5-flash (Lebih Kuat)</option>
                                    <option value="gemini-2.0-flash">gemini-2.0-flash</option>
                                    <option value="gemini-1.5-flash">gemini-1.5-flash</option>
                                    <option value="gemini-2.5-pro">gemini-2.5-pro (Paling Kuat)</option>
                                </select>
                            </div>
                            <div style="display: flex; align-items: center; gap: 10px; margin-top: 10px;">
                                <input type="checkbox" name="smart_try" id="smart_try_cb" value="1" checked style="transform: scale(1.2); accent-color: #3b82f6;">
                                <label for="smart_try_cb" style="font-size: 0.85rem; color: var(--text-color); cursor: pointer; margin: 0;">
                                    <strong>Smart Try</strong> — Otomatis coba model berikutnya jika gagal/rate limit
                                </label>
                            </div>
                            <p style="font-size: 0.75rem; color: var(--secondary-text); margin: 6px 0 0 0;">💡 Urutan fallback: 2.5-flash-lite → 2.5-flash → 2.0-flash → 1.5-flash → 2.5-pro</p>
                        </div>

                        <div class="form-group" id="api_key_group">
                            <label>API KEY GEMINI / OPENROUTER</label>
                            <div class="api-input-container">
                                <input type="text" id="api_key_input" name="api_key" placeholder="Masukkan API Key Anda di sini" required value="{{ api_key_val }}">
                                <button type="button" class="btn-check-api" onclick="checkAPIKey()">Cek API</button>
                            </div>
                            <div id="api_status" class="api-status-badge"></div>
                        </div>

                        <div class="form-group">
                            <label>DOKUMEN SKRIPSI UNTUK DIPARAFRASE (.DOCX)</label>
                            <div class="file-upload-wrapper">
                                <span class="upload-icon">📄</span>
                                <div class="file-name-label" id="ai-orig-label">Pilih file atau seret ke sini</div>
                                <input type="file" name="original_doc" accept=".docx" required onchange="updateLabel(this, 'ai-orig-label')">
                            </div>
                        </div>

                        <div class="form-group">
                            <label>LAPORAN PDF TURNITIN (.PDF) <span style="font-weight: normal; color: var(--secondary-text); font-size: 0.8rem; text-transform: none; letter-spacing: 0;">(Opsional - unggah untuk menargetkan bagian yang terdeteksi plagiasi saja)</span></label>
                            <div class="file-upload-wrapper">
                                <span class="upload-icon">📕</span>
                                <div class="file-name-label" id="ai-pdf-label">Pilih file atau seret ke sini</div>
                                <input type="file" name="turnitin_pdf" accept=".pdf" onchange="updateLabel(this, 'ai-pdf-label')">
                            </div>
                        </div>

                        <div class="form-group">
                            <label>PERSENTASE TURNITIN AWAL (%) <span style="font-weight: normal; color: var(--secondary-text); font-size: 0.8rem; text-transform: none; letter-spacing: 0;">(Opsional - jika kosong, dideteksi otomatis dari PDF, default 30%)</span></label>
                            <input type="number" name="initial_similarity" min="0" max="100" step="0.1" placeholder="Contoh: 29.0" style="width: 100%;">
                        </div>

                        <div id="realtime-predictor-card" style="display: none; background: rgba(56, 189, 248, 0.08); border: 1px solid rgba(56, 189, 248, 0.2); border-radius: 10px; padding: 15px; margin: 15px 0; color: var(--text-color); font-family: inherit;">
                            <div style="font-weight: bold; font-size: 1rem; color: #38bdf8; display: flex; align-items: center; gap: 8px; margin-bottom: 8px;">
                                <span>📊</span> Prediksi Turnitin Real-time
                            </div>
                            <div style="display: grid; grid-template-columns: repeat(2, 1fr); gap: 10px; font-size: 0.9rem;">
                                <div>Total Kata Dokumen: <span id="pred-docx-words" style="font-weight: bold; color: var(--text-color);">-</span></div>
                                <div>Persentase Turnitin: <span id="pred-initial-sim" style="font-weight: bold; color: var(--text-color);">-</span></div>
                                <div>Estimasi Kata Plagiat: <span id="pred-plag-words" style="font-weight: bold; color: #f43f5e;">-</span></div>
                                <div>Prediksi Hasil Akhir: <span id="pred-final-sim" style="font-weight: bold; color: #10b981;">-</span></div>
                            </div>
                            <div style="font-size: 0.75rem; color: var(--secondary-text); margin-top: 8px; font-style: italic;">
                                *Prediksi dihitung berdasarkan asumsi parafrase radikal mendetoksifikasi 100% kemiripan yang ditargetkan.
                            </div>
                        </div>

                        <div class="form-group">
                            <label>NAMA FILE OUTPUT <span style="font-weight: normal; color: var(--secondary-text); font-size: 0.8rem; text-transform: none; letter-spacing: 0;">(Opsional - kosongkan untuk nama otomatis)</span></label>
                            <div style="display: flex; align-items: center; gap: 8px;">
                                <input type="text" name="custom_filename" id="ai_custom_filename" placeholder="Contoh: Skripsi_Final_Revisi" style="width: 100%; font-family: inherit;" oninput="sanitizeFilename(this)">
                                <span style="color: var(--secondary-text); font-size: 0.85rem; white-space: nowrap;">.docx</span>
                            </div>
                        </div>

                        <button type="submit" class="btn-submit">Mulai Parafrase AI</button>
                    </form>
                </div>

                <!-- TAB 3: PDF COMPRESSOR -->
                <div id="compress" class="tab-content">
                    <form action="/process_compress" method="POST" enctype="multipart/form-data">
                        <div class="form-group">
                            <label>DOKUMEN PDF UNTIIK DIKOMPRES (.PDF)</label>
                            <div class="file-upload-wrapper">
                                <span class="upload-icon">📕</span>
                                <div class="file-name-label" id="compress-pdf-label">Pilih file atau seret ke sini</div>
                                <input type="file" name="pdf_file" accept=".pdf" required onchange="updateLabel(this, 'compress-pdf-label')">
                            </div>
                        </div>

                        <button type="submit" class="btn-submit">Kompres PDF</button>
                    </form>
                </div>

                <!-- TAB 4: PDF TO WORD -->
                <div id="pdf2word" class="tab-content">
                    <form action="/process_pdf2word" method="POST" enctype="multipart/form-data">
                        <div class="form-group">
                            <label>PILIH FILE PDF UNTUK DIKONVERSI</label>
                            <div class="file-upload-wrapper">
                                <span class="upload-icon">📕</span>
                                <div class="file-name-label" id="convert-pdf-label">Seret & letakkan file PDF Anda di sini atau klik untuk mencari</div>
                                <input type="file" name="pdf_file" accept=".pdf" required onchange="updateLabel(this, 'convert-pdf-label')">
                            </div>
                        </div>

                        <button type="submit" class="btn-submit">Konversi ke Word</button>
                    </form>
                </div>

                <!-- TAB 5: WORD TO PDF -->
                <div id="word2pdf" class="tab-content">
                    <form action="/process_word2pdf" method="POST" enctype="multipart/form-data">
                        <div class="form-group">
                            <label>PILIH FILE WORD UNTUK DIKONVERSI</label>
                            <div class="file-upload-wrapper">
                                <span class="upload-icon">📄</span>
                                <div class="file-name-label" id="word2pdf-doc-label">Seret & letakkan file Word (.docx) Anda di sini atau klik untuk mencari</div>
                                <input type="file" name="doc_file" accept=".docx" required onchange="updateLabel(this, 'word2pdf-doc-label')">
                            </div>
                        </div>

                        <button type="submit" class="btn-submit">Konversi ke PDF</button>
                    </form>
                </div>

                <!-- TAB 6: PDF SOLVER -->
                <div id="pdf_solver" class="tab-content">
                    <form action="/process_pdf_solver" method="POST" enctype="multipart/form-data">
                        <div class="form-group">
                            <label>MESIN SOLVER</label>
                            <select name="engine" id="engine_select_solver" onchange="toggleAPIKeyFieldSolver()">
                                <option value="gemini">Gemini AI (Membutuhkan Google Key)</option>
                                <option value="openrouter">OpenRouter AI (Kompatibel Free & Credits Key)</option>
                            </select>
                        </div>

                        <div class="form-group" id="api_key_group_solver">
                            <label>API KEY GEMINI / OPENROUTER</label>
                            <div class="api-input-container">
                                <input type="text" id="api_key_input_solver" name="api_key" placeholder="Masukkan API Key Anda di sini" required value="{{ api_key_val }}">
                                <button type="button" class="btn-check-api" onclick="checkAPIKeySolver()">Cek API</button>
                            </div>
                            <div id="api_status_solver" class="api-status-badge"></div>
                        </div>

                        <div class="form-group">
                            <label>DOKUMEN LAPORAN PDF TURNITIN (.PDF)</label>
                            <div class="file-upload-wrapper">
                                <span class="upload-icon">📕</span>
                                <div class="file-name-label" id="pdf-solver-label">Pilih file atau seret ke sini</div>
                                <input type="file" name="pdf_file" accept=".pdf" required onchange="updateLabel(this, 'pdf-solver-label')">
                            </div>
                        </div>

                        <button type="submit" class="btn-submit">Mulai Analisis & Ekstraksi Plagiarisme</button>
                    </form>
                </div>

                <!-- TAB 7: COMPARE DOCS -->
                <div id="compare_docs" class="tab-content">
                    <form action="/process_compare_docs" method="POST" enctype="multipart/form-data">
                        <div class="form-group">
                            <label>DOKUMEN ASLI (.DOCX)</label>
                            <div class="file-upload-wrapper">
                                <span class="upload-icon">📄</span>
                                <div class="file-name-label" id="compare-orig-label">Pilih file atau seret ke sini</div>
                                <input type="file" name="original_doc" accept=".docx" required onchange="updateLabel(this, 'compare-orig-label')">
                            </div>
                        </div>

                        <div class="form-group">
                            <label>DOKUMEN PARAFTASE (.DOCX)</label>
                            <div class="file-upload-wrapper">
                                <span class="upload-icon">✍️</span>
                                <div class="file-name-label" id="compare-para-label">Pilih file atau seret ke sini</div>
                                <input type="file" name="paraphrased_doc" accept=".docx" required onchange="updateLabel(this, 'compare-para-label')">
                            </div>
                        </div>

                        <button type="submit" class="btn-submit">Bandingkan Dokumen</button>
                    </form>
                </div>
            </div>

            <!-- PARAPHRASE RULES SECTION -->
            <div class="rules-title">Aturan Parafrase (Aktif & Diperbarui)</div>
            <div class="rules-container">
                <ul>
                    <li><strong>Makna & Substansi (Anti-Halusinasi):</strong> Selalu pertahankan makna asli 100% secara mutlak, jangan menambah informasi baru, jangan membuat asumsi/opini, dan dilarang keras berhalusinasi atau mengarang fakta ilmiah.</li>
                    <li><strong>Struktur & Pola Kalimat:</strong> Parafrase seluruh kalimat secara menyeluruh (bukan sekadar ganti sinonim), variasikan posisi subjek-predikat-objek, konjungsi, serta gunakan pola aktif/pasif secara natural.</li>
                    <li><strong>Preservasi Angka & Istilah Teknis:</strong> Wajib mempertahankan seluruh data numerik, persentase, tanggal, rumus, satuan ukuran, nama algoritma/teori, singkatan resmi, nama variabel, dan istilah ilmiah/teknis.</li>
                    <li><strong>Proteksi Sitasi & Referensi:</strong> Pertahankan secara utuh nama penulis, tahun, format sitasi (APA/MLA/IEEE), nomor referensi, serta dilarang keras memodifikasi daftar pustaka (*bibliography*).</li>
                    <li><strong>Kualitas Bahasa & EYD V:</strong> Tulis ulang secara ringkas dan bebas kata mubazir sesuai kaidah PUEBI/EYD V, serta pastikan hasil akhir terdengar mengalir alami layaknya tulisan manusia (*humanized style*).</li>
                    <li><strong>Integritas Struktur Dokumen:</strong> Menjaga keutuhan tata letak halaman (*layout*), margin, orientasi, tabel, gambar, footnote, hyperlink, penomoran halaman, dan seluruh format Word bawaan dokumen asli.</li>
                </ul>
            </div>

            <!-- 4 FEATURES GRID BELOW CARD -->
            <div class="features-grid">
                <div class="feature-card">
                    <span class="feature-icon">🧠</span>
                    <div class="feature-title">Ditenagai AI</div>
                    <div class="feature-desc">Teknologi kecerdasan buatan tingkat lanjut untuk memproses parafrase alami.</div>
                </div>
                <div class="feature-card">
                    <span class="feature-icon">🛡️</span>
                    <div class="feature-title">Aman & Privat</div>
                    <div class="feature-desc">Dokumen diproses aman di server lokal dan tidak dipublikasikan ke publik.</div>
                </div>
                <div class="feature-card">
                    <span class="feature-icon">⚡</span>
                    <div class="feature-title">Sangat Cepat</div>
                    <div class="feature-desc">Proses konversi dan analisis dokumen selesai dalam hitungan detik.</div>
                </div>
                <div class="feature-card">
                    <span class="feature-icon">💎</span>
                    <div class="feature-title">Kualitas Tinggi</div>
                    <div class="feature-desc">Mempertahankan tata letak asli, sitasi akademik, dan integritas naskah.</div>
                </div>
            </div>

            <footer>
                <p>&copy; 2026 Ruang Kreasi. Dibuat dengan kecintaan terhadap produktivitas akademik.</p>
            </footer>
        </div>

        <!-- RIGHT SIDEBAR (SERVER STORAGE) -->
        <div class="right-sidebar">
            <div class="sidebar-header">
                <h2>Penyimpanan Server</h2>
                <p>Kelola file dokumen di server. Klik nama file untuk melihat opsi penggunaan cepat.</p>
            </div>
            
            <div class="sidebar-section">
                <h3>Word Documents (.docx)</h3>
                <div class="file-list">
                    {% if docx_files %}
                        {% for file in docx_files %}
                            <div class="file-item">
                                <span class="file-name" onclick="showFileMenu(event, '{{ file }}', true)" title="Klik untuk menu aksi">{{ file }}</span>
                                <div class="file-actions">
                                    <a href="/download/{{ file }}" title="Download File">
                                        <svg width="15" height="15" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"/><polyline points="7 10 12 15 17 10"/><line x1="12" y1="15" x2="12" y2="3"/></svg>
                                    </a>
                                </div>
                            </div>
                        {% endfor %}
                    {% else %}
                        <div class="empty-files">Tidak ada file Word</div>
                    {% endif %}
                </div>
            </div>

            <div class="sidebar-section">
                <h3>PDF Documents (.pdf)</h3>
                <div class="file-list">
                    {% if pdf_files %}
                        {% for file in pdf_files %}
                            <div class="file-item">
                                <span class="file-name" onclick="showFileMenu(event, '{{ file }}', false)" title="Klik untuk menu aksi">{{ file }}</span>
                                <div class="file-actions">
                                    <a href="/download/{{ file }}" title="Download File">
                                        <svg width="15" height="15" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"/><polyline points="7 10 12 15 17 10"/><line x1="12" y1="15" x2="12" y2="3"/></svg>
                                    </a>
                                </div>
                            </div>
                        {% endfor %}
                    {% else %}
                        <div class="empty-files">Tidak ada file PDF</div>
                    {% endif %}
                </div>
            </div>
            
            <button class="view-all-btn" onclick="alert('Semua berkas tersimpan aman di direktori uploads/')">
                <span>Lihat Semua File</span>
                <svg width="15" height="15" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><line x1="5" y1="12" x2="19" y2="12"/><polyline points="12 5 19 12 12 19"/></svg>
            </button>
        </div>
    </div>

    <script>
        function updateLabel(input, labelId) {
            const label = document.getElementById(labelId);
            if (input.files && input.files.length > 0) {
                label.innerText = input.files[0].name;
                label.style.color = 'var(--text-color)';
            } else {
                label.innerText = 'Pilih file atau seret ke sini';
                label.style.color = 'var(--secondary-text)';
            }
        }

        function sanitizeFilename(input) {
            // Hanya izinkan huruf, angka, spasi, dash, underscore, titik
            input.value = input.value.replace(/[^\w\-. ]/g, '');
        }

        function switchTab(tabId) {
            document.querySelectorAll('.tab-btn').forEach(btn => btn.classList.remove('active'));
            document.querySelectorAll('.tab-content').forEach(content => content.classList.remove('active'));
            
            const activeBtn = Array.from(document.querySelectorAll('.tab-btn')).find(btn => btn.getAttribute('onclick').includes(tabId));
            if (activeBtn) activeBtn.classList.add('active');
            
            const activeContent = document.getElementById(tabId);
            if (activeContent) activeContent.classList.add('active');

            // Switch navigation sidebar state
            document.querySelectorAll('.nav-link').forEach(link => {
                link.classList.remove('active');
                if (link.getAttribute('onclick').includes(tabId)) {
                    link.classList.add('active');
                }
            });

            // Save tab state in URL hash
            window.location.hash = tabId;
        }

        function toggleAPIKeyField() {
            const engine = document.getElementById('engine_select').value;
            const apiKeyContainer = document.getElementById('api_key_group');
            const apiKeyInput = document.getElementById('api_key_input');
            const geminiOpts = document.getElementById('gemini_options_group');
            if (engine === 'local') {
                apiKeyContainer.style.display = 'none';
                apiKeyInput.removeAttribute('required');
                if (geminiOpts) geminiOpts.style.display = 'none';
            } else {
                apiKeyContainer.style.display = 'block';
                apiKeyInput.setAttribute('required', 'required');
                if (geminiOpts) geminiOpts.style.display = (engine === 'gemini') ? 'block' : 'none';
            }
        }

        function toggleLocalOptions() {
            const engine = document.getElementById('engine_select').value;
            const localGroup = document.getElementById('local_options_group');
            if (localGroup) {
                localGroup.style.display = (engine === 'local') ? 'block' : 'none';
            }
        }

        function toggleAPIKeyFieldSolver() {
            const engine = document.getElementById('engine_select_solver').value;
            const apiKeyContainer = document.getElementById('api_key_group_solver');
            const apiKeyInput = document.getElementById('api_key_input_solver');
            if (engine === 'local') {
                apiKeyContainer.style.display = 'none';
                apiKeyInput.removeAttribute('required');
            } else {
                apiKeyContainer.style.display = 'block';
                apiKeyInput.setAttribute('required', 'required');
            }
        }

        function toggleManualRefUpload() {
            const checkbox = document.getElementById('use_active_session');
            const refGroup = document.getElementById('manual_ref_group');
            const refInput = document.getElementById('manual_ref_input');
            if (checkbox && checkbox.checked) {
                if (refGroup) refGroup.style.display = 'none';
                if (refInput) refInput.removeAttribute('required');
            } else {
                if (refGroup) refGroup.style.display = 'block';
                if (refInput) refInput.setAttribute('required', 'required');
            }
        }

        // Initialize active tab based on Hash or parameter
        window.addEventListener('DOMContentLoaded', () => {
            const urlParams = new URLSearchParams(window.location.search);
            const activeTabParam = urlParams.get('active_tab');
            const hash = window.location.hash.substring(1);
            
            if (activeTabParam) {
                switchTab(activeTabParam);
            } else if (hash) {
                switchTab(hash);
            }
            toggleAPIKeyField();
            toggleAPIKeyFieldSolver();
            toggleManualRefUpload();
        });

        function checkAPIKey() {
            const apiKey = document.getElementById('api_key_input').value.trim();
            const statusDiv = document.getElementById('api_status');
            
            if (!apiKey) {
                statusDiv.style.display = 'block';
            }

            statusDiv.style.display = 'block';
            statusDiv.style.background = 'rgba(255, 255, 255, 0.03)';
            statusDiv.style.borderColor = 'rgba(255, 255, 255, 0.08)';
            statusDiv.style.color = 'var(--secondary-text)';
            statusDiv.innerText = '⌛ Memeriksa status API Key...';

            fetch('/check_api', {
                method: 'POST',
                headers: {
                    'Content-Type': 'application/json',
                },
                body: JSON.stringify({ api_key: apiKey })
            })
            .then(response => response.json())
            .then(data => {
                if (data.status === 'success') {
                    statusDiv.style.background = 'rgba(16, 185, 129, 0.1)';
                    statusDiv.style.borderColor = 'rgba(16, 185, 129, 0.2)';
                    statusDiv.style.color = '#34d399';
                    statusDiv.innerText = '🟢 API Key Aktif & Valid! Siap digunakan.';
                } else {
                    statusDiv.style.background = 'rgba(239, 68, 68, 0.1)';
                    statusDiv.style.borderColor = 'rgba(239, 68, 68, 0.2)';
                    statusDiv.style.color = '#f87171';
                    statusDiv.innerText = '🔴 API Key Tidak Valid: ' + data.message;
                }
            })
            .catch(error => {
                statusDiv.style.background = 'rgba(239, 68, 68, 0.1)';
                statusDiv.style.borderColor = 'rgba(239, 68, 68, 0.2)';
                statusDiv.style.color = '#f87171';
                statusDiv.innerText = '🔴 Gagal terhubung ke server untuk mengecek API Key.';
            });
        }

        function checkAPIKeySolver() {
            const apiKey = document.getElementById('api_key_input_solver').value.trim();
            const statusDiv = document.getElementById('api_status_solver');
            
            if (!apiKey) {
                statusDiv.style.display = 'block';
                statusDiv.style.background = 'rgba(239, 68, 68, 0.1)';
                statusDiv.style.borderColor = 'rgba(239, 68, 68, 0.2)';
                statusDiv.style.color = '#f87171';
                statusDiv.innerText = '⚠️ Silakan masukkan API Key terlebih dahulu!';
                return;
            }

            statusDiv.style.display = 'block';
            statusDiv.style.background = 'rgba(255, 255, 255, 0.03)';
            statusDiv.style.borderColor = 'rgba(255, 255, 255, 0.08)';
            statusDiv.style.color = 'var(--secondary-text)';
            statusDiv.innerText = '⌛ Memeriksa status API Key...';

            fetch('/check_api', {
                method: 'POST',
                headers: {
                    'Content-Type': 'application/json',
                },
                body: JSON.stringify({ api_key: apiKey })
            })
            .then(response => response.json())
            .then(data => {
                if (data.status === 'success') {
                    statusDiv.style.background = 'rgba(16, 185, 129, 0.1)';
                    statusDiv.style.borderColor = 'rgba(16, 185, 129, 0.2)';
                    statusDiv.style.color = '#34d399';
                    statusDiv.innerText = '🟢 API Key Aktif & Valid! Siap digunakan.';
                } else {
                    statusDiv.style.background = 'rgba(239, 68, 68, 0.1)';
                    statusDiv.style.borderColor = 'rgba(239, 68, 68, 0.2)';
                    statusDiv.style.color = '#f87171';
                    statusDiv.innerText = '🔴 API Key Tidak Valid: ' + data.message;
                }
            })
            .catch(error => {
                statusDiv.style.background = 'rgba(239, 68, 68, 0.1)';
                statusDiv.style.borderColor = 'rgba(239, 68, 68, 0.2)';
                statusDiv.style.color = '#f87171';
                statusDiv.innerText = '🔴 Gagal terhubung ke server untuk mengecek API Key.';
            });
        }

        function showFileMenu(e, filename, isDocx) {
            e.stopPropagation();
            const existing = document.querySelector('.file-action-menu');
            if (existing) existing.remove();

            const menu = document.createElement('div');
            menu.className = 'file-action-menu';
            menu.style.position = 'absolute';
            menu.style.background = '#0d1423';
            menu.style.border = '1px solid rgba(255, 255, 255, 0.08)';
            menu.style.borderRadius = '12px';
            menu.style.boxShadow = '0 10px 25px rgba(0,0,0,0.3)';
            menu.style.zIndex = '1000';
            menu.style.padding = '8px 0';
            menu.style.minWidth = '220px';
            menu.style.top = (e.pageY + 10) + 'px';
            menu.style.left = (e.pageX - 180) + 'px';

            let items = [];
            if (isDocx) {
                items = [
                    { text: '📄 Set as Original (Manual)', action: () => selectServerFile('manual', 'original_doc', filename) },
                    { text: '✍️ Set as Reference (Manual)', action: () => selectServerFile('manual', 'reference_doc', filename) },
                    { text: '🤖 Set as Original (AI Pro)', action: () => selectServerFile('ai', 'original_doc', filename) },
                    { text: '📄 Set as Original (Compare)', action: () => selectServerFile('compare_docs', 'original_doc', filename) },
                    { text: '✍️ Set as Paraphrased (Compare)', action: () => selectServerFile('compare_docs', 'paraphrased_doc', filename) },
                    { text: '🔄 Set as Document to Convert (Word to PDF)', action: () => selectServerFile('word2pdf', 'doc_file', filename) }
                ];
            } else {
                items = [
                    { text: '📕 Set as PDF Solver Input', action: () => selectServerFile('pdf_solver', 'pdf_file', filename) },
                    { text: '🔄 Set as PDF to Convert (PDF to Word)', action: () => selectServerFile('pdf2word', 'pdf_file', filename) },
                    { text: '📕 Set as Turnitin PDF Report (AI Pro)', action: () => selectServerFile('ai', 'turnitin_pdf', filename) }
                ];
            }

            items.forEach(item => {
                const btn = document.createElement('button');
                btn.type = 'button';
                btn.style.width = '100%';
                btn.style.padding = '10px 16px';
                btn.style.textAlign = 'left';
                btn.style.background = 'none';
                btn.style.border = 'none';
                btn.style.cursor = 'pointer';
                btn.style.fontSize = '0.8rem';
                btn.style.fontFamily = 'inherit';
                btn.innerText = item.text;
                btn.onmouseenter = () => btn.style.background = 'rgba(255,255,255,0.05)';
                btn.onmouseleave = () => btn.style.background = 'none';
                btn.onclick = () => {
                    item.action();
                    menu.remove();
                };
                menu.appendChild(btn);
            });

            document.body.appendChild(menu);
            
            const closeHandler = () => {
                menu.remove();
                document.removeEventListener('click', closeHandler);
            };
            setTimeout(() => {
                document.addEventListener('click', closeHandler);
            }, 10);
        }

        function selectServerFile(tabId, fieldName, filename) {
            switchTab(tabId);
            
            const form = document.querySelector(`#${tabId} form`);
            if (!form) return;
            
            let hiddenInput = form.querySelector(`input[name="server_${fieldName}"]`);
            if (!hiddenInput) {
                hiddenInput = document.createElement('input');
                hiddenInput.type = 'hidden';
                hiddenInput.name = `server_${fieldName}`;
                form.appendChild(hiddenInput);
            }
            hiddenInput.value = filename;
            
            let labelId = '';
            if (tabId === 'manual' && fieldName === 'original_doc') labelId = 'manual-orig-label';
            if (tabId === 'manual' && fieldName === 'reference_doc') labelId = 'manual-ref-label';
            if (tabId === 'ai' && fieldName === 'original_doc') labelId = 'ai-orig-label';
            if (tabId === 'ai' && fieldName === 'turnitin_pdf') labelId = 'ai-pdf-label';
            if (tabId === 'compare_docs' && fieldName === 'original_doc') labelId = 'compare-orig-label';
            if (tabId === 'compare_docs' && fieldName === 'paraphrased_doc') labelId = 'compare-para-label';
            if (tabId === 'pdf_solver' && fieldName === 'pdf_file') labelId = 'pdf-solver-label';
            if (tabId === 'word2pdf' && fieldName === 'doc_file') labelId = 'word2pdf-doc-label';
            if (tabId === 'pdf2word' && fieldName === 'pdf_file') labelId = 'convert-pdf-label';
            
            if (labelId) {
                const label = document.getElementById(labelId);
                if (label) {
                    label.innerHTML = `
                        <div style="display: flex; align-items: center; justify-content: space-between; width: 100%;">
                            <span style="color: #34d399; font-weight: 600;">📁 Server: ${filename}</span>
                            <button type="button" onclick="resetServerFile('${tabId}', '${fieldName}', '${labelId}', event)" style="background: none; border: none; color: #f87171; font-weight: bold; cursor: pointer; padding: 0 5px; font-size: 1.2rem; line-height: 1;">×</button>
                        </div>
                    `;
                }
            }
            
            const fileInput = form.querySelector(`input[name="${fieldName}"]`);
            if (fileInput) {
                fileInput.removeAttribute('required');
            }
            if (window.updatePrediction) {
                window.updatePrediction();
            }
        }

        window.resetServerFile = function(tabId, fieldName, labelId, event) {
            if (event) {
                event.stopPropagation();
                event.preventDefault();
            }
            const form = document.querySelector(`#${tabId} form`);
            if (!form) return;
            
            const hiddenInput = form.querySelector(`input[name="server_${fieldName}"]`);
            if (hiddenInput) {
                hiddenInput.value = '';
            }
            
            const label = document.getElementById(labelId);
            if (label) {
                label.innerHTML = 'Choose a file or drag it here';
            }
            
            const fileInput = form.querySelector(`input[name="${fieldName}"]`);
            if (fileInput) {
                fileInput.value = '';
                fileInput.setAttribute('required', 'required');
            }
            if (window.updatePrediction) {
                window.updatePrediction();
            }
        }

        function toggleDarkMode() {
            document.body.classList.toggle('dark-mode');
            const isDark = document.body.classList.contains('dark-mode');
            localStorage.setItem('dark-mode', isDark ? 'true' : 'false');
        }

        function toggleLeftSidebar() {
            const sidebar = document.querySelector('.left-sidebar');
            const overlay = document.getElementById('sidebar-overlay');
            if (window.innerWidth <= 1024) {
                const isOpen = sidebar.classList.toggle('active');
                if (overlay) overlay.classList.toggle('active', isOpen);
                // Close right sidebar if open
                document.querySelector('.right-sidebar').classList.remove('active');
            } else {
                sidebar.classList.toggle('collapsed');
                updateLayoutGrid();
            }
        }

        function toggleRightSidebar() {
            const sidebar = document.querySelector('.right-sidebar');
            const overlay = document.getElementById('sidebar-overlay');
            if (window.innerWidth <= 1024) {
                const isOpen = sidebar.classList.toggle('active');
                if (overlay) overlay.classList.toggle('active', isOpen);
                // Close left sidebar if open
                document.querySelector('.left-sidebar').classList.remove('active');
            } else {
                sidebar.classList.toggle('collapsed');
                updateLayoutGrid();
            }
        }

        function closeMobileDrawers() {
            document.querySelector('.left-sidebar').classList.remove('active');
            document.querySelector('.right-sidebar').classList.remove('active');
            const overlay = document.getElementById('sidebar-overlay');
            if (overlay) overlay.classList.remove('active');
        }

        function updateLayoutGrid() {
            if (window.innerWidth <= 1024) return;
            const leftCollapsed = document.querySelector('.left-sidebar').classList.contains('collapsed');
            const rightCollapsed = document.querySelector('.right-sidebar').classList.contains('collapsed');
            const layout = document.querySelector('.main-layout');
            
            let leftCol = leftCollapsed ? '0px' : 'var(--sidebar-width)';
            let rightCol = rightCollapsed ? '0px' : 'var(--storage-width)';
            
            layout.style.gridTemplateColumns = `${leftCol} 1fr ${rightCol}`;
        }

        // Apply dark mode on load if stored
        if (localStorage.getItem('dark-mode') === 'true') {
            document.body.classList.add('dark-mode');
        }

        // Show loading state on form submit
        let loadingInterval = null;
        window.addEventListener('DOMContentLoaded', () => {
            document.querySelectorAll('form').forEach(form => {
                form.addEventListener('submit', function(e) {
                    let valid = true;
                    form.querySelectorAll('[required]').forEach(input => {
                        if (!input.value && !form.querySelector(`input[name="server_${input.name}"]`)?.value) {
                            valid = false;
                        }
                    });
                    
                    if (!valid) return;

                    const overlay = document.getElementById('loading-overlay');
                    const subtitle = document.getElementById('loading-subtitle');
                    if (overlay) {
                        // overlay.classList.add('active'); // DISABLED - USER REQUESTED TO REMOVE LOADING EFFECT
                        
                        const loadingTexts = [
                            "Menghubungkan ke mesin bypass...",
                            "Sedang menyuap robot Turnitin...",
                            "Membaca jampi-jampi bebas plagiasi...",
                            "Membungkus kata-kata biar tidak terendus AI...",
                            "Menghapus jejak kecerdasan buatan...",
                            "Merakit kalimat agar terlihat jenius di mata dosen...",
                            "Menyusun skripsi, menghindari revisi...",
                            "Membuat paragraf Anda tidak terdeteksi oleh radar Turnitin...",
                            "Menguji ketahanan mental dosen penguji...",
                            "Merapikan sitasi, mengamankan wisuda..."
                        ];
                        
                        let idx = 0;
                        if (loadingInterval) clearInterval(loadingInterval);
                        loadingInterval = setInterval(() => {
                            idx = (idx + 1) % loadingTexts.length;
                            subtitle.innerText = loadingTexts[idx];
                        }, 2500);
                    }
                });
            });

            // Real-time predictor logic
            const docInput = document.querySelector('#ai input[name="original_doc"]');
            const pdfInput = document.querySelector('#ai input[name="turnitin_pdf"]');
            const simInput = document.querySelector('#ai input[name="initial_similarity"]');
            const predCard = document.getElementById('realtime-predictor-card');
            
            async function updatePrediction() {
                const formData = new FormData();
                let hasFiles = false;
                
                if (docInput && docInput.files && docInput.files.length > 0) {
                    formData.append('original_doc', docInput.files[0]);
                    hasFiles = true;
                }
                if (pdfInput && pdfInput.files && pdfInput.files.length > 0) {
                    formData.append('turnitin_pdf', pdfInput.files[0]);
                    hasFiles = true;
                }
                if (simInput && simInput.value) {
                    formData.append('initial_similarity', simInput.value);
                }
                
                // Also support selected server files if any, scoped to #ai
                const serverDoc = document.querySelector('#ai input[name="server_original_doc"]');
                const serverPdf = document.querySelector('#ai input[name="server_turnitin_pdf"]');
                if (serverDoc && serverDoc.value) {
                    formData.append('server_original_doc', serverDoc.value);
                    hasFiles = true;
                }
                if (serverPdf && serverPdf.value) {
                    formData.append('server_turnitin_pdf', serverPdf.value);
                    hasFiles = true;
                }
                
                if (!hasFiles) {
                    predCard.style.display = 'none';
                    return;
                }
                
                try {
                    const res = await fetch('/api/predict', {
                        method: 'POST',
                        body: formData
                    });
                    if (res.ok) {
                        const data = await res.json();
                        document.getElementById('pred-docx-words').innerText = data.docx_words.toLocaleString() + ' kata';
                        document.getElementById('pred-initial-sim').innerText = data.initial_similarity + '%';
                        document.getElementById('pred-plag-words').innerText = Math.round(data.plagiarized_words).toLocaleString() + ' kata';
                        document.getElementById('pred-final-sim').innerText = data.predicted_similarity.toFixed(1) + '% (Bebas Plagiasi)';
                        predCard.style.display = 'block';
                        // Auto fill similarity input if detected from PDF
                        if (data.detected_from_pdf && !simInput.value) {
                            simInput.value = data.initial_similarity;
                        }
                    }
                } catch (err) {
                    console.error("Prediction failed:", err);
                }
            }
            
            // Expose globally so selectServerFile and resetServerFile can call it
            window.updatePrediction = updatePrediction;
            
            if (docInput) docInput.addEventListener('change', updatePrediction);
            if (pdfInput) pdfInput.addEventListener('change', updatePrediction);
            if (simInput) simInput.addEventListener('input', updatePrediction);
        });
    </script>

    <!-- COOL FULLSCREEN LOADING OVERLAY -->
    <div id="loading-overlay" class="loading-overlay">
        <div class="loading-container">
            <div class="glow-spinner"></div>
            <div class="loading-title">Ruang Kreasi sedang Bekerja</div>
            <div id="loading-subtitle" class="loading-subtitle">Menghubungkan ke mesin bypass...</div>
        </div>
    </div>
</body>
</html>
"""

def should_keep_replacement(orig, para):
    orig_clean = orig.strip()
    para_clean = para.strip()
    
    # 1. Jika teks asli hanya berupa angka (misal: "2025"), abaikan
    if orig_clean.isdigit():
        return False
        
    # 2. Jika teks asli hanya berupa angka di dalam kurung/siku (misal: "(2025)" atau "[1]"), abaikan
    if re.match(r'^[\(\[\s]*\d+[\)\]\s]*$', orig_clean):
        return False
        
    # 3. Jika teks asli mengandung huruf tapi hasil parafrase sama sekali tidak mengandung huruf (misal: "4. Gizi" -> "4"), abaikan
    has_letters_orig = bool(re.search(r'[a-zA-Z]', orig_clean))
    has_letters_para = bool(re.search(r'[a-zA-Z]', para_clean))
    if has_letters_orig and not has_letters_para:
        return False
        
    # 4. Jika teks terlalu pendek (misal kurang dari 3 karakter), abaikan
    if len(orig_clean) < 3:
        return False
        
    return True

def parse_revisi(path):
    doc = Document(path)
    replacements = []
    current_original = None
    current_paraphrase = None
    
    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
        
        orig_match = re.search(r'\*\*Teks Asli:\*\*\s*(.*)$', txt)
        para_match = re.search(r'\*\*Hasil Parafrase:\*\*\s*(.*)$', txt)
        
        if orig_match:
            val = orig_match.group(1).strip()
            if val.startswith('"') and val.endswith('"'):
                val = val[1:-1]
            elif val.startswith('“') and val.endswith('”'):
                val = val[1:-1]
            elif val.startswith("'") and val.endswith("'"):
                val = val[1:-1]
            current_original = val.strip()
        elif para_match:
            val = para_match.group(1).strip()
            if val.startswith('**'):
                val = val[2:]
            if val.endswith('**'):
                val = val[:-2]
            val = val.strip()
            if val.startswith('"') and val.endswith('"'):
                val = val[1:-1]
            elif val.startswith('“') and val.endswith('”'):
                val = val[1:-1]
            elif val.startswith("'") and val.endswith("'"):
                val = val[1:-1]
            current_paraphrase = val.strip()
            
            if current_original:
                current_paraphrase = current_paraphrase.replace("**", "")
                if should_keep_replacement(current_original, current_paraphrase):
                    replacements.append((current_original, current_paraphrase))
                current_original = None
                current_paraphrase = None
                
    return replacements

def _count_docx_words(docx_path):
    try:
        doc = Document(docx_path)
        count = 0
        for p in doc.paragraphs:
            count += len(re.findall(r'\b\w+\b', p.text))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        count += len(re.findall(r'\b\w+\b', p.text))
        return max(count, 1)
    except Exception as e:
        print(f"Failed to count docx words: {e}")
        return 1

def _extract_pdf_similarity(pdf_path):
    try:
        reader = PdfReader(pdf_path)
        for page in reversed(reader.pages):
            text = page.extract_text() or ""
            patterns = [
                r'(\d+)\s*%\s*(?:SIMILARITY\s+INDEX|originality\s+report|index\s+similaritas)',
                r'(?:SIMILARITY\s+INDEX|originality\s+report|index\s+similaritas)[^\d]{0,20}(\d+)\s*%',
            ]
            for pat in patterns:
                match = re.search(pat, text, re.IGNORECASE)
                if match:
                    val = int(match.group(1))
                    if 0 <= val <= 100:
                        return val
    except Exception as e:
        print(f"Failed to extract Turnitin similarity: {e}")
    return None

def make_fuzzy_regex(orig_text):
    words = re.findall(r'\b\w+\b', orig_text)
    if not words:
        return None
    escaped_words = [re.escape(w) for w in words]
    pattern = r'\b' + r'\b[\s\S]{0,15}?\b'.join(escaped_words) + r'\b'
    try:
        return re.compile(pattern, re.IGNORECASE)
    except:
        return None

def do_smart_replacements(original_doc_path, replacements, output_path):
    doc = Document(original_doc_path)
    
    def clean(text):
        return re.sub(r'\s+', ' ', text).strip()
        
    def safe_replace(p, orig, para):
        # 1. Coba ganti di level run satu per satu untuk menjaga style
        replaced_in_run = False
        for r in p.runs:
            if orig in r.text:
                r.text = r.text.replace(orig, para)
                replaced_in_run = True
                
        if replaced_in_run:
            return True
            
        # 2. Jika gagal (karena teks terpecah di beberapa run), ganti teks paragraf & wariskan style run pertama
        orig_text = p.text
        if orig in orig_text:
            new_text = orig_text.replace(orig, para)
            if p.runs:
                first_run = p.runs[0]
                font_name = first_run.font.name
                font_size = first_run.font.size
                bold = first_run.bold
                italic = first_run.italic
                underline = first_run.underline
                color = first_run.font.color.rgb if (first_run.font and first_run.font.color) else None
                
                first_run.text = new_text
                for r in p.runs[1:]:
                    r.text = ""
                    
                first_run.font.name = font_name
                first_run.font.size = font_size
                first_run.bold = bold
                first_run.italic = italic
                first_run.underline = underline
                if color:
                    first_run.font.color.rgb = color
            else:
                p.text = new_text
            return True
            
        # 3. Cek jika versi clean ada (spasi berlebih)
        orig_clean = clean(orig)
        p_clean = clean(p.text)
        if orig_clean in p_clean:
            words = [re.escape(w) for w in orig_clean.split()]
            pattern = r'\s+'.join(words)
            new_text, count = re.subn(pattern, para, p.text)
            if count > 0:
                if p.runs:
                    first_run = p.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    bold = first_run.bold
                    italic = first_run.italic
                    underline = first_run.underline
                    color = first_run.font.color.rgb if (first_run.font and first_run.font.color) else None
                    
                    first_run.text = new_text
                    for r in p.runs[1:]:
                        r.text = ""
                        
                    first_run.font.name = font_name
                    first_run.font.size = font_size
                    first_run.bold = bold
                    first_run.italic = italic
                    first_run.underline = underline
                    if color:
                        first_run.font.color.rgb = color
                else:
                    p.text = new_text
                return True
        return False

    sorted_reps = sorted(replacements, key=lambda x: len(x[0]), reverse=True)
    replaced_count = 0
    replaced_words_count = 0
    
    for p in doc.paragraphs:
        p_clean = clean(p.text)
        for orig, para in sorted_reps:
            orig_clean = clean(orig)
            matched_str = None
            if orig_clean in p_clean:
                matched_str = orig
            else:
                fuzzy_rx = make_fuzzy_regex(orig)
                if fuzzy_rx:
                    m = fuzzy_rx.search(p.text)
                    if m:
                        matched_str = m.group(0)

            if matched_str:
                orig_words_len = len(re.findall(r'\b\w+\b', orig))
                if len(p_clean) - len(clean(matched_str)) < 10:
                    # Ganti total teks paragraf (tetap warisi format run pertama)
                    if p.runs:
                        first_run = p.runs[0]
                        font_name = first_run.font.name
                        font_size = first_run.font.size
                        bold = first_run.bold
                        italic = first_run.italic
                        underline = first_run.underline
                        color = first_run.font.color.rgb if (first_run.font and first_run.font.color) else None
                        
                        first_run.text = para
                        for r in p.runs[1:]:
                            r.text = ""
                            
                        first_run.font.name = font_name
                        first_run.font.size = font_size
                        first_run.bold = bold
                        first_run.italic = italic
                        first_run.underline = underline
                        if color:
                            first_run.font.color.rgb = color
                    else:
                        p.text = para
                    replaced_count += 1
                    replaced_words_count += orig_words_len
                    p_clean = clean(para)
                else:
                    # Ganti sebagian/substring
                    if safe_replace(p, matched_str, para):
                        replaced_count += 1
                        replaced_words_count += orig_words_len
                        p_clean = clean(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p_clean = clean(p.text)
                    for orig, para in sorted_reps:
                        orig_clean = clean(orig)
                        matched_str = None
                        if orig_clean in p_clean:
                            matched_str = orig
                        else:
                            fuzzy_rx = make_fuzzy_regex(orig)
                            if fuzzy_rx:
                                m = fuzzy_rx.search(p.text)
                                if m:
                                    matched_str = m.group(0)

                        if matched_str:
                            orig_words_len = len(re.findall(r'\b\w+\b', orig))
                            if len(p_clean) - len(clean(matched_str)) < 10:
                                if p.runs:
                                    first_run = p.runs[0]
                                    font_name = first_run.font.name
                                    font_size = first_run.font.size
                                    bold = first_run.bold
                                    italic = first_run.italic
                                    underline = first_run.underline
                                    color = first_run.font.color.rgb if (first_run.font and first_run.font.color) else None
                                    
                                    first_run.text = para
                                    for r in p.runs[1:]:
                                        r.text = ""
                                        
                                    first_run.font.name = font_name
                                    first_run.font.size = font_size
                                    first_run.bold = bold
                                    first_run.italic = italic
                                    first_run.underline = underline
                                    if color:
                                        first_run.font.color.rgb = color
                                else:
                                    p.text = para
                                replaced_count += 1
                                replaced_words_count += orig_words_len
                                p_clean = clean(para)
                            else:
                                if safe_replace(p, matched_str, para):
                                    replaced_count += 1
                                    replaced_words_count += orig_words_len
                                    p_clean = clean(p.text)
                                    
    doc.save(output_path)
    return replaced_count, replaced_words_count

@app.route('/')
def index():
    success_msg = request.args.get('success_msg')
    error_msg = request.args.get('error_msg')
    result_file = request.args.get('result_file')
    orig_file = request.args.get('orig_file', '')
    api_key_val = request.args.get('api_key_val', '')
    
    import json
    active_reps_path = os.path.join(app.config['UPLOAD_FOLDER'], 'active_replacements.json')
    has_active_replacements = False
    active_replacements_count = 0
    if os.path.exists(active_reps_path):
        try:
            with open(active_reps_path, 'r') as f:
                reps = json.load(f)
                if reps:
                    has_active_replacements = True
                    active_replacements_count = len(reps)
        except:
            pass
            
    # List files in upload folder for the side select storage panel
    docx_files = []
    pdf_files = []
    try:
        upload_dir = app.config['UPLOAD_FOLDER']
        if os.path.exists(upload_dir):
            for file in os.listdir(upload_dir):
                if file.startswith('.') or file == 'active_replacements.json':
                    continue
                path = os.path.join(upload_dir, file)
                if os.path.isfile(path):
                    if file.lower().endswith('.docx'):
                        docx_files.append(file)
                    elif file.lower().endswith('.pdf'):
                        pdf_files.append(file)
    except:
        pass

    return render_template_string(HTML_TEMPLATE, 
                                  success_msg=success_msg, 
                                  error_msg=error_msg, 
                                  result_file=result_file, 
                                  orig_file=orig_file,
                                  api_key_val=api_key_val,
                                  has_active_replacements=has_active_replacements,
                                  active_replacements_count=active_replacements_count,
                                  docx_files=docx_files,
                                  pdf_files=pdf_files)


@app.route('/api/predict', methods=['POST'])
def api_predict():
    import tempfile
    
    # Resolve files
    doc_file = request.files.get('original_doc')
    pdf_file = request.files.get('turnitin_pdf')
    server_doc = request.form.get('server_original_doc')
    server_pdf = request.form.get('server_turnitin_pdf')
    initial_sim = request.form.get('initial_similarity')
    
    docx_words = 0
    pdf_similarity = None
    detected_from_pdf = False
    
    # Temporarily save doc to read
    if doc_file and doc_file.filename != '':
        try:
            fd, temp_path = tempfile.mkstemp(suffix='.docx')
            os.close(fd)
            doc_file.save(temp_path)
            docx_words = _count_docx_words(temp_path)
            try: os.remove(temp_path)
            except: pass
        except Exception as e:
            print(f"Error reading docx in prediction: {e}")
    elif server_doc:
        doc_path = os.path.join(app.config['UPLOAD_FOLDER'], server_doc)
        if os.path.exists(doc_path):
            docx_words = _count_docx_words(doc_path)
            
    # Read PDF
    if pdf_file and pdf_file.filename != '':
        try:
            fd, temp_path = tempfile.mkstemp(suffix='.pdf')
            os.close(fd)
            pdf_file.save(temp_path)
            pdf_similarity = _extract_pdf_similarity(temp_path)
            if pdf_similarity is not None:
                detected_from_pdf = True
            try: os.remove(temp_path)
            except: pass
        except Exception as e:
            print(f"Error reading pdf in prediction: {e}")
    elif server_pdf:
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], server_pdf)
        if os.path.exists(pdf_path):
            pdf_similarity = _extract_pdf_similarity(pdf_path)
            if pdf_similarity is not None:
                detected_from_pdf = True
                
    # Fallback/override logic
    initial_sim_val = 30.0
    if initial_sim:
        try:
            initial_sim_val = float(initial_sim)
        except:
            pass
    elif pdf_similarity is not None:
        initial_sim_val = float(pdf_similarity)
        
    plagiarized_words = docx_words * (initial_sim_val / 100.0)
    
    # Real-time prediction assumes complete radical paraphrasing reduces the targeted similarity to 0%
    predicted_similarity = 0.0
    
    return jsonify({
        "docx_words": docx_words,
        "initial_similarity": initial_sim_val,
        "plagiarized_words": plagiarized_words,
        "predicted_similarity": predicted_similarity,
        "detected_from_pdf": detected_from_pdf
    })

@app.route('/check_api', methods=['POST'])
def check_api():
    data = request.get_json()
    api_key = data.get('api_key')
    if not api_key:
        return jsonify({'status': 'error', 'message': 'API Key kosong'})
        
    try:
        if api_key.startswith('sk-or-'):
            import requests
            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            }
            payload = {
                "model": "openrouter/free",
                "messages": [{"role": "user", "content": "Tes"}],
                "max_tokens": 10
            }
            res = requests.post("https://openrouter.ai/api/v1/chat/completions", headers=headers, json=payload)
            if res.status_code == 200:
                return jsonify({'status': 'success'})
            else:
                return jsonify({'status': 'error', 'message': f'OpenRouter Error: {res.text}'})
        else:
            client = genai.Client(api_key=api_key)
            models_to_try = [
                'gemini-2.5-flash-lite', 'gemini-2.5-flash',
                'gemini-2.0-flash', 'gemini-1.5-flash',
                'gemini-1.5-pro', 'gemini-2.5-pro'
            ]
            response = None
            last_err = None

            for model_name in models_to_try:
                try:
                    response = client.models.generate_content(
                        model=model_name,
                        contents="Tes koneksi. Balas dengan kata OK saja."
                    )
                    if response and response.text:
                        break
                except Exception as err:
                    last_err = err

            if response and response.text:
                return jsonify({'status': 'success'})
            else:
                return jsonify({'status': 'error', 'message': f'Respons kosong atau error dari API. Terakhir error: {last_err}'})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


def resolve_file(form_server_param, file_key, allowed_ext):
    server_file = request.form.get(form_server_param)
    if server_file:
        filename = server_file
        path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if not os.path.exists(path):
            raise ValueError(f"File server '{filename}' tidak ditemukan.")
        return filename, path, False
    else:
        if file_key not in request.files:
            raise ValueError(f"Silakan pilih file untuk '{file_key}'.")
        f = request.files[file_key]
        if f.filename == '':
            raise ValueError("Nama file tidak valid.")
        if not f.filename.lower().endswith(allowed_ext):
            raise ValueError(f"Format file harus berupa {allowed_ext}!")
        path = os.path.join(app.config['UPLOAD_FOLDER'], f.filename)
        f.save(path)
        return f.filename, path, True

@app.route('/process_manual', methods=['POST'])
def process_manual():
    use_active_session = request.form.get('use_active_session') == 'on'
    
    try:
        orig_name, orig_path, is_temp_orig = resolve_file('server_original_doc', 'original_doc', '.docx')
    except Exception as e:
        return redirect(url_for('index', error_msg=str(e), active_tab="manual"))
        
    reps = []
    active_reps_path = os.path.join(app.config['UPLOAD_FOLDER'], 'active_replacements.json')
    
    if use_active_session:
        import json
        if os.path.exists(active_reps_path):
            try:
                with open(active_reps_path, 'r') as f:
                    reps = json.load(f)
            except Exception as e:
                return redirect(url_for('index', error_msg=f"Error reading session replacements: {str(e)}", active_tab="manual"))
        if not reps:
            return redirect(url_for('index', error_msg="No active session replacements found.", active_tab="manual"))
    else:
        try:
            ref_name, ref_path, is_temp_ref = resolve_file('server_reference_doc', 'reference_doc', '.docx')
            reps = parse_revisi(ref_path)
            if is_temp_ref:
                try: os.remove(ref_path)
                except: pass
        except Exception as e:
            if is_temp_orig:
                try: os.remove(orig_path)
                except: pass
            return redirect(url_for('index', error_msg=str(e), active_tab="manual"))
            
    if reps:
        try:
            out_filename = "PARAFRASED_" + orig_name
            out_path = os.path.join(app.config['UPLOAD_FOLDER'], out_filename)
            
            replaced_count, _ = do_smart_replacements(orig_path, reps, out_path)
            
            if use_active_session:
                try: os.remove(active_reps_path)
                except: pass
                
            return redirect(url_for('index', success_msg=f"Success! Replaced {replaced_count} matches.", result_file=out_filename, orig_file=orig_name, active_tab="manual"))
        except Exception as e:
            return redirect(url_for('index', error_msg=f"Error occurred during replacement: {str(e)}", active_tab="manual"))
            
    return redirect(url_for('index', error_msg="Unknown error.", active_tab="manual"))

def parse_response_to_replacements(response_text):
    import re
    reps = []
    current_original = None
    for line in response_text.split('\n'):
        txt = line.strip()
        if not txt: continue
        orig_match = re.search(r'\*\*Teks Asli:\*\*\s*(.*)$', txt)
        para_match = re.search(r'\*\*Hasil Parafrase:\*\*\s*(.*)$', txt)
        if orig_match:
            val = orig_match.group(1).strip()
            if val.startswith('"') and val.endswith('"'): val = val[1:-1]
            current_original = val.strip()
        elif para_match:
            val = para_match.group(1).strip()
            if val.startswith('**'): val = val[2:]
            if val.endswith('**'): val = val[:-2]
            val = val.strip()
            if val.startswith('"') and val.endswith('"'): val = val[1:-1]
            current_paraphrase = val.strip()
            if current_original:
                current_paraphrase = current_paraphrase.replace("**", "")
                reps.append([current_original, current_paraphrase])
                current_original = None
    return reps

@app.route('/process_ai', methods=['POST'])
def process_ai():
    engine = request.form.get('engine', 'gemini')
    api_key = request.form.get('api_key')
    
    if engine != 'local' and api_key:
        api_key = api_key.strip()
        if api_key.startswith('sk-or-'):
            engine = 'openrouter'
        elif api_key.startswith('AIzaSy') or len(api_key) == 39:
            engine = 'gemini'
            
    if engine != 'local' and not api_key:
        return redirect(url_for('index', error_msg="API Key is required for AI modes.", active_tab="ai"))
        
    try:
        orig_name, orig_path, is_temp_orig = resolve_file('server_original_doc', 'original_doc', '.docx')
    except Exception as e:
        return redirect(url_for('index', error_msg=str(e), api_key_val=api_key, active_tab="ai"))
        
    has_pdf = False
    pdf_name = None
    pdf_path = None
    is_temp_pdf = False
    
    if request.form.get('server_turnitin_pdf') or ('turnitin_pdf' in request.files and request.files['turnitin_pdf'].filename != ''):
        try:
            pdf_name, pdf_path, is_temp_pdf = resolve_file('server_turnitin_pdf', 'turnitin_pdf', '.pdf')
            has_pdf = True
        except Exception as pdf_err:
            if is_temp_orig:
                try: os.remove(orig_path)
                except: pass
            return redirect(url_for('index', error_msg=f"Error reading PDF Report: {str(pdf_err)}", api_key_val=api_key, active_tab="ai"))
            
    initial_sim = None
    form_sim = request.form.get('initial_similarity')
    if form_sim:
        try:
            initial_sim = float(form_sim)
        except:
            pass
            
    if initial_sim is None and has_pdf:
        initial_sim = _extract_pdf_similarity(pdf_path)
        
    if initial_sim is None:
        initial_sim = 30.0
        
    docx_words = _count_docx_words(orig_path)

    if has_pdf:
        sanitized_path = os.path.join(app.config['UPLOAD_FOLDER'], "SANITIZED_" + pdf_name)
        try:
            reader = PdfReader(pdf_path)
            
            # Cari batasan halaman BAB I sampai Daftar Pustaka
            start_page = 0
            end_page = len(reader.pages) - 1
            
            BAB_START_RE = re.compile(
                r'(BAB\s+[IiVv1-5]+|BAB\s+SATU|BAB\s+DUA|BAB\s+TIGA|BAB\s+EMPAT|BAB\s+LIMA|CHAPTER\s+[IiVv1-5]+)',
                re.IGNORECASE
            )
            BAB_END_RE = re.compile(
                r'(DAFTAR\s+PUSTAKA|DAFTAR\s+ACUAN|REFERENSI|BIBLIOGRAPHY|LAMPIRAN|APPENDIX)',
                re.IGNORECASE
            )
            
            found_start = False
            for idx, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                # Cari BAB I
                if not found_start and BAB_START_RE.search(text):
                    start_page = idx
                    found_start = True
                # Cari Daftar Pustaka
                if found_start and BAB_END_RE.search(text):
                    end_page = idx
                    break
            
            print(f"PDF Scan range detected: page {start_page} to {end_page}")
            
            writer = PdfWriter()
            for idx in range(start_page, end_page + 1):
                writer.add_page(reader.pages[idx])
                
            with open(sanitized_path, "wb") as f:
                writer.write(f)
            os.remove(pdf_path)
            pdf_path = sanitized_path
        except Exception as clean_err:
            print(f"Sanitization / Range limit failed: {clean_err}")
            
        try:
            response_text = None
            last_err = None
            
            # --- Prompt untuk mode Gemini (bisa lihat visual PDF) ---
            gemini_prompt = (
                "Berikut adalah dokumen PDF Turnitin. Silakan analisis HANYA pada bagian teks yang memiliki warna highlight/sorotan plagiasi (plagiarisme).\n\n"
                "PENTING: Parafrase harus dilakukan secara RADIKAL dan SANGAT KREATIF agar lolos dari deteksi Turnitin (kemiripan 0%). Jangan hanya mengganti kata per kata! Struktur kalimat harus dirombak total (misal: aktif menjadi pasif, ubah urutan klausa, gabungkan/pecah kalimat, gunakan sinonim akademis yang tepat), namun makna esensial ilmiah harus tetap sama 100%.\n\n"
                "Anda WAJIB memberikan output dengan format persis seperti ini untuk setiap temuan plagiasi (tanpa kata pengantar, penutup, atau komentar apa pun):\n\n"
                "### Halaman [Nomor Halaman]\n"
                "**Teks Asli:** [Teks asli yang tersorot plagiasi di PDF]\n"
                "**Hasil Parafrase:** [Teks baru hasil parafrase radikal Anda]\n"
            )

            
            if engine == 'local':
                # Mode lokal: ekstrak paragraf dari DOCX skripsi, parafrase multi-pass
                local_passes = max(1, min(3, int(request.form.get('local_passes', 1))))
                import tempfile as _tmp, shutil as _shutil
                
                current_path = orig_path
                temp_pass_paths = []
                
                for pass_num in range(local_passes):
                    local_session = ParaphraseSession()
                    doc_local = Document(current_path)
                    full_text = ""
                    for para_idx, para in enumerate(doc_local.paragraphs):
                        txt = para.text.strip()
                        if len(txt) > 30:  # hanya paragraf substantif
                            if para_idx % 10 == 0 or para_idx == len(doc_local.paragraphs) - 1:
                                print(f"[Local AI] Pass {pass_num+1}/{local_passes} | Processing paragraph {para_idx}/{len(doc_local.paragraphs)}...")
                            para_result = offline_paraphrase(txt, session=local_session)
                            full_text += f"* **Teks Asli:** \"{txt}\"\n* **Hasil Parafrase:** \"**{para_result}**\"\n\n"
                    
                    if pass_num < local_passes - 1:
                        # Simpan hasil pass ini ke file sementara, jadikan input pass berikutnya
                        reps_pass = parse_response_to_replacements(full_text)
                        fd_pass, path_pass = _tmp.mkstemp(suffix='.docx')
                        os.close(fd_pass)
                        do_smart_replacements(current_path, reps_pass, path_pass)
                        temp_pass_paths.append(path_pass)
                        current_path = path_pass
                
                response_text = full_text
                
                # Bersihkan file temp pass
                for tp in temp_pass_paths:
                    try: os.remove(tp)
                    except: pass
                
                # Simpan path yang sudah diproses multi-pass sebagai orig_path untuk do_smart_replacements akhir
                if temp_pass_paths:
                    # Gunakan pass terakhir-1 sebagai source (full_text sudah dari pass terakhir)
                    pass  # current_path sudah dikembalikan ke original, full_text dari pass terakhir
                
            elif engine == 'openrouter':
                # Mode OpenRouter: ekstrak paragraf dari DOCX skripsi, kirim batch ke AI
                doc_or = Document(orig_path)
                docx_paragraphs = []
                for para in doc_or.paragraphs:
                    txt = para.text.strip()
                    if len(txt) > 30:
                        docx_paragraphs.append(txt)
                
                # Kirim dalam batch agar tidak melebihi token limit
                BATCH_SIZE = 20
                full_text = ""
                import requests as req_lib
                headers_or = {
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json"
                }
                or_prompt_template = (
                    "Parafrase setiap kalimat/paragraf Indonesia berikut secara akademik. "
                    "Output HANYA dengan format berikut untuk setiap item (tanpa komentar lain):\n"
                    "**Teks Asli:** [teks asli]\n"
                    "**Hasil Parafrase:** [teks parafrase]\n\n"
                    "Daftar teks yang harus diparafrase:\n"
                )
                for batch_start in range(0, len(docx_paragraphs), BATCH_SIZE):
                    batch = docx_paragraphs[batch_start:batch_start + BATCH_SIZE]
                    batch_text = ""
                    for i, p in enumerate(batch):
                        batch_text += f"{i+1}. {p}\n"
                    payload_or = {
                        "model": "openrouter/free",
                        "messages": [
                            {"role": "system", "content": PARAPHRASE_SYSTEM_PROMPT},
                            {"role": "user", "content": or_prompt_template + batch_text}
                        ],
                        "max_tokens": 4000
                    }
                    res_or = req_lib.post("https://openrouter.ai/api/v1/chat/completions", headers=headers_or, json=payload_or)
                    if res_or.status_code == 200:
                        res_json_or = res_or.json()
                        if "choices" in res_json_or and len(res_json_or["choices"]) > 0:
                            full_text += res_json_or["choices"][0]["message"]["content"] + "\n"
                        elif "error" in res_json_or:
                            raise ValueError(f"OpenRouter Error: {res_json_or['error']}")
                    else:
                        raise ValueError(f"OpenRouter HTTP Error {res_or.status_code}: {res_or.text}")
                response_text = full_text
                
            else:
                # Mode Gemini: kirim PDF Turnitin secara visual (bisa baca highlight warna)
                _client = genai.Client(api_key=api_key)
                uploaded_file = _client.files.upload(
                    file=pdf_path,
                    config={'mime_type': 'application/pdf'}
                )
                # Mode Gemini: kirim PDF Turnitin secara visual (bisa baca highlight warna)
                import time
                for _ in range(30):
                    _f = _client.files.get(name=uploaded_file.name)
                    if _f.state.name == "ACTIVE":
                        uploaded_file = _f
                        break
                    elif _f.state.name == "FAILED":
                        raise ValueError("Gagal memproses file PDF di server Google API.")
                    time.sleep(2)

                # Smart Try: gunakan model pilihan user, fallback ke model lain jika gagal
                selected_model = request.form.get('gemini_model', 'gemini-2.5-flash-lite')
                smart_try = request.form.get('smart_try', '0') == '1'
                
                fallback_models = [
                    'gemini-2.5-flash-lite', 'gemini-2.5-flash',
                    'gemini-2.0-flash', 'gemini-1.5-flash',
                    'gemini-2.5-pro'
                ]
                # Susun urutan: model pilihan dulu, lalu sisanya sebagai fallback
                if smart_try:
                    models_to_try = [selected_model] + [m for m in fallback_models if m != selected_model]
                else:
                    models_to_try = [selected_model]  # Hanya pakai model yang dipilih, tidak fallback
                for model_name in models_to_try:
                    try:
                        response = _client.models.generate_content(
                            model=model_name,
                            config=genai_types.GenerateContentConfig(
                                system_instruction=SYSTEM_INSTRUCTION_PDF_SOLVER
                            ),
                            contents=[uploaded_file, gemini_prompt]
                        )
                        if response and response.text:
                            response_text = response.text
                            break
                    except Exception as model_err:
                        last_err = model_err
                        # Jika rate limit (429), tunggu sesuai saran API lalu lanjut ke model berikutnya
                        err_str = str(model_err)
                        if '429' in err_str or 'RESOURCE_EXHAUSTED' in err_str:
                            import re as _re
                            delay_match = _re.search(r'retry.*?(\d+)s', err_str, _re.IGNORECASE)
                            wait_sec = int(delay_match.group(1)) if delay_match else 10
                            wait_sec = max(2, min(wait_sec, 15)) # Cap sleep to 15s max
                            print(f"Rate limit on {model_name}, waiting {wait_sec}s...")
                            import time as _time; _time.sleep(wait_sec)
                try:
                    _client.files.delete(name=uploaded_file.name)
                except:
                    pass
                    
            if not response_text:
                raise ValueError(f"Paraphrase engine returned empty response or failed. Last error: {last_err}")
                
            reps = []
            current_original = None
            current_paraphrase = None
            for line in response_text.split('\n'):
                txt = line.strip()
                if not txt:
                    continue
                orig_match = re.search(r'(?:\*\*|\*|)?Teks\s+Asli(?:\*\*|\*|)?\s*:\s*(.*)$', txt, re.IGNORECASE)
                para_match = re.search(r'(?:\*\*|\*|)?Hasil\s+Parafrase(?:\*\*|\*|)?\s*:\s*(.*)$', txt, re.IGNORECASE)
                if orig_match:
                    val = orig_match.group(1).strip()
                    if val.startswith('**'): val = val[2:]
                    if val.endswith('**'): val = val[:-2]
                    val = val.strip()
                    if val.startswith('"') and val.endswith('"'): val = val[1:-1]
                    if val.startswith('“') and val.endswith('”'): val = val[1:-1]
                    if val.startswith("'") and val.endswith("'"): val = val[1:-1]
                    current_original = val.strip()
                elif para_match:
                    val = para_match.group(1).strip()
                    if val.startswith('**'): val = val[2:]
                    if val.endswith('**'): val = val[:-2]
                    val = val.strip()
                    if val.startswith('"') and val.endswith('"'): val = val[1:-1]
                    current_paraphrase = val.strip()
                    if current_original:
                        current_paraphrase = current_paraphrase.replace("**", "")
                        if should_keep_replacement(current_original, current_paraphrase):
                            reps.append([current_original, current_paraphrase])
                        current_original = None
                        current_paraphrase = None
                        
            # Write to debug log file
            try:
                debug_path = os.path.join(app.config['UPLOAD_FOLDER'], 'debug_log.txt')
                with open(debug_path, 'w', encoding='utf-8') as df:
                    df.write("=== RAW RESPONSE ===\n")
                    df.write(response_text)
                    df.write("\n\n=== PARSED REPLACEMENTS ===\n")
                    for o, p_item in reps:
                        df.write(f"ORIG: {repr(o)}\nPARA: {repr(p_item)}\n---\n")
            except Exception as debug_err:
                print(f"Failed to write debug log: {debug_err}")
                        
            if is_temp_pdf:
                try: os.remove(pdf_path)
                except: pass
                
            if reps:
                out_filename = "PARAFRASED_" + orig_name
                out_path = os.path.join(app.config['UPLOAD_FOLDER'], out_filename)
                replaced_count, replaced_words_count = do_smart_replacements(orig_path, reps, out_path)
                
                # Jangan hapus berkas asli agar bisa dibandingkan (preview)
                # if is_temp_orig:
                #     try: os.remove(orig_path)
                #     except: pass
                
                # Calculate predicted similarity score reduction
                predicted_sim = max(0.0, initial_sim - (replaced_words_count / docx_words * 100))
                predicted_sim = round(predicted_sim, 1)
                
                success_msg = (
                    f"Targeted AI Paraphrase successful! Replaced {replaced_count} matches. "
                    f"Predicted Turnitin Similarity reduced from {initial_sim}% to {predicted_sim}% "
                    f"(paraphrased {replaced_words_count} of {docx_words} total words)."
                )
                return redirect(url_for('index', success_msg=success_msg, result_file=out_filename, orig_file=orig_name, api_key_val=api_key, active_tab="ai"))
            else:
                raise ValueError("Tidak ada paragraf plagiasi yang berhasil diekstrak dari laporan PDF.")
        except Exception as e:
            if is_temp_pdf:
                try: os.remove(pdf_path)
                except: pass
            if is_temp_orig:
                try: os.remove(orig_path)
                except: pass
            return redirect(url_for('index', error_msg=f"Targeted Paraphrase Error: {str(e)}", api_key_val=api_key, active_tab="ai"))
            
    try:
        try:
            doc = Document(orig_path)
        except Exception as docx_err:
            raise Exception(f"File dokumen tidak bisa dibuka atau rusak: {docx_err}")

        replaced_count = 0

        # ---- Deteksi batas BAB I - BAB V ----
        import re as _re2
        BAB_START_PATTERN = _re2.compile(
            r'^(BAB\s+[IiVv1-5]+|BAB\s+SATU|BAB\s+DUA|BAB\s+TIGA|BAB\s+EMPAT|BAB\s+LIMA'
            r'|CHAPTER\s+[IiVv1-5]+)',
            _re2.IGNORECASE
        )
        BAB_END_PATTERN = _re2.compile(
            r'^(DAFTAR\s+PUSTAKA|DAFTAR\s+ACUAN|REFERENSI|BIBLIOGRAPHY'
            r'|LAMPIRAN|APPENDIX|PENUTUP\s+DOKUMEN)',
            _re2.IGNORECASE
        )
        ROMAN_TO_INT = {'I': 1, 'II': 2, 'III': 3, 'IV': 4, 'V': 5,
                        '1': 1, '2': 2, '3': 3, '4': 4, '5': 5}

        def _bab_number(txt):
            """Extract integer 1-5 from 'BAB X' text, return 0 if not parseable."""
            m = _re2.search(r'(BAB|CHAPTER)\s+([IVX1-5]+)', txt, _re2.IGNORECASE)
            if m:
                token = m.group(2).upper()
                return ROMAN_TO_INT.get(token, 0)
            return 0

        # Cek apakah dokumen memiliki heading BAB
        has_bab_headings = any(BAB_START_PATTERN.match(para.text.strip()) for para in doc.paragraphs)
        inside_bab15 = not has_bab_headings  # Jika tidak ada heading BAB sama sekali, anggap True (proses semua)
        reached_end   = False  # True kalau sudah lewat BAB V

        # Session dengan Consistency Checker + Frequency Penalty + Dynamic Rewrite Level
        doc_session = ParaphraseSession()

        for p in doc.paragraphs:
            original_text = p.text.strip()
            txt_upper = original_text.upper()

            # Cek apakah ini heading BAB baru
            if BAB_START_PATTERN.match(original_text):
                bab_num = _bab_number(original_text)
                if 1 <= bab_num <= 5:
                    inside_bab15 = True
                    reached_end  = False
                else:
                    # BAB VI ke atas → hentikan
                    inside_bab15 = False
                    reached_end  = True
                # Update session aggression untuk BAB ini
                doc_session.set_section(original_text)

            # Cek apakah ini Daftar Pustaka / Lampiran
            if BAB_END_PATTERN.match(original_text):
                inside_bab15 = False
                reached_end  = True
                doc_session.set_section(original_text)  # aggression → 0.0

            # Lewati paragraf jika berada di luar jangkauan scan (sebelum BAB I atau setelah Daftar Pustaka)
            if not inside_bab15 or reached_end:
                continue

            # Skip kalau baris terlalu pendek atau heading bab
            if len(original_text) <= 20 or original_text.isupper():
                continue

            try:
                para_text = None
                last_err = None

                if engine == 'local':
                    para_text = offline_paraphrase(original_text, session=doc_session)

                elif engine == 'openrouter':
                    import requests
                    headers = {
                        "Authorization": f"Bearer {api_key}",
                        "Content-Type": "application/json"
                    }
                    payload = {
                        "model": "openrouter/free",
                        "messages": [
                            {"role": "system", "content": PARAPHRASE_SYSTEM_PROMPT},
                            {"role": "user", "content": original_text}
                        ],
                        "max_tokens": 1000
                    }
                    res = requests.post("https://openrouter.ai/api/v1/chat/completions", headers=headers, json=payload)
                    if res.status_code == 200:
                        res_json = res.json()
                        if "choices" in res_json and len(res_json["choices"]) > 0:
                            para_text = res_json["choices"][0]["message"]["content"].strip()
                        elif "error" in res_json:
                            raise ValueError(f"OpenRouter Error: {res_json['error']}")
                        else:
                            raise ValueError(f"OpenRouter response missing 'choices': {res_json}")
                    else:
                        raise ValueError(f"OpenRouter HTTP Error {res.status_code}: {res.text}")
                else:
                    _client = genai.Client(api_key=api_key)
                    models_to_try = [
                        'gemini-2.0-flash', 'gemini-2.5-flash',
                        'gemini-2.5-flash-lite', 'gemini-1.5-flash',
                        'gemini-1.5-pro'
                    ]
                    for model_name in models_to_try:
                        try:
                            response = _client.models.generate_content(
                                model=model_name,
                                config=genai_types.GenerateContentConfig(
                                    system_instruction=PARAPHRASE_SYSTEM_PROMPT
                                ),
                                contents=original_text
                            )
                            if response and response.text:
                                para_text = response.text.strip()
                                break
                        except Exception as m_err:
                            last_err = m_err
                            # Jika rate limit (429), tunggu sesuai saran API
                            err_str = str(m_err)
                            if '429' in err_str or 'RESOURCE_EXHAUSTED' in err_str:
                                import re as _re
                                delay_match = _re.search(r'retry.*?(\d+)s', err_str, _re.IGNORECASE)
                                wait_sec = int(delay_match.group(1)) if delay_match else 10
                                wait_sec = max(2, min(wait_sec, 15)) # Cap sleep to 15s max
                                print(f"Rate limit on {model_name}, waiting {wait_sec}s...")
                                import time as _time; _time.sleep(wait_sec)

                if not para_text:
                    print(f"Skipping paragraph due to API failure. Last error: {last_err}")
                    continue

                if len(p.runs) > 0:
                    p.runs[0].text = para_text
                    for r in p.runs[1:]:
                        r.text = ""
                else:
                    p.text = para_text
                replaced_count += 1
            except Exception as api_err:
                print(f"API Error at paragraph: {original_text[:50]}... Error: {api_err}")

                    
        out_filename = "AI_PARAFRASED_" + orig_name
        out_path = os.path.join(app.config['UPLOAD_FOLDER'], out_filename)

        # ---- TRIM OUTPUT: hapus paragraf sebelum BAB I ----
        # Temukan index paragraf BAB I pertama
        bab1_idx = None
        for idx, p in enumerate(doc.paragraphs):
            t = p.text.strip()
            if re.match(r'^BAB\s+(I|1|SATU)\b', t, re.IGNORECASE):
                bab1_idx = idx
                break

        # Hapus semua paragraf sebelum BAB I (cover, abstrak, kata pengantar, daftar isi)
        if bab1_idx and bab1_idx > 0:
            paras_to_delete = list(doc.paragraphs[:bab1_idx])
            for p in paras_to_delete:
                try:
                    p._element.getparent().remove(p._element)
                except Exception:
                    pass

        doc.save(out_path)

        # Jangan hapus berkas asli agar bisa dibandingkan (preview)
        # if is_temp_orig:
        #     try:
        #         os.remove(orig_path)
        #     except:
        #         pass

        # Hitung similarity score estimasi
        all_orig_text = " ".join(
            p.text for p in Document(os.path.join(app.config['UPLOAD_FOLDER'], orig_name)).paragraphs
            if len(p.text.strip()) > 20
        ) if orig_name else ""
        all_para_text = " ".join(
            p.text for p in Document(out_path).paragraphs
            if len(p.text.strip()) > 20
        )
        change_score = _estimate_change_score(all_orig_text, all_para_text)

        return redirect(url_for('index',
            success_msg=f"✅ Berhasil memparafrase {replaced_count} paragraf! "
                        f"Estimasi perubahan kata: ~{change_score}% kata baru. "
                        f"Output: BAB I s/d Daftar Pustaka.",
            result_file=out_filename, orig_file=orig_name,
            api_key_val=api_key, active_tab="ai"))

        
    except Exception as e:
        try:
            if os.path.exists(orig_path):
                os.remove(orig_path)
        except:
            pass
        return redirect(url_for('index', error_msg=f"Error occurred: {str(e)}", api_key_val=api_key, active_tab="ai"))

@app.route('/process_compress', methods=['POST'])
def process_compress():
    if 'pdf_file' not in request.files:
        return redirect(url_for('index', error_msg="No file part", active_tab="compress"))
        
    pdf_file = request.files['pdf_file']
    if pdf_file.filename == '':
        return redirect(url_for('index', error_msg="No selected file", active_tab="compress"))
        
    if pdf_file and pdf_file.filename.lower().endswith('.pdf'):
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], pdf_file.filename)
        out_filename = "COMPRESSED_" + pdf_file.filename
        out_path = os.path.join(app.config['UPLOAD_FOLDER'], out_filename)
        
        pdf_file.save(pdf_path)
        
        try:
            reader = PdfReader(pdf_path)
            writer = PdfWriter()
            
            for page in reader.pages:
                new_page = writer.add_page(page)
                new_page.compress_content_streams()
                
            for page in writer.pages:
                for img in page.images:
                    img.replace(img.image, quality=60)
                
            with open(out_path, "wb") as f:
                writer.write(f)
                
            orig_size = os.path.getsize(pdf_path)
            new_size = os.path.getsize(out_path)
            
            try:
                os.remove(pdf_path)
            except:
                pass
                
            reduction = ((orig_size - new_size) / orig_size) * 100
            success_msg = f"Successfully compressed! Size reduced from {orig_size/1024:.1f} KB to {new_size/1024:.1f} KB ({reduction:.1f}% reduction)."
            return redirect(url_for('index', success_msg=success_msg, result_file=out_filename, active_tab="compress"))
        except Exception as e:
            try:
                if os.path.exists(pdf_path):
                    os.remove(pdf_path)
            except:
                pass
            return redirect(url_for('index', error_msg=f"Error occurred: {str(e)}", active_tab="compress"))
            
    return redirect(url_for('index', error_msg="Invalid file format. Please upload a PDF.", active_tab="compress"))

@app.route('/process_pdf2word', methods=['POST'])
def process_pdf2word():
    if 'pdf_file' not in request.files:
        return redirect(url_for('index', error_msg="No file part", active_tab="pdf2word"))
        
    pdf_file = request.files['pdf_file']
    if pdf_file.filename == '':
        return redirect(url_for('index', error_msg="No selected file", active_tab="pdf2word"))
        
    if pdf_file and pdf_file.filename.lower().endswith('.pdf'):
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], pdf_file.filename)
        docx_filename = pdf_file.filename[:-4] + ".docx"
        docx_path = os.path.join(app.config['UPLOAD_FOLDER'], docx_filename)
        
        pdf_file.save(pdf_path)
        
        try:
            from pdf2docx import Converter
            cv = Converter(pdf_path)
            cv.convert(docx_path, start=0, end=None)
            cv.close()
            
            # Post-process the generated docx file to fix tab stops, margins, and layout rules
            try:
                from docx import Document
                import re
                doc = Document(docx_path)
                for p in doc.paragraphs:
                    # Clear tab stops to prevent weird ruler alignments
                    p.paragraph_format.tab_stops.clear()
                    
                    # Clean up run-level tabs and multiple spaces
                    for r in p.runs:
                        if r.text:
                            # Replace tab characters with a single space
                            r.text = r.text.replace('\t', ' ')
                            # Replace multiple spaces with a single space
                            r.text = re.sub(r' +', ' ', r.text)
                doc.save(docx_path)
            except Exception as post_err:
                print(f"Post-processing formatting failed: {post_err}")
                
            try:
                os.remove(pdf_path)
            except:
                pass
                
            return redirect(url_for('index', success_msg="Conversion successful!", result_file=docx_filename, active_tab="pdf2word"))
        except Exception as e:
            try:
                if os.path.exists(pdf_path):
                    os.remove(pdf_path)
            except:
                pass
            return redirect(url_for('index', error_msg=f"Error occurred during conversion: {str(e)}", active_tab="pdf2word"))
            
    return redirect(url_for('index', error_msg="Invalid file format. Please upload a PDF.", active_tab="pdf2word"))

@app.route('/process_word2pdf', methods=['POST'])
def process_word2pdf():
    try:
        docx_name, docx_path, is_temp_docx = resolve_file('server_doc_file', 'doc_file', '.docx')
    except Exception as e:
        return redirect(url_for('index', error_msg=str(e), active_tab="word2pdf"))
        
    pdf_filename = docx_name[:-5] + ".pdf"
    pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], pdf_filename)
    
    try:
        import pythoncom
        pythoncom.CoInitialize()
        
        import win32com.client
        word = win32com.client.Dispatch('Word.Application')
        word.Visible = False
        word.DisplayAlerts = 0
        word.ScreenUpdating = False
        try:
            doc = word.Documents.Open(docx_path)
            doc.SaveAs(pdf_path, FileFormat=17) # 17 is wdFormatPDF
            doc.Close()
        finally:
            word.Quit()
            
        if is_temp_docx:
            try:
                os.remove(docx_path)
            except:
                pass
            
        return redirect(url_for('index', success_msg="Conversion successful!", result_file=pdf_filename, active_tab="word2pdf"))
    except Exception as e:
        if is_temp_docx:
            try:
                if os.path.exists(docx_path):
                    os.remove(docx_path)
            except:
                pass
        return redirect(url_for('index', error_msg=f"Error occurred during conversion: {str(e)}", active_tab="word2pdf"))
            
    return redirect(url_for('index', error_msg="Invalid file format. Please upload a Word Document (.docx).", active_tab="word2pdf"))

@app.route('/process_pdf_solver', methods=['POST'])
def process_pdf_solver():
    engine = request.form.get('engine', 'gemini')
    api_key = request.form.get('api_key')
    
    if engine != 'local' and api_key:
        api_key = api_key.strip()
        if api_key.startswith('sk-or-'):
            engine = 'openrouter'
        elif api_key.startswith('AIzaSy') or len(api_key) == 39:
            engine = 'gemini'
            
    if engine != 'local' and not api_key:
        return redirect(url_for('index', error_msg="API Key is required for AI modes.", active_tab="pdf_solver"))
        
    try:
        pdf_name, pdf_path, is_temp_pdf = resolve_file('server_pdf_file', 'pdf_file', '.pdf')
    except Exception as e:
        return redirect(url_for('index', error_msg=str(e), api_key_val=api_key, active_tab="pdf_solver"))
        
    sanitized_path = os.path.join(app.config['UPLOAD_FOLDER'], "SANITIZED_" + pdf_name)
    try:
        reader = PdfReader(pdf_path)
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        with open(sanitized_path, "wb") as f:
            writer.write(f)
        os.remove(pdf_path)
        pdf_path = sanitized_path
    except Exception as clean_err:
        print(f"Sanitization failed: {clean_err}")
        if os.path.exists(sanitized_path):
            try: os.remove(sanitized_path)
            except: pass
            
    try:
        response_text = None
        last_err = None
        
        prompt = (
            "Berikut adalah dokumen PDF Turnitin. Silakan analisis HANYA pada bagian teks yang memiliki warna highlight/sorotan plagiasi. "
            "Abaikan teks yang bersih (tanpa warna). Hasilkan output berupa daftar teks asli dan hasil parafrasenya sesuai aturan format Markdown yang telah ditetapkan."
        )
        
        if engine == 'local':
            reader = PdfReader(pdf_path)
            full_text = ""
            for idx, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    full_text += f"\n### Halaman {idx+1}\n"
                    for line in text.split('\n'):
                        if line.strip():
                            para = offline_paraphrase(line.strip())
                            full_text += f"* **Teks Asli:** \"{line.strip()}\"\n* **Hasil Parafrase:** \"**{para}**\"\n\n"
            response_text = full_text
            
        elif engine == 'openrouter':
            # Extract text from the PDF file instead of sending it as base64 image_url
            reader = PdfReader(pdf_path)
            pdf_text_content = ""
            for idx, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    pdf_text_content += f"\n### Halaman {idx + 1}\n{page_text}\n"
                    
            import requests
            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            }
            payload = {
                "model": "openrouter/free",
                "messages": [
                    {"role": "system", "content": SYSTEM_INSTRUCTION_PDF_SOLVER},
                    {
                        "role": "user",
                        "content": f"{prompt}\n\nBerikut adalah konten teks dari PDF Turnitin yang harus dianalisis:\n{pdf_text_content}"
                    }
                ],
                "max_tokens": 4000
            }
            res = requests.post("https://openrouter.ai/api/v1/chat/completions", headers=headers, json=payload)
            if res.status_code == 200:
                res_json = res.json()
                if "choices" in res_json and len(res_json["choices"]) > 0:
                    response_text = res_json["choices"][0]["message"]["content"]
                elif "error" in res_json:
                    raise ValueError(f"OpenRouter Error: {res_json['error']}")
                else:
                    raise ValueError(f"OpenRouter response missing 'choices': {res_json}")
            else:
                raise ValueError(f"OpenRouter HTTP Error {res.status_code}: {res.text}")
        else:
            _client = genai.Client(api_key=api_key)
            uploaded_file = _client.files.upload(
                file=pdf_path,
                config={'mime_type': 'application/pdf'}
            )

            import time
            for _ in range(30):
                _f = _client.files.get(name=uploaded_file.name)
                if _f.state.name == "ACTIVE":
                    uploaded_file = _f
                    break
                elif _f.state.name == "FAILED":
                    raise ValueError("Gagal memproses file PDF di server Google API.")
                time.sleep(2)

            models_to_try = [
                'gemini-2.0-flash', 'gemini-2.5-flash',
                'gemini-2.5-flash-lite', 'gemini-1.5-flash',
                'gemini-1.5-pro'
            ]

            for model_name in models_to_try:
                try:
                    response = _client.models.generate_content(
                        model=model_name,
                        config=genai_types.GenerateContentConfig(
                            system_instruction=SYSTEM_INSTRUCTION_PDF_SOLVER
                        ),
                        contents=[uploaded_file, prompt]
                    )
                    if response and response.text:
                        response_text = response.text
                        break
                except Exception as model_err:
                    print(f"Failed using model {model_name}: {model_err}")
                    last_err = model_err
                    # Jika rate limit (429), tunggu sesuai saran API
                    err_str = str(model_err)
                    if '429' in err_str or 'RESOURCE_EXHAUSTED' in err_str:
                        import re as _re
                        delay_match = _re.search(r'retry.*?(\d+)s', err_str, _re.IGNORECASE)
                        wait_sec = int(delay_match.group(1)) if delay_match else 10
                        wait_sec = max(2, min(wait_sec, 15)) # Cap sleep to 15s max
                        print(f"Rate limit on {model_name}, waiting {wait_sec}s...")
                        import time as _time; _time.sleep(wait_sec)

            try:
                _client.files.delete(name=uploaded_file.name)
            except:
                pass
            
        if not response_text:
            raise ValueError(f"Paraphrase engine returned empty response or failed. Last error: {last_err}")
            
        # Parse replacements from response_text and save to session json file
        import json
        try:
            reps = []
            current_original = None
            current_paraphrase = None
            for line in response_text.split('\n'):
                txt = line.strip()
                if not txt:
                    continue
                orig_match = re.search(r'\*\*Teks Asli:\*\*\s*(.*)$', txt)
                para_match = re.search(r'\*\*Hasil Parafrase:\*\*\s*(.*)$', txt)
                if orig_match:
                    val = orig_match.group(1).strip()
                    if val.startswith('"') and val.endswith('"'): val = val[1:-1]
                    current_original = val.strip()
                elif para_match:
                    val = para_match.group(1).strip()
                    if val.startswith('**'): val = val[2:]
                    if val.endswith('**'): val = val[:-2]
                    val = val.strip()
                    if val.startswith('"') and val.endswith('"'): val = val[1:-1]
                    current_paraphrase = val.strip()
                    if current_original:
                        current_paraphrase = current_paraphrase.replace("**", "")
                        reps.append([current_original, current_paraphrase])
                        current_original = None
                        current_paraphrase = None
            
            if reps:
                active_reps_path = os.path.join(app.config['UPLOAD_FOLDER'], 'active_replacements.json')
                with open(active_reps_path, 'w') as f:
                    json.dump(reps, f)
        except Exception as parse_save_err:
            print(f"Failed to save active replacements: {parse_save_err}")
            
        out_filename = "ruang_kreasi_result.docx"
        out_path = os.path.join(app.config['UPLOAD_FOLDER'], out_filename)
        
        doc = Document()
        paragraphs = response_text.split('\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.save(out_path)
        
        if is_temp_pdf:
            try:
                os.remove(pdf_path)
            except:
                pass
            
        return redirect(url_for('index', success_msg="PDF successfully analyzed! Paraphrases generated successfully.", result_file=out_filename, api_key_val=api_key, active_tab="pdf_solver"))
        
    except Exception as e:
        try:
            if is_temp_pdf and os.path.exists(pdf_path):
                os.remove(pdf_path)
        except:
            pass
        return redirect(url_for('index', error_msg=f"Error occurred: {str(e)}", api_key_val=api_key, active_tab="pdf_solver"))

COMPARE_RESULT_TEMPLATE = """
<!DOCTYPE html>
<html lang="id">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Hasil Perbandingan - Ruang Kreasi</title>
    <link href="https://fonts.googleapis.com/css2?family=Outfit:wght@400;600;700;800&family=Plus+Jakarta+Sans:wght@400;500;600;700&display=swap" rel="stylesheet">
    <style>
        :root {
            --bg-color: #fafafa;
            --container-bg: #ffffff;
            --card-bg: #ffffff;
            --text-color: #171717;
            --secondary-text: #737373;
            --accent-gradient: #171717;
            --border-color: #e5e5e5;
            
            --red-highlight: #fdf2f2;
            --red-border: #fde8e8;
            --red-text: #9b1c1c;
            
            --green-highlight: #eafaf1;
            --green-border: #d4ebdf;
            --green-text: #0f5132;
        }

        body {
            font-family: 'Plus Jakarta Sans', sans-serif;
            background-color: var(--bg-color);
            color: var(--text-color);
            margin: 0;
            padding: 40px 20px;
            font-size: 14px;
            min-height: 100vh;
        }

        .header {
            max-width: 1200px;
            margin: 0 auto 35px auto;
            display: flex;
            justify-content: space-between;
            align-items: center;
            border-bottom: 1px solid var(--border-color);
            padding-bottom: 25px;
        }

        .header-info h1 {
            font-family: 'Outfit', sans-serif;
            font-size: 1.8rem;
            margin: 0;
            font-weight: 800;
            color: var(--text-color);
        }

        .header-info p {
            margin: 6px 0 0 0;
            color: var(--secondary-text);
            font-size: 0.9rem;
        }

        .btn-back {
            display: inline-flex;
            align-items: center;
            gap: 8px;
            padding: 10px 20px;
            border: 1px solid #d4d4d4;
            border-radius: 10px;
            background: #ffffff;
            color: var(--text-color);
            font-weight: 600;
            text-decoration: none;
            transition: all 0.2s;
            cursor: pointer;
        }

        .btn-back:hover {
            background: #f5f5f5;
            border-color: #a3a3a3;
        }

        .comparison-container {
            max-width: 1200px;
            margin: 0 auto;
            display: flex;
            flex-direction: column;
            gap: 20px;
        }

        .grid-header {
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 25px;
            font-weight: 700;
            font-size: 0.9rem;
            text-transform: uppercase;
            letter-spacing: 0.08em;
            padding: 10px 15px;
            color: var(--text-color);
            border-bottom: 2px solid var(--border-color);
        }

        @media (max-width: 768px) {
            .grid-header {
                display: none;
            }
            .row {
                grid-template-columns: 1fr !important;
            }
            .header {
                flex-direction: column;
                align-items: flex-start;
                gap: 15px;
            }
        }

        .row {
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 25px;
            background: var(--container-bg);
            border: 1px solid var(--border-color);
            border-radius: 14px;
            padding: 20px;
            box-shadow: 0 1px 3px rgba(0, 0, 0, 0.05);
            transition: all 0.2s ease;
        }

        .row:hover {
            border-color: #a3a3a3;
        }

        .col {
            padding: 15px;
            border-radius: 10px;
            line-height: 1.6;
            white-space: pre-wrap;
            word-break: break-word;
            font-size: 0.95rem;
        }

        .col-orig {
            background: var(--red-highlight);
            color: var(--red-text);
            border: 1px solid var(--red-border);
        }

        .col-para {
            background: var(--green-highlight);
            color: var(--green-text);
            border: 1px solid var(--green-border);
        }

        .para-num {
            font-family: 'Outfit', sans-serif;
            font-size: 0.75rem;
            color: var(--secondary-text);
            font-weight: 700;
            margin-bottom: 8px;
            text-transform: uppercase;
            letter-spacing: 0.05em;
            display: flex;
            align-items: center;
            gap: 6px;
        }

        .no-diff {
            text-align: center;
            padding: 60px;
            background: var(--container-bg);
            border: 1px solid var(--border-color);
            border-radius: 14px;
            font-size: 1.1rem;
            color: var(--secondary-text);
            box-shadow: 0 1px 3px rgba(0, 0, 0, 0.05);
        }
    </style>
</head>
<body>
    <div class="header">
        <div class="header-info">
            <h1>Hasil Perbandingan Dokumen</h1>
            <p>Membandingkan <strong>{{ orig_name }}</strong> (Kiri) dan <strong>{{ para_name }}</strong> (Kanan)</p>
        </div>
        <a href="/" class="btn-back">
            <svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><line x1="19" y1="12" x2="5" y2="12"/><polyline points="12 19 5 12 12 5"/></svg>
            Kembali ke Dashboard
        </a>
    </div>

    <div class="comparison-container">
        {% if diffs %}
            <div class="grid-header">
                <div>Original Text (Skripsi Asli)</div>
                <div>Paraphrased Text (Hasil Parafrase)</div>
            </div>
            
            {% for diff in diffs %}
            <div class="row">
                <div>
                    <div class="para-num">
                        <svg width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><circle cx="12" cy="12" r="10"/><line x1="12" y1="8" x2="12" y2="16"/><line x1="8" y1="12" x2="16" y2="12"/></svg>
                        Paragraf {{ diff.paragraph_num }}
                    </div>
                    <div class="col col-orig">{{ diff.original }}</div>
                </div>
                <div>
                    <div class="para-num">
                        <svg width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><polyline points="20 6 9 17 4 12"/></svg>
                        Paragraf {{ diff.paragraph_num }}
                    </div>
                    <div class="col col-para">{{ diff.paraphrased }}</div>
                </div>
            </div>
            {% endfor %}
        {% else %}
            <div class="no-diff">
                🎉 Dokumen identik! Tidak ditemukan adanya perbedaan kata/kalimat di antara kedua file tersebut.
            </div>
        {% endif %}
    </div>
</body>
</html>
"""

@app.route('/process_compare_docs', methods=['POST'])
def process_compare_docs():
    try:
        orig_name, orig_path, is_temp_orig = resolve_file('server_original_doc', 'original_doc', '.docx')
    except Exception as e:
        return redirect(url_for('index', error_msg=str(e), active_tab="compare_docs"))
        
    try:
        para_name, para_path, is_temp_para = resolve_file('server_paraphrased_doc', 'paraphrased_doc', '.docx')
    except Exception as e:
        if is_temp_orig:
            try: os.remove(orig_path)
            except: pass
        return redirect(url_for('index', error_msg=str(e), active_tab="compare_docs"))
        
    try:
        doc_orig = Document(orig_path)
        doc_para = Document(para_path)
        
        diffs = []
        for idx, (p_orig, p_para) in enumerate(zip(doc_orig.paragraphs, doc_para.paragraphs)):
            txt_orig = p_orig.text.strip()
            txt_para = p_para.text.strip()
            if txt_orig != txt_para:
                diffs.append({
                    'paragraph_num': idx + 1,
                    'original': txt_orig,
                    'paraphrased': txt_para
                })
                
        if is_temp_orig:
            try: os.remove(orig_path)
            except: pass
        if is_temp_para:
            try: os.remove(para_path)
            except: pass
            
        return render_template_string(COMPARE_RESULT_TEMPLATE, diffs=diffs, orig_name=orig_name, para_name=para_name)
    except Exception as e:
        if is_temp_orig:
            try: os.remove(orig_path)
            except: pass
        if is_temp_para:
            try: os.remove(para_path)
            except: pass
        return redirect(url_for('index', error_msg=f"Error occurred during comparison: {str(e)}", active_tab="compare_docs"))

@app.route('/download/<filename>')
def download(filename):
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    return redirect(url_for('index', error_msg="File not found."))

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(debug=False, host='0.0.0.0', port=port)
